import sys
import os
import re
import tempfile
import json
import gc
import shutil
from datetime import datetime
import qtawesome as qta
import requests

from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout,
                             QHBoxLayout, QLabel, QLineEdit, QPushButton,
                             QFileDialog, QMessageBox, QProgressBar,
                             QScrollArea, QFrame, QDialog,
                             QListWidget, QListWidgetItem,
                             QGridLayout, QGroupBox, QStackedWidget,
                             QSplitter, QTextEdit, QComboBox,
                             QSizePolicy, QGraphicsDropShadowEffect)
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QTimer, QPropertyAnimation, QEasingCurve, QPoint, QRegularExpression
from PyQt6.QtGui import (QPalette, QColor, QShortcut, QKeySequence, QFont,
                         QPixmap, QDoubleValidator, QRegularExpressionValidator)

import matplotlib

matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.backends.backend_qtagg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sklearn.linear_model import LinearRegression
import numpy as np

# ============================================================
# ДИРЕКТОРИЯ
# ============================================================
if getattr(sys, 'frozen', False):
    BASE_DIR = os.path.dirname(sys.executable)
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

PATIENTS_DIR = os.path.join(BASE_DIR, "patients")
os.makedirs(PATIENTS_DIR, exist_ok=True)

# ============================================================
# НАЗВАНИЯ КАНАЛОВ
# ============================================================
CHANNEL_LABELS = {
    "angles": {
        "default": ["Канал 1", "Канал 2", "Канал 3", "Канал 4"],
        "ПОВОРОТ СТОПЫ": ["Угол левой стопы", "Угол правой стопы"],
        "ПОВОРОТ ГОЛЕНИ": ["Угол левой голени", "Угол правой голени"],
        "ПОВОРОТ БЕДРА": ["Угол левого бедра", "Угол левой голени", "Угол правого бедра", "Угол правой голени"],
        "ПРИСЕДАНИЯ": ["Угол левого бедра", "Угол левой голени", "Угол левой стопы", "Угол правого бедра",
                       "Угол правой голени", "Угол правой стопы"],
        "ХОДЬБА": ["Угол левого бедра", "Угол левой голени", "Угол левой стопы", "Угол правого бедра",
                   "Угол правой голени", "Угол правой стопы"],
    },
    "forces": {
        "default": ["Канал 1", "Канал 2", "Канал 3", "Канал 4"],
        "ПОВОРОТ СТОПЫ": [
            "Сила на левом носке", "Сила на левой пятке", "Сила на левой голени",
            "Сила на правом носке", "Сила на правой пятке", "Сила на правой голени"
        ],
        "ПОВОРОТ ГОЛЕНИ": [
            "Сила на левом бедре", "Сила на левой голени",
            "Сила на правом бедре", "Сила на правой голени"
        ],
        "ПОВОРОТ БЕДРА": [
            "Сила на левом бедре", "Сила на левой голени",
            "Сила на правом бедре", "Сила на правой голени"
        ],
        "ПРИСЕДАНИЯ": [
            "Сила на левом бедре", "Сила на левом носке", "Сила на левой пятке", "Сила на левой голени",
            "Сила на правом бедре", "Сила на правом носке", "Сила на правой пятке", "Сила на правой голени"
        ],
        "ХОДЬБА": [
            "Сила на левом бедре", "Сила на левом носке", "Сила на левой пятке", "Сила на левой голени",
            "Сила на правом бедре", "Сила на правом носке", "Сила на правой пятке", "Сила на правой голени"
        ],
    }
}


def get_channel_label(exercise_name: str, channel_idx: int, is_angle: bool = True) -> str:
    key = "angles" if is_angle else "forces"
    name_upper = exercise_name.upper()

    if "БЕДРО" in name_upper and key == "angles":
        labels = CHANNEL_LABELS[key]["ПОВОРОТ БЕДРА"]
        return labels[channel_idx - 1] if 0 <= channel_idx - 1 < len(labels) else f"Канал {channel_idx}"

    for keyword, labels in CHANNEL_LABELS[key].items():
        if keyword in name_upper:
            if 0 <= channel_idx - 1 < len(labels):
                return labels[channel_idx - 1]
            break

    default = CHANNEL_LABELS[key]["default"]
    return default[channel_idx - 1] if 0 <= channel_idx - 1 < len(default) else f"Канал {channel_idx}"


# ============================================================
# ФУНКЦИЯ ПОСТРОЕНИЯ ГРАФИКОВ
# ============================================================
def save_graphs_for_exercise(exercise, output_dir):
    measurements = exercise['measurements']
    if not measurements:
        return None, None

    def parse_time_to_seconds(time_str):
        dt = datetime.strptime(time_str, '%Y-%m-%d %H:%M:%S.%f')
        return dt.hour * 3600 + dt.minute * 60 + dt.second + dt.microsecond / 1e6

    first_time_sec = parse_time_to_seconds(measurements[0][0])
    times = [parse_time_to_seconds(tm) - first_time_sec for tm, _, _ in measurements]

    n_angles = len(measurements[0][1])
    n_forces = len(measurements[0][2])

    angles_by_channel = [[] for _ in range(n_angles)]
    forces_by_channel = [[] for _ in range(n_forces)]

    for _, angles, forces in measurements:
        for ch in range(n_angles):
            angles_by_channel[ch].append(angles[ch])
        for ch in range(n_forces):
            forces_by_channel[ch].append(forces[ch] * 9.81 / 1000)

    angle_graphs = []
    force_graphs = []

    # Графики углов
    for ch in range(n_angles):
        label = get_channel_label(exercise['name'], ch + 1, True)

        fig, ax = plt.subplots(figsize=(12, 5), dpi=150)
        ax.plot(times, angles_by_channel[ch], linewidth=2.5, color='#4a9eff')
        ax.set_ylabel(label + " (градусы)", fontsize=12, fontweight='bold')
        ax.set_xlabel('Время (секунды)', fontsize=11)
        ax.grid(True, linestyle='--', alpha=0.7)
        angle_png = os.path.join(output_dir, f'angle_ch_{ch + 1}.png')
        plt.savefig(angle_png, dpi=150, bbox_inches='tight')
        plt.close(fig)
        angle_graphs.append((label, angle_png))
        gc.collect()

    # Графики сил
    exercise_upper = exercise['name'].upper()

    if "ПОВОРОТ СТОПЫ" in exercise_upper:
        # Левая стопа: носок + пятка на одном графике
        if n_forces >= 2:
            fig, ax = plt.subplots(figsize=(12, 5), dpi=150)
            ax.plot(times, forces_by_channel[0], linewidth=2.5, color='#ff6b6b', label='Носок')
            ax.plot(times, forces_by_channel[1], linewidth=2.5, color='#ffaa44', label='Пятка')
            ax.set_ylabel('Сила (Н)', fontsize=12, fontweight='bold')
            ax.set_xlabel('Время (секунды)', fontsize=11)
            ax.grid(True, linestyle='--', alpha=0.7)
            ax.legend()
            force_png = os.path.join(output_dir, f'force_left_foot.png')
            plt.savefig(force_png, dpi=150, bbox_inches='tight')
            plt.close(fig)
            force_graphs.append(("Левая стопа (носок/пятка)", force_png))
            gc.collect()

        # Левая голень
        if n_forces >= 3:
            fig, ax = plt.subplots(figsize=(12, 5), dpi=150)
            ax.plot(times, forces_by_channel[2], linewidth=2.5, color='#ff6b6b')
            ax.set_ylabel('Сила (Н)', fontsize=12, fontweight='bold')
            ax.set_xlabel('Время (секунды)', fontsize=11)
            ax.grid(True, linestyle='--', alpha=0.7)
            force_png = os.path.join(output_dir, f'force_left_shin.png')
            plt.savefig(force_png, dpi=150, bbox_inches='tight')
            plt.close(fig)
            force_graphs.append(("Левая голень", force_png))
            gc.collect()

        # Правая стопа
        if n_forces >= 5:
            fig, ax = plt.subplots(figsize=(12, 5), dpi=150)
            ax.plot(times, forces_by_channel[3], linewidth=2.5, color='#ff6b6b', label='Носок')
            ax.plot(times, forces_by_channel[4], linewidth=2.5, color='#ffaa44', label='Пятка')
            ax.set_ylabel('Сила (Н)', fontsize=12, fontweight='bold')
            ax.set_xlabel('Время (секунды)', fontsize=11)
            ax.grid(True, linestyle='--', alpha=0.7)
            ax.legend()
            force_png = os.path.join(output_dir, f'force_right_foot.png')
            plt.savefig(force_png, dpi=150, bbox_inches='tight')
            plt.close(fig)
            force_graphs.append(("Правая стопа (носок/пятка)", force_png))
            gc.collect()

        # Правая голень
        if n_forces >= 6:
            fig, ax = plt.subplots(figsize=(12, 5), dpi=150)
            ax.plot(times, forces_by_channel[5], linewidth=2.5, color='#ff6b6b')
            ax.set_ylabel('Сила (Н)', fontsize=12, fontweight='bold')
            ax.set_xlabel('Время (секунды)', fontsize=11)
            ax.grid(True, linestyle='--', alpha=0.7)
            force_png = os.path.join(output_dir, f'force_right_shin.png')
            plt.savefig(force_png, dpi=150, bbox_inches='tight')
            plt.close(fig)
            force_graphs.append(("Правая голень", force_png))
            gc.collect()

    else:
        for ch in range(n_forces):
            label = get_channel_label(exercise['name'], ch + 1, False)

            fig, ax = plt.subplots(figsize=(12, 5), dpi=150)
            ax.plot(times, forces_by_channel[ch], linewidth=2.5, color='#ff6b6b')

            max_val = max(forces_by_channel[ch])
            min_val = min(forces_by_channel[ch])
            max_idx = forces_by_channel[ch].index(max_val)
            min_idx = forces_by_channel[ch].index(min_val)
            ax.plot(times[max_idx], max_val, 'o', color='lime', markersize=8)
            ax.plot(times[min_idx], min_val, 'o', color='red', markersize=8)
            ax.annotate(f'{max_val:.1f} Н', xy=(times[max_idx], max_val), xytext=(0, 8),
                        textcoords='offset points', ha='center', fontsize=9, color='lime')
            ax.annotate(f'{min_val:.1f} Н', xy=(times[min_idx], min_val), xytext=(0, -15),
                        textcoords='offset points', ha='center', fontsize=9, color='red')

            ax.set_ylabel(label + " (Н)", fontsize=12, fontweight='bold')
            ax.set_xlabel('Время (секунды)', fontsize=11)
            ax.grid(True, linestyle='--', alpha=0.7)

            force_png = os.path.join(output_dir, f'force_ch_{ch + 1}.png')
            plt.savefig(force_png, dpi=150, bbox_inches='tight')
            plt.close(fig)
            force_graphs.append((label, force_png))
            gc.collect()

    graphs_info = {
        'angles': angle_graphs,
        'forces': force_graphs
    }
    with open(os.path.join(output_dir, 'graphs_info.json'), 'w', encoding='utf-8') as f:
        json.dump(graphs_info, f, ensure_ascii=False, indent=2)

    all_forces = [value for channel in forces_by_channel for value in channel]
    max_force = max(all_forces) if all_forces else 0
    with open(os.path.join(output_dir, 'max_force.txt'), 'w') as f:
        f.write(str(max_force))

    # Для обратной совместимости
    if n_angles > 0:
        fig, axes = plt.subplots(n_angles, 1, figsize=(12, 3.5 * n_angles), sharex=True, dpi=150)
        if n_angles == 1:
            axes = [axes]
        for ch, ax in enumerate(axes):
            ax.plot(times, angles_by_channel[ch], linewidth=2.5, color='#4a9eff')
            ax.set_ylabel(get_channel_label(exercise['name'], ch + 1, True), fontsize=10)
            ax.grid(True, linestyle='--', alpha=0.7)
        axes[-1].set_xlabel('Время (секунды)', fontsize=11)
        plt.tight_layout()
        angles_png = os.path.join(output_dir, 'angles.png')
        plt.savefig(angles_png, dpi=150, bbox_inches='tight')
        plt.close(fig)
    else:
        angles_png = None

    if n_forces > 0:
        fig, axes = plt.subplots(n_forces, 1, figsize=(12, 3.5 * n_forces), sharex=True, dpi=150)
        if n_forces == 1:
            axes = [axes]
        for ch, ax in enumerate(axes):
            ax.plot(times, forces_by_channel[ch], linewidth=2.5, color='#ff6b6b')
            ax.set_ylabel(get_channel_label(exercise['name'], ch + 1, False), fontsize=10)
            ax.grid(True, linestyle='--', alpha=0.7)
        axes[-1].set_xlabel('Время (секунды)', fontsize=11)
        plt.tight_layout()
        forces_png = os.path.join(output_dir, 'forces.png')
        plt.savefig(forces_png, dpi=150, bbox_inches='tight')
        plt.close(fig)
    else:
        forces_png = None

    gc.collect()
    return angles_png, forces_png


# ============================================================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ И КЛАССЫ
# ============================================================
def docx_to_txt(docx_path):
    doc = Document(docx_path)
    text = '\n'.join([para.text for para in doc.paragraphs])
    fd, temp_path = tempfile.mkstemp(suffix='.txt', prefix='rehab_', text=True)
    with os.fdopen(fd, 'w', encoding='utf-8') as f:
        f.write(text)
    return temp_path


class RobustParser:
    @staticmethod
    def parse(txt_path):
        content = None
        for enc in ['utf-8', 'cp1251', 'latin-1']:
            try:
                with open(txt_path, 'r', encoding=enc) as f:
                    content = f.read()
                break
            except:
                continue
        if content is None:
            return []

        pattern = re.compile(
            r'(?P<time>\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}:\d{2}(?:\.\d+)?)\s+'
            r'(?P<angles>[-\d,;.]+)\s+'
            r'(?P<forces>[-\d,;.]+)'
        )
        all_measurements = []
        for match in pattern.finditer(content):
            time_str = match.group('time')
            angles = [float(x) for x in re.split(r'[,;]+', match.group('angles')) if x.strip()]
            forces = [float(x) for x in re.split(r'[,;]+', match.group('forces')) if x.strip()]
            if angles and forces:
                all_measurements.append((time_str, angles, forces, match.start()))

        if not all_measurements:
            return []

        title_spans = []
        for match in re.finditer(r'\*?\*?===\s*(.+?)\s*===\*?\*?', content):
            title_spans.append((match.start(), match.group(1).strip()))
        title_spans.sort(key=lambda x: x[0])
        title_spans.append((len(content), None))

        exercises = []
        for i in range(len(title_spans) - 1):
            start_pos, title = title_spans[i]
            end_pos = title_spans[i + 1][0]
            if title is None:
                continue
            block = content[start_pos:end_pos]

            start_time_val = None
            params = {}
            start_match = re.search(r'\*?Начало:\s*(.+?)(?=\n|$)', block, re.IGNORECASE)
            if start_match:
                param_line = start_match.group(1)
                date_match = re.search(r'(\d{4}-\d{2}-\d{2})', param_line)
                time_match = re.search(r'(\d{2}:\d{2}:\d{2})', param_line)
                if date_match and time_match:
                    start_time_val = f"{date_match.group(1)} {time_match.group(1)}"
                elif date_match:
                    start_time_val = date_match.group(1)
                for part in param_line.split('|'):
                    if ':' in part:
                        k, v = part.split(':', 1)
                        params[k.strip()] = v.strip()

            block_meas = [m for m in all_measurements if start_pos <= m[3] < end_pos]
            if not block_meas:
                continue

            if start_time_val:
                safe_start = start_time_val.replace(':', '-').replace(' ', '_')
            else:
                safe_start = datetime.now().strftime('%Y%m%d_%H%M%S')
            folder_name = f"{title}_{safe_start}"
            exercises.append({
                'name': title,
                'start_time': start_time_val,
                'params': params,
                'measurements': [(t, a, f) for t, a, f, _ in block_meas],
                'folder_name': folder_name
            })
        return exercises


def create_exercise_report(exercise, output_dir, angles_png, forces_png):
    doc = Document()
    title = doc.add_heading(f'Отчёт по упражнению: {exercise["name"]}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Время начала: {exercise["start_time"]}')
    if exercise['params']:
        doc.add_heading('Параметры', level=2)
        for k, v in exercise['params'].items():
            doc.add_paragraph(f'{k}: {v}')
    doc.add_paragraph(f'Общее количество измерений: {len(exercise["measurements"])}')
    doc.add_heading('Графики', level=2)
    doc.add_paragraph('Углы по каналам:')
    doc.add_picture(angles_png, width=Inches(6.0))
    doc.add_paragraph('Силы по каналам:')
    doc.add_picture(forces_png, width=Inches(6.0))
    report_path = os.path.join(output_dir, 'report.docx')
    doc.save(report_path)
    return report_path


def create_full_report(patient_name, patient_dir):
    exercises_amp_data = {}
    exercises_info = []

    for folder in os.listdir(patient_dir):
        folder_path = os.path.join(patient_dir, folder)
        if not os.path.isdir(folder_path):
            continue
        if os.path.exists(os.path.join(folder_path, 'angles.png')):
            ex_name = folder.split('_', 1)[0] if '_' in folder else folder
            exercises_info.append({
                'name': ex_name,
                'angles_png': os.path.join(folder_path, 'angles.png'),
                'forces_png': os.path.join(folder_path, 'forces.png')
            })
            max_force_path = os.path.join(folder_path, 'max_force.txt')
            if os.path.exists(max_force_path):
                with open(max_force_path, 'r') as f:
                    max_force = float(f.read().strip())
                date_str = ""
                for part in folder.split('_'):
                    if re.match(r'\d{4}-\d{2}-\d{2}', part):
                        date_str = part
                        break
                if date_str:
                    try:
                        dt = datetime.strptime(date_str, '%Y-%m-%d')
                        date_str = dt.strftime('%d.%m.%Y')
                    except:
                        pass
                else:
                    date_str = "неизвестно"
                if ex_name not in exercises_amp_data:
                    exercises_amp_data[ex_name] = {}
                exercises_amp_data[ex_name][date_str] = max(exercises_amp_data[ex_name].get(date_str, 0), max_force)

    if not exercises_info:
        return

    doc = Document()
    title = doc.add_heading(f'Полный отчёт по пациенту: {patient_name}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    doc.add_page_break()

    for ex in exercises_info:
        doc.add_heading(f'Упражнение: {ex["name"]}', level=2)
        doc.add_picture(ex['angles_png'], width=Inches(5.0))
        doc.add_picture(ex['forces_png'], width=Inches(5.0))
        doc.add_page_break()

    if exercises_amp_data:
        doc.add_heading('Динамика амплитуд упражнений', level=1)
        for ex_name, date_force_dict in exercises_amp_data.items():
            items = sorted(date_force_dict.items(),
                           key=lambda x: datetime.strptime(x[0], '%d.%m.%Y') if re.match(r'\d{2}\.\d{2}\.\d{4}',
                                                                                         x[0]) else datetime.min)
            dates = [d[0] for d in items]
            forces = [d[1] for d in items]

            plt.figure(figsize=(8, 5), dpi=100, facecolor='white')
            plt.plot(dates, forces, marker='o', linestyle='-', linewidth=2, markersize=8, color='#5a9eff')
            plt.xlabel('Дата')
            plt.ylabel('Максимальная сила (Н)')
            plt.title(f'{ex_name} – динамика максимальной силы', fontsize=12)
            plt.xticks(rotation=45)
            plt.grid(True, linestyle='--', alpha=0.7)
            plt.tight_layout()
            temp_png = os.path.join(tempfile.gettempdir(), f"amp_report_{ex_name}.png")
            plt.savefig(temp_png, dpi=100, bbox_inches='tight')
            plt.close('all')
            gc.collect()

            doc.add_heading(f'{ex_name}', level=3)
            doc.add_picture(temp_png, width=Inches(6.0))
            if os.path.exists(temp_png):
                os.remove(temp_png)

    full_path = os.path.join(patient_dir, 'full_report.docx')
    doc.save(full_path)
    gc.collect()


# ============================================================
# PROCESSING THREAD
# ============================================================
class ProcessingThread(QThread):
    finished_signal = pyqtSignal(bool, str, str)

    def __init__(self, docx_path, patient_folder):
        super().__init__()
        self.docx_path = docx_path
        self.patient_folder = patient_folder

    def run(self):
        try:
            txt_path = docx_to_txt(self.docx_path)
            exercises = RobustParser.parse(txt_path)
            os.unlink(txt_path)
            gc.collect()

            patient_dir = os.path.join(PATIENTS_DIR, self.patient_folder)

            for ex in exercises:
                ex_dir = os.path.join(patient_dir, ex['folder_name'])
                os.makedirs(ex_dir, exist_ok=True)
                angles_png, forces_png = save_graphs_for_exercise(ex, ex_dir)
                if angles_png and forces_png:
                    create_exercise_report(ex, ex_dir, angles_png, forces_png)
                gc.collect()

            create_full_report(self.patient_folder, patient_dir)
            gc.collect()

            self.finished_signal.emit(True, f"Отчёт успешно добавлен к пациенту {self.patient_folder}", patient_dir)
        except Exception as e:
            self.finished_signal.emit(False, f"Ошибка при обработке:\n{str(e)}", "")


# ============================================================
# AI
# ============================================================
class LMStudioThread(QThread):
    finished = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, prompt, url="http://localhost:1234", model=""):
        super().__init__()
        self.prompt = prompt
        self.url = url.rstrip('/') + "/v1/chat/completions"
        self.model = model.strip() or "oreal-deepseek-r1-distill-qwen-7b"

    def run(self):
        try:
            system_prompt = (
                "Ты — профессиональный ассистент-реабилитолог по восстановлению нижних конечностей. "
                "Отвечай ИСКЛЮЧИТЕЛЬНО на русском языке. "
                "Никогда не используй английский, китайский или любой другой язык. "
                "Не вставляй иностранные слова, иероглифы или символы. "
                "Все ответы должны быть написаны только на чистом русском языке."
            )
            payload = {
                "model": self.model,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": self.prompt}
                ],
                "temperature": 0.7,
                "max_tokens": 2048
            }
            response = requests.post(self.url, json=payload, timeout=None)
            response.raise_for_status()
            data = response.json()
            answer = data["choices"][0]["message"]["content"]
            self.finished.emit(answer)
        except requests.exceptions.ConnectionError:
            self.error.emit("Не удалось подключиться к LM Studio.\n\nЗапустите LM Studio → Local Server → Start Server")
        except Exception as e:
            self.error.emit(f"Ошибка: {str(e)}")


class AIAnalysisWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent_page = parent
        self.current_exercise = None
        self.lmstudio_thread = None
        self.setup_ui()

    def setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(8)

        header = QHBoxLayout()
        self.title_label = QLabel("AI-анализ упражнения")
        self.title_label.setStyleSheet("font-size: 15px; font-weight: 700; color: #5a9eff;")
        header.addWidget(self.title_label)
        header.addStretch()

        self.refresh_btn = QPushButton("Обновить анализ")
        self.refresh_btn.setFixedHeight(36)
        self.refresh_btn.clicked.connect(self.update_analysis)
        header.addWidget(self.refresh_btn)
        layout.addLayout(header)

        self.result_text = QTextEdit()
        self.result_text.setReadOnly(True)
        self.result_text.setMaximumHeight(280)
        layout.addWidget(self.result_text)

        self.status_label = QLabel("Выберите упражнение")
        self.status_label.setStyleSheet("color: #888; font-size: 11px;")
        self.status_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.status_label)

    def set_exercise(self, exercise_name):
        self.current_exercise = exercise_name
        self.title_label.setText(f"AI-анализ: {exercise_name}")
        self.update_analysis()

    def update_analysis(self):
        if not self.parent_page or not self.parent_page.current_patient or not self.current_exercise:
            self.status_label.setText("❌ Выберите упражнение")
            return

        self.status_label.setText("⏳ Анализ...")
        self.refresh_btn.setEnabled(False)
        self.result_text.clear()

        try:
            patient_dir = os.path.join(PATIENTS_DIR, self.parent_page.current_patient)
            data = collect_analysis_data(self.parent_page.current_patient, patient_dir, self.current_exercise)
            if not data or not data.get('exercises'):
                self.result_text.setText("Нет данных по этому упражнению")
                self.refresh_btn.setEnabled(True)
                return

            prompt = self.create_prompt(data)
            self.lmstudio_thread = LMStudioThread(prompt)
            self.lmstudio_thread.finished.connect(self.on_response)
            self.lmstudio_thread.error.connect(self.on_error)
            self.lmstudio_thread.start()
        except Exception:
            self.status_label.setText("Ошибка")
            self.refresh_btn.setEnabled(True)

    def create_prompt(self, data):
        ex = data['exercises'][0] if data['exercises'] else {}
        prompt = f"Упражнение: {data['patient_name']} — {ex.get('name', '')}\n"
        prompt += f"Выполнений: {ex.get('measurements_count', 0)}\n"
        prompt += f"Максимальная сила: {ex.get('max_force', 0):.0f} Н\n"
        prompt += f"Тренд: {ex.get('trend', '')}\n\n"
        prompt += "Дай краткий анализ динамики и рекомендации по реабилитации только по этому упражнению."
        return prompt

    def on_response(self, response):
        self.result_text.setText(response)
        self.status_label.setText("✅ Готово")
        self.refresh_btn.setEnabled(True)

    def on_error(self, error_msg):
        self.status_label.setText("⚠️ Ошибка")
        self.result_text.setText(error_msg)
        self.refresh_btn.setEnabled(True)


def collect_analysis_data(patient_name, patient_dir, exercise_name=None):
    exercises_data = {}
    for folder in os.listdir(patient_dir):
        folder_path = os.path.join(patient_dir, folder)
        if not os.path.isdir(folder_path) or not os.path.exists(os.path.join(folder_path, 'max_force.txt')):
            continue
        ex_name = folder.split('_', 1)[0] if '_' in folder else folder
        if exercise_name and ex_name != exercise_name:
            continue
        with open(os.path.join(folder_path, 'max_force.txt'), 'r') as f:
            max_force = float(f.read().strip())
        date_str = ""
        for part in folder.split('_'):
            if re.match(r'\d{4}-\d{2}-\d{2}', part):
                date_str = part
                break
        date_obj = datetime.now()
        if date_str:
            try:
                date_obj = datetime.strptime(date_str, '%Y-%m-%d')
            except:
                pass
        if ex_name not in exercises_data:
            exercises_data[ex_name] = []
        exercises_data[ex_name].append((date_obj, max_force))

    for ex_name in exercises_data:
        exercises_data[ex_name].sort(key=lambda x: x[0])

    exercises_list = []
    for ex_name, data in exercises_data.items():
        if len(data) >= 2:
            dates_num = np.array([(d[0] - data[0][0]).days for d in data]).reshape(-1, 1)
            forces = np.array([d[1] for d in data])
            model = LinearRegression()
            model.fit(dates_num, forces)
            trend_slope = model.coef_[0]
            trend = "рост" if trend_slope > 0 else "снижение" if trend_slope < 0 else "стабильно"
        else:
            trend = "мало данных"

        exercises_list.append({
            'name': ex_name,
            'latest_force': data[-1][1] if data else 0,
            'trend': trend,
            'measurements_count': len(data),
            'max_force': max([d[1] for d in data]) if data else 0,
            'min_force': min([d[1] for d in data]) if data else 0
        })

    return {
        'patient_name': patient_name,
        'exercises': exercises_list,
        'summary': {}
    }


# ============================================================
# CARDS
# ============================================================
class AnimatedCard(QWidget):
    def __init__(self, title, subtitle, icon="", parent=None):
        super().__init__(parent)
        self.setFixedHeight(98)
        layout = QVBoxLayout(self)
        layout.setContentsMargins(18, 12, 18, 10)
        layout.setSpacing(4)

        top = QHBoxLayout()
        if icon:
            self.icon_label = QLabel(icon)
            self.icon_label.setStyleSheet("font-size: 26px; background: transparent;")
            top.addWidget(self.icon_label)
        self.title_label = QLabel(title)
        self.title_label.setStyleSheet("font-size: 17px; font-weight: 700; color: #f0f0ff; background: transparent;")
        top.addWidget(self.title_label)
        top.addStretch()
        layout.addLayout(top)

        self.subtitle_label = QLabel(subtitle)
        self.subtitle_label.setStyleSheet("font-size: 13px; color: #b0b0d0; background: transparent;")
        self.subtitle_label.setWordWrap(True)
        self.subtitle_label.setMinimumHeight(48)
        layout.addWidget(self.subtitle_label)

        shadow = QGraphicsDropShadowEffect()
        shadow.setBlurRadius(10)
        shadow.setXOffset(0)
        shadow.setYOffset(2)
        shadow.setColor(QColor(0, 0, 0, 40))
        self.setGraphicsEffect(shadow)

        self.setStyleSheet("background-color: #25253a; border-radius: 14px;")


class PatientCardWidget(AnimatedCard):
    double_clicked = pyqtSignal(str)
    clicked = pyqtSignal(str)

    def __init__(self, display_name, folder_name, subtitle="", parent=None):
        super().__init__(display_name, subtitle, "", parent)
        self.folder_name = folder_name
        self.selected = False

    def mousePressEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            self.clicked.emit(self.folder_name)
        super().mousePressEvent(event)

    def mouseDoubleClickEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            self.double_clicked.emit(self.folder_name)
        super().mouseDoubleClickEvent(event)

    def set_selected(self, selected: bool):
        self.selected = selected
        if selected:
            self.setStyleSheet("""
                PatientCardWidget {
                    background-color: #1a2a4a;
                    border: 4px solid #00d0ff;
                    border-radius: 14px;
                }
            """)
        else:
            self.setStyleSheet("""
                PatientCardWidget {
                    background-color: #25253a;
                    border-radius: 14px;
                }
            """)
        self.update()
        self.repaint()


class ExerciseCardWidget(AnimatedCard):
    def __init__(self, name, subtitle="", icon="📊", parent=None):
        super().__init__(name, subtitle, icon, parent)

    def get_name(self):
        return self.title_label.text()


# ============================================================
# ДИАЛОГИ
# ============================================================
class PatientInfoDialog(QDialog):
    def __init__(self, existing_name="", existing_birth="", existing_complaint="",
                 existing_height="", existing_weight="", existing_upper="", existing_middle="", existing_lower="",
                 parent=None):
        super().__init__(parent)
        self.setWindowTitle(" ")
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.Dialog)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
        self.setMinimumWidth(560)
        self.setMinimumHeight(620)

        self.main_frame = QFrame(self)
        self.main_frame.setObjectName("mainFrame")
        self.main_frame.setStyleSheet("""
            QFrame#mainFrame {
                background-color: #1f2533;
                border-radius: 24px;
                border: 1px solid #3a4459;
            }
        """)
        shadow = QGraphicsDropShadowEffect()
        shadow.setBlurRadius(30)
        shadow.setXOffset(0)
        shadow.setYOffset(10)
        shadow.setColor(QColor(0, 0, 0, 160))
        self.main_frame.setGraphicsEffect(shadow)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.addWidget(self.main_frame)

        inner_layout = QVBoxLayout(self.main_frame)
        inner_layout.setContentsMargins(40, 30, 40, 30)
        inner_layout.setSpacing(18)

        title = QLabel("Добавить нового пациента" if not existing_name else "Редактировать карточку пациента")
        title.setStyleSheet("font-size: 24px; font-weight: 700; color: #5a9eff;")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        inner_layout.addWidget(title)

        surname, first_name, patronymic = self._split_full_name(existing_name)

        form = QGridLayout()
        form.setSpacing(12)

        self.surname_edit = QLineEdit(surname)
        form.addWidget(QLabel("Фамилия: *"), 0, 0)
        form.addWidget(self.surname_edit, 0, 1)

        self.first_name_edit = QLineEdit(first_name)
        form.addWidget(QLabel("Имя: *"), 1, 0)
        form.addWidget(self.first_name_edit, 1, 1)

        self.patronymic_edit = QLineEdit(patronymic)
        form.addWidget(QLabel("Отчество (при наличии)"), 2, 0)
        form.addWidget(self.patronymic_edit, 2, 1)

        self.birth_edit = QLineEdit(existing_birth)
        self.birth_edit.setPlaceholderText("01.01.1990")
        self.birth_edit.textChanged.connect(self.auto_format_birth_date)
        form.addWidget(QLabel("Дата рождения (ДД.ММ.ГГГГ): *"), 3, 0)
        form.addWidget(self.birth_edit, 3, 1)

        self.height_edit = QLineEdit(existing_height)
        self.height_edit.setPlaceholderText("170")
        form.addWidget(QLabel("Рост (см): *"), 4, 0)
        form.addWidget(self.height_edit, 4, 1)

        self.weight_edit = QLineEdit(existing_weight)
        self.weight_edit.setPlaceholderText("70.5")
        form.addWidget(QLabel("Вес (кг): *"), 5, 0)
        form.addWidget(self.weight_edit, 5, 1)

        self.upper_link_edit = QLineEdit(existing_upper)
        self.upper_link_edit.setPlaceholderText("45")
        form.addWidget(QLabel("Длина верхнего звена (см): *"), 6, 0)
        form.addWidget(self.upper_link_edit, 6, 1)

        self.middle_link_edit = QLineEdit(existing_middle)
        self.middle_link_edit.setPlaceholderText("40")
        form.addWidget(QLabel("Длина среднего звена (см): *"), 7, 0)
        form.addWidget(self.middle_link_edit, 7, 1)

        self.lower_link_edit = QLineEdit(existing_lower)
        self.lower_link_edit.setPlaceholderText("35")
        form.addWidget(QLabel("Длина нижнего звена (см): *"), 8, 0)
        form.addWidget(self.lower_link_edit, 8, 1)

        self.complaint_combo = QComboBox()
        self.complaint_combo.setEditable(True)
        self.complaint_combo.addItems([
            "Травма коленного сустава", "Перелом голени / лодыжки",
            "Артроз тазобедренного сустава", "Артроз коленного сустава",
            "Ампутация нижней конечности", "Постоперационная реабилитация",
            "Протезирование коленного сустава", "Протезирование тазобедренного сустава",
            "Повреждение связок голеностопа", "Другое (указать)"
        ])
        if existing_complaint:
            self.complaint_combo.setCurrentText(existing_complaint)
        form.addWidget(QLabel("Причина обращения:"), 9, 0)
        form.addWidget(self.complaint_combo, 9, 1)

        inner_layout.addLayout(form)

        name_validator = QRegularExpressionValidator(QRegularExpression(r"^[а-яА-ЯёЁa-zA-Z]+$"))
        self.surname_edit.setValidator(name_validator)
        self.first_name_edit.setValidator(name_validator)
        self.patronymic_edit.setValidator(name_validator)

        number_validator = QDoubleValidator(0, 999.99, 2)
        number_validator.setNotation(QDoubleValidator.Notation.StandardNotation)
        self.height_edit.setValidator(number_validator)
        self.weight_edit.setValidator(number_validator)
        self.upper_link_edit.setValidator(number_validator)
        self.middle_link_edit.setValidator(number_validator)
        self.lower_link_edit.setValidator(number_validator)

        btn_layout = QHBoxLayout()
        self.save_btn = QPushButton("Сохранить")
        self.save_btn.setFixedHeight(50)
        self.save_btn.clicked.connect(self.accept)
        self.cancel_btn = QPushButton("Отмена")
        self.cancel_btn.setFixedHeight(50)
        self.cancel_btn.clicked.connect(self.reject)

        btn_layout.addWidget(self.save_btn)
        btn_layout.addWidget(self.cancel_btn)
        inner_layout.addLayout(btn_layout)

    def _split_full_name(self, full_name):
        if not full_name:
            return "", "", ""
        parts = full_name.strip().split()
        surname = parts[0] if parts else ""
        first_name = parts[1] if len(parts) > 1 else ""
        patronymic = " ".join(parts[2:]) if len(parts) > 2 else ""
        return surname, first_name, patronymic

    def auto_format_birth_date(self, text):
        self.birth_edit.blockSignals(True)
        cleaned = re.sub(r'\D', '', text)
        formatted = ""
        if len(cleaned) > 0:
            formatted = cleaned[:2]
            if len(cleaned) > 2:
                formatted += '.' + cleaned[2:4]
            if len(cleaned) > 4:
                formatted += '.' + cleaned[4:8]
                if len(cleaned) >= 8:
                    self.birth_edit.clearFocus()
        self.birth_edit.setText(formatted)
        self.birth_edit.setCursorPosition(len(formatted))
        self.birth_edit.blockSignals(False)

    def get_info(self):
        surname = self.surname_edit.text().strip()
        first_name = self.first_name_edit.text().strip()
        if not surname or not first_name:
            QMessageBox.warning(self, "Ошибка", "Фамилия и Имя — обязательные поля!")
            return None, None, None, None, None, None, None, None
        birth = self.birth_edit.text().strip()
        if not birth:
            QMessageBox.warning(self, "Ошибка", "Дата рождения — обязательное поле!")
            return None, None, None, None, None, None, None, None
        height = self.height_edit.text().strip()
        weight = self.weight_edit.text().strip()
        if not height or not weight:
            QMessageBox.warning(self, "Ошибка", "Рост и Вес — обязательные поля!")
            return None, None, None, None, None, None, None, None
        upper = self.upper_link_edit.text().strip()
        middle = self.middle_link_edit.text().strip()
        lower = self.lower_link_edit.text().strip()
        if not upper or not middle or not lower:
            QMessageBox.warning(self, "Ошибка", "Длины звеньев — обязательные поля!")
            return None, None, None, None, None, None, None, None

        full_name = f"{surname} {first_name}"
        if self.patronymic_edit.text().strip():
            full_name += f" {self.patronymic_edit.text().strip()}"

        return (full_name, birth, self.complaint_combo.currentText().strip(),
                height, weight, upper, middle, lower)


class DeleteConfirmDialog(QDialog):
    def __init__(self, patient_name, parent=None):
        super().__init__(parent)
        self.setWindowTitle(" ")
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.Dialog)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
        self.setMinimumWidth(480)
        self.setMinimumHeight(220)

        self.main_frame = QFrame(self)
        self.main_frame.setObjectName("mainFrame")
        self.main_frame.setStyleSheet("""
            QFrame#mainFrame {
                background-color: #1f2533;
                border-radius: 24px;
                border: 1px solid #3a4459;
            }
        """)
        shadow = QGraphicsDropShadowEffect()
        shadow.setBlurRadius(30)
        shadow.setXOffset(0)
        shadow.setYOffset(10)
        shadow.setColor(QColor(0, 0, 0, 160))
        self.main_frame.setGraphicsEffect(shadow)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.addWidget(self.main_frame)

        inner = QVBoxLayout(self.main_frame)
        inner.setContentsMargins(40, 30, 40, 30)
        inner.setSpacing(20)

        title = QLabel("Удаление пациента")
        title.setStyleSheet("font-size: 22px; font-weight: 700; color: #ff6b6b;")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        inner.addWidget(title)

        msg = QLabel(f"Вы действительно хотите удалить пациента\n«{patient_name}»\nи все его данные?")
        msg.setStyleSheet("font-size: 16px; color: #e0e0e0; text-align: center;")
        msg.setWordWrap(True)
        msg.setAlignment(Qt.AlignmentFlag.AlignCenter)
        inner.addWidget(msg)

        btn_layout = QHBoxLayout()
        self.yes_btn = QPushButton("Да, удалить")
        self.yes_btn.setFixedHeight(48)
        self.yes_btn.setStyleSheet("""
            QPushButton { background-color: #ff4444; color: white; border-radius: 12px; font-weight: 600; }
            QPushButton:hover { background-color: #ff6666; }
        """)
        self.yes_btn.clicked.connect(self.accept)

        self.no_btn = QPushButton("Отмена")
        self.no_btn.setFixedHeight(48)
        self.no_btn.setStyleSheet("""
            QPushButton { background-color: #2a2a3a; border: 2px solid #5a9eff; border-radius: 12px; color: white; font-weight: 600; }
            QPushButton:hover { background-color: #3a3a4e; }
        """)
        self.no_btn.clicked.connect(self.reject)

        btn_layout.addWidget(self.yes_btn)
        btn_layout.addWidget(self.no_btn)
        inner.addLayout(btn_layout)


class DeleteReportDialog(QDialog):
    def __init__(self, patient_name, available_dates, parent=None):
        super().__init__(parent)
        self.setWindowTitle(" ")
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.Dialog)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
        self.setMinimumWidth(480)
        self.setMinimumHeight(320)

        self.main_frame = QFrame(self)
        self.main_frame.setObjectName("mainFrame")
        self.main_frame.setStyleSheet("""
            QFrame#mainFrame {
                background-color: #1f2533;
                border-radius: 24px;
                border: 1px solid #3a4459;
            }
        """)
        shadow = QGraphicsDropShadowEffect()
        shadow.setBlurRadius(30)
        shadow.setXOffset(0)
        shadow.setYOffset(10)
        shadow.setColor(QColor(0, 0, 0, 160))
        self.main_frame.setGraphicsEffect(shadow)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.addWidget(self.main_frame)

        inner = QVBoxLayout(self.main_frame)
        inner.setContentsMargins(40, 30, 40, 30)
        inner.setSpacing(20)

        title = QLabel("Удаление отчёта")
        title.setStyleSheet("font-size: 22px; font-weight: 700; color: #ff6b6b;")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        inner.addWidget(title)

        msg = QLabel(f"Выберите дату отчёта для удаления\n(пациент: {patient_name})")
        msg.setStyleSheet("font-size: 15px; color: #e0e0e0; text-align: center;")
        msg.setWordWrap(True)
        msg.setAlignment(Qt.AlignmentFlag.AlignCenter)
        inner.addWidget(msg)

        inner.addWidget(QLabel("Дата выполнения упражнений:"))
        self.date_combo = QComboBox()
        self.date_combo.setStyleSheet("""
            QComboBox {
                background-color: #2a2a3a;
                color: white;
                border: 1px solid #5a9eff;
                border-radius: 8px;
                padding: 10px;
                font-size: 14px;
            }
            QComboBox::drop-down {
                border: none;
            }
            QComboBox::down-arrow {
                image: none;
                border: none;
            }
        """)
        for date in sorted(available_dates, reverse=True):
            self.date_combo.addItem(date)
        inner.addWidget(self.date_combo)

        if not available_dates:
            no_dates_label = QLabel("Нет доступных отчётов для удаления")
            no_dates_label.setStyleSheet("color: #ff8888; font-size: 14px;")
            no_dates_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            inner.addWidget(no_dates_label)

        btn_layout = QHBoxLayout()
        self.delete_btn = QPushButton("Удалить отчёт")
        self.delete_btn.setFixedHeight(48)
        self.delete_btn.setStyleSheet("""
            QPushButton { background-color: #ff4444; color: white; border-radius: 12px; font-weight: 600; }
            QPushButton:hover { background-color: #ff6666; }
        """)
        self.delete_btn.clicked.connect(self.accept)

        self.cancel_btn = QPushButton("Отмена")
        self.cancel_btn.setFixedHeight(48)
        self.cancel_btn.setStyleSheet("""
            QPushButton { background-color: #2a2a3a; border: 2px solid #5a9eff; border-radius: 12px; color: white; font-weight: 600; }
            QPushButton:hover { background-color: #3a3a4e; }
        """)
        self.cancel_btn.clicked.connect(self.reject)

        btn_layout.addWidget(self.delete_btn)
        btn_layout.addWidget(self.cancel_btn)
        inner.addLayout(btn_layout)

        if not available_dates:
            self.delete_btn.setEnabled(False)
            self.delete_btn.setStyleSheet("background-color: #555; color: #aaa; border-radius: 12px;")

    def get_selected_date(self):
        return self.date_combo.currentText()


class PatientFilterDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Фильтры поиска пациентов")
        self.setMinimumWidth(420)
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.Dialog)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)

        self.main_frame = QFrame(self)
        self.main_frame.setObjectName("mainFrame")
        self.main_frame.setStyleSheet("""
            QFrame#mainFrame {
                background-color: #1f2533;
                border-radius: 20px;
                border: 1px solid #3a4459;
            }
        """)
        shadow = QGraphicsDropShadowEffect()
        shadow.setBlurRadius(25)
        shadow.setXOffset(0)
        shadow.setYOffset(8)
        shadow.setColor(QColor(0, 0, 0, 140))
        self.main_frame.setGraphicsEffect(shadow)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.addWidget(self.main_frame)

        inner = QVBoxLayout(self.main_frame)
        inner.setContentsMargins(30, 25, 30, 25)
        inner.setSpacing(18)

        title = QLabel("Фильтры")
        title.setStyleSheet("font-size: 20px; font-weight: 700; color: #5a9eff;")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        inner.addWidget(title)

        inner.addWidget(QLabel("Дата рождения (ДД.ММ.ГГГГ):"))
        self.birth_edit = QLineEdit()
        self.birth_edit.setPlaceholderText("01.01.1990")
        inner.addWidget(self.birth_edit)

        inner.addWidget(QLabel("Причина обращения:"))
        self.complaint_edit = QLineEdit()
        self.complaint_edit.setPlaceholderText("например: травма колена")
        inner.addWidget(self.complaint_edit)

        btn_layout = QHBoxLayout()
        self.apply_btn = QPushButton("Применить")
        self.apply_btn.setFixedHeight(48)
        self.apply_btn.clicked.connect(self.accept)

        self.cancel_btn = QPushButton("Отмена")
        self.cancel_btn.setFixedHeight(48)
        self.cancel_btn.clicked.connect(self.reject)

        btn_layout.addWidget(self.apply_btn)
        btn_layout.addWidget(self.cancel_btn)
        inner.addLayout(btn_layout)

    def get_filters(self):
        return {
            'birth': self.birth_edit.text().strip().lower(),
            'complaint': self.complaint_edit.text().strip().lower()
        }


# ============================================================
# PatientsListPage
# ============================================================
class PatientsListPage(QWidget):
    patient_selected = pyqtSignal(str)

    def __init__(self, main_window, parent=None):
        super().__init__(parent)
        self.main_window = main_window
        self.all_patients = []
        self.filtered_patients = []
        self.selected_folder = None
        self.current_filters = None

        layout = QVBoxLayout(self)
        layout.setContentsMargins(30, 30, 30, 30)
        layout.setSpacing(18)

        title = QLabel("SWSU ROBOTICS")
        title.setStyleSheet("font-size: 32px; font-weight: 900; color: #5a9eff;")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title)

        self.add_patient_btn = QPushButton("Добавить нового пациента")
        self.add_patient_btn.setFixedHeight(54)
        self.add_patient_btn.clicked.connect(self.main_window.add_new_patient)
        self.add_patient_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #1e90ff, stop:1 #0b7be0);
                color: white;
                border-radius: 12px;
                font-size: 17px;
                font-weight: 600;
                border: none;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #40b0ff, stop:1 #1e90ff);
            }
            QPushButton:pressed {
                background: #0a5eb8;
            }
        """)
        layout.addWidget(self.add_patient_btn)

        search_layout = QHBoxLayout()
        self.search_edit = QLineEdit()
        self.search_edit.setPlaceholderText("🔍 Поиск по ФИО...")
        self.search_edit.textChanged.connect(self.filter_patients)
        search_layout.addWidget(self.search_edit)

        self.filter_btn = QPushButton("🔧 Фильтры")
        self.filter_btn.clicked.connect(self.show_filter_dialog)
        search_layout.addWidget(self.filter_btn)
        layout.addLayout(search_layout)

        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)

        self.list_container = QWidget()
        self.list_layout = QVBoxLayout(self.list_container)
        self.list_layout.setContentsMargins(0, 0, 0, 0)
        self.list_layout.setSpacing(14)
        self.scroll.setWidget(self.list_container)
        layout.addWidget(self.scroll)

        self.load_patients()

    def load_patients(self):
        patients = []
        ignore_folders = ['__pycache__', 'dist', 'build', 'venv', 'myenv', 'env', '.git', '.idea']
        for item in os.listdir(PATIENTS_DIR):
            item_path = os.path.join(PATIENTS_DIR, item)
            if not os.path.isdir(item_path) or item.startswith('.') or item in ignore_folders:
                continue
            if os.path.exists(os.path.join(item_path, 'info.txt')):
                patients.append(item)
        self.all_patients = sorted(patients)
        self.filter_patients()

    def show_filter_dialog(self):
        dlg = PatientFilterDialog(self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            self.current_filters = dlg.get_filters()
        else:
            self.current_filters = None
        self.filter_patients()

    def filter_patients(self):
        text = self.search_edit.text().strip().lower()
        filters = self.current_filters

        result = []
        for folder in self.all_patients:
            patient_dir = os.path.join(PATIENTS_DIR, folder)
            info_path = os.path.join(patient_dir, 'info.txt')
            display_name = folder
            birth = ""
            complaint = ""
            if os.path.exists(info_path):
                try:
                    with open(info_path, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                        display_name = data.get('name', folder)
                        birth = data.get('birth_date', '')
                        complaint = data.get('complaint', '').lower()
                except:
                    pass

            if text and text not in display_name.lower() and text not in folder.lower():
                continue
            if filters:
                if filters.get('birth') and filters['birth'] not in birth.lower():
                    continue
                if filters.get('complaint') and filters['complaint'] not in complaint:
                    continue

            result.append((folder, display_name, birth))

        self.filtered_patients = result
        self.update_list()

    def update_list(self):
        for i in reversed(range(self.list_layout.count())):
            widget = self.list_layout.itemAt(i).widget()
            if widget:
                widget.deleteLater()
        gc.collect()

        for folder, display_name, birth in self.filtered_patients:
            patient_dir = os.path.join(PATIENTS_DIR, folder)
            if not os.path.exists(patient_dir):
                continue
            exercises_count = sum(
                1 for f in os.listdir(patient_dir)
                if os.path.isdir(os.path.join(patient_dir, f)) and os.path.exists(
                    os.path.join(patient_dir, f, 'angles.png'))
            )
            subtitle = f"Дата рождения: {birth}" if birth else ""
            if exercises_count:
                if subtitle:
                    subtitle += f" • {exercises_count} упражнений"
                else:
                    subtitle = f"{exercises_count} упражнений"

            card = PatientCardWidget(display_name, folder, subtitle)
            card.double_clicked.connect(self.on_patient_double_clicked)
            card.clicked.connect(self.on_patient_clicked)

            if folder == self.selected_folder:
                card.set_selected(True)

            self.list_layout.addWidget(card)

        spacer = QWidget()
        spacer.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self.list_layout.addWidget(spacer)

    def on_patient_clicked(self, folder_name):
        self.selected_folder = folder_name
        self.update_list()

    def on_patient_double_clicked(self, folder_name):
        self.selected_folder = folder_name
        self.update_list()
        self.patient_selected.emit(folder_name)

    def set_selected_patient(self, folder_name: str):
        self.selected_folder = folder_name
        self.update_list()

    def refresh(self):
        self.load_patients()


# ============================================================
# PatientWithExercisesPage
# ============================================================
class PatientWithExercisesPage(QWidget):
    back_to_list = pyqtSignal()
    exercise_selected = pyqtSignal(str, str)

    def __init__(self, main_window, parent=None):
        super().__init__(parent)
        self.main_window = main_window
        self.current_patient = None
        self.exercise_variants = {}

        layout = QHBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(20)

        self.left_panel = QFrame()
        self.left_panel.setFrameShape(QFrame.Shape.StyledPanel)
        self.left_panel.setStyleSheet("background-color: #2a2a3a; border-radius: 18px;")
        left_layout = QVBoxLayout(self.left_panel)
        left_layout.setContentsMargins(22, 22, 22, 22)

        self.back_btn = QPushButton("←")
        self.back_btn.setFixedSize(50, 44)
        self.back_btn.setStyleSheet("""
            QPushButton {
                background-color: #ff4444;
                color: white;
                font-size: 24px;
                font-weight: bold;
                border: none;
                border-radius: 8px;
            }
            QPushButton:hover {
                background-color: #ff6666;
            }
        """)
        self.back_btn.clicked.connect(self.back_to_list.emit)
        left_layout.addWidget(self.back_btn)

        self.patient_info_group = QGroupBox("Карточка пациента")
        self.patient_info_group.setStyleSheet(
            "QGroupBox { color: white; font-size: 15px; border: 1px solid #444; border-radius: 14px; margin-top: 12px; }")
        info_layout = QGridLayout(self.patient_info_group)
        self.patient_name_label = QLabel()
        self.patient_name_label.setStyleSheet("font-size: 21px; font-weight: bold; color: #5a9eff;")
        self.patient_birth_label = QLabel()
        self.patient_height_label = QLabel()
        self.patient_weight_label = QLabel()
        self.patient_upper_link_label = QLabel()
        self.patient_middle_link_label = QLabel()
        self.patient_lower_link_label = QLabel()
        self.patient_complaint_label = QLabel()
        self.patient_exercises_count_label = QLabel()

        info_layout.addWidget(QLabel("ФИО:"), 0, 0)
        info_layout.addWidget(self.patient_name_label, 0, 1)
        info_layout.addWidget(QLabel("Дата рождения:"), 1, 0)
        info_layout.addWidget(self.patient_birth_label, 1, 1)
        info_layout.addWidget(QLabel("Рост:"), 2, 0)
        info_layout.addWidget(self.patient_height_label, 2, 1)
        info_layout.addWidget(QLabel("Вес:"), 3, 0)
        info_layout.addWidget(self.patient_weight_label, 3, 1)
        info_layout.addWidget(QLabel("Верхнее звено:"), 4, 0)
        info_layout.addWidget(self.patient_upper_link_label, 4, 1)
        info_layout.addWidget(QLabel("Среднее звено:"), 5, 0)
        info_layout.addWidget(self.patient_middle_link_label, 5, 1)
        info_layout.addWidget(QLabel("Нижнее звено:"), 6, 0)
        info_layout.addWidget(self.patient_lower_link_label, 6, 1)
        info_layout.addWidget(QLabel("Причина обращения:"), 7, 0)
        info_layout.addWidget(self.patient_complaint_label, 7, 1)
        info_layout.addWidget(QLabel("Упражнений:"), 8, 0)
        info_layout.addWidget(self.patient_exercises_count_label, 8, 1)
        left_layout.addWidget(self.patient_info_group)

        btn_layout = QVBoxLayout()
        btn_layout.setSpacing(12)

        self.add_report_btn = QPushButton("📋  Добавить отчёт")
        self.add_report_btn.setFixedHeight(50)
        self.add_report_btn.setStyleSheet("""
            QPushButton {
                background-color: #2a2a3a;
                border: 2px solid #5a9eff;
                border-radius: 12px;
                color: white;
                font-weight: 600;
            }
            QPushButton:hover {
                background-color: #3a3a4e;
                border: 2px solid #6ab0ff;
            }
        """)
        self.add_report_btn.clicked.connect(self.add_report_to_patient)
        btn_layout.addWidget(self.add_report_btn)

        self.edit_btn = QPushButton("✏️  Редактировать карточку")
        self.edit_btn.setFixedHeight(50)
        self.edit_btn.setStyleSheet("""
            QPushButton {
                background-color: #2a2a3a;
                border: 2px solid #5a9eff;
                border-radius: 12px;
                color: white;
                font-weight: 600;
            }
            QPushButton:hover {
                background-color: #3a3a4e;
                border: 2px solid #6ab0ff;
            }
        """)
        self.edit_btn.clicked.connect(self.edit_current_patient)
        btn_layout.addWidget(self.edit_btn)

        self.delete_report_btn = QPushButton("🗑  Удалить отчёт по дате")
        self.delete_report_btn.setFixedHeight(50)
        self.delete_report_btn.setStyleSheet("""
            QPushButton {
                background-color: #2a2a3a;
                border: 2px solid #ffaa44;
                border-radius: 12px;
                color: white;
                font-weight: 600;
            }
            QPushButton:hover {
                background-color: #3a3a4e;
                border: 2px solid #ffcc66;
            }
        """)
        self.delete_report_btn.clicked.connect(self.delete_report_by_date)
        btn_layout.addWidget(self.delete_report_btn)

        self.delete_patient_btn = QPushButton("🗑  Удалить пациента")
        self.delete_patient_btn.setFixedHeight(50)
        self.delete_patient_btn.setStyleSheet("""
            QPushButton {
                background-color: #2a2a3a;
                border: 2px solid #ff6b6b;
                border-radius: 12px;
                color: white;
                font-weight: 600;
            }
            QPushButton:hover {
                background-color: #3a3a4e;
                border: 2px solid #ff8a8a;
            }
        """)
        self.delete_patient_btn.clicked.connect(self.delete_patient)
        btn_layout.addWidget(self.delete_patient_btn)

        left_layout.addLayout(btn_layout)
        left_layout.addStretch()

        self.right_panel = QWidget()
        right_layout = QVBoxLayout(self.right_panel)
        right_layout.setContentsMargins(0, 0, 0, 0)

        self.ex_list_widget = QListWidget()
        self.ex_list_widget.setStyleSheet("""
            QListWidget {
                background: transparent;
                border: none;
            }
            QListWidget::item {
                background: transparent;
            }
        """)
        self.ex_list_widget.setSpacing(14)
        self.ex_list_widget.itemDoubleClicked.connect(self.on_exercise_double_click)
        right_layout.addWidget(self.ex_list_widget)

        splitter = QSplitter(Qt.Orientation.Horizontal)
        splitter.addWidget(self.left_panel)
        splitter.addWidget(self.right_panel)
        splitter.setSizes([400, 800])
        layout.addWidget(splitter)

    def edit_current_patient(self):
        if self.current_patient:
            self.main_window.edit_patient(self.current_patient)

    def set_patient(self, patient_name):
        self.current_patient = patient_name
        self.load_patient_info()
        self.load_exercises()

    def add_report_to_patient(self):
        if self.current_patient:
            self.main_window.add_report_to_existing(self.current_patient)

    def load_patient_info(self):
        patient_dir = os.path.join(PATIENTS_DIR, self.current_patient)
        info_path = os.path.join(patient_dir, 'info.txt')
        name = self.current_patient
        birth = ""
        complaint = ""
        height = ""
        weight = ""
        upper = ""
        middle = ""
        lower = ""
        if os.path.exists(info_path):
            try:
                with open(info_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    name = data.get('name', self.current_patient)
                    birth = data.get('birth_date', '')
                    complaint = data.get('complaint', '')
                    height = data.get('height_cm', '')
                    weight = data.get('weight_kg', '')
                    upper = data.get('upper_link_cm', '')
                    middle = data.get('middle_link_cm', '')
                    lower = data.get('lower_link_cm', '')
            except:
                pass
        exercises_count = sum(
            1 for f in os.listdir(patient_dir)
            if
            os.path.isdir(os.path.join(patient_dir, f)) and os.path.exists(os.path.join(patient_dir, f, 'angles.png'))
        )
        self.patient_name_label.setText(name)
        self.patient_birth_label.setText(birth if birth else "не указана")
        self.patient_height_label.setText(f"{height} см" if height else "не указан")
        self.patient_weight_label.setText(f"{weight} кг" if weight else "не указан")
        self.patient_upper_link_label.setText(f"{upper} см" if upper else "не указана")
        self.patient_middle_link_label.setText(f"{middle} см" if middle else "не указана")
        self.patient_lower_link_label.setText(f"{lower} см" if lower else "не указана")
        self.patient_complaint_label.setText(complaint if complaint else "не указана")
        self.patient_exercises_count_label.setText(str(exercises_count))

    def load_exercises(self):
        if not self.current_patient:
            return
        patient_dir = os.path.join(PATIENTS_DIR, self.current_patient)
        variants = {}

        folder_list = []
        for folder in os.listdir(patient_dir):
            folder_path = os.path.join(patient_dir, folder)
            if not os.path.isdir(folder_path) or not os.path.exists(os.path.join(folder_path, 'angles.png')):
                continue
            date_match = re.search(r'(\d{4}-\d{2}-\d{2})', folder)
            if date_match:
                try:
                    dt = datetime.strptime(date_match.group(1), '%Y-%m-%d')
                    folder_list.append((dt, folder))
                except:
                    folder_list.append((datetime.min, folder))

        folder_list.sort(key=lambda x: x[0], reverse=True)

        for _, folder in folder_list:
            folder_path = os.path.join(patient_dir, folder)
            ex_name = folder.split('_', 1)[0] if '_' in folder else folder
            date_str = ""
            parts = folder.split('_')
            for part in parts:
                if re.match(r'\d{4}-\d{2}-\d{2}', part):
                    try:
                        dt = datetime.strptime(part, '%Y-%m-%d')
                        date_str = dt.strftime('%d.%m.%Y')
                    except:
                        date_str = part
                    break
            else:
                date_str = "неизвестно"

            if ex_name not in variants:
                variants[ex_name] = []
            variants[ex_name].append((date_str, folder))

        self.exercise_variants = variants
        self.update_exercise_list()

    def update_exercise_list(self):
        self.ex_list_widget.clear()
        for ex_name in sorted(self.exercise_variants.keys()):
            variants = self.exercise_variants.get(ex_name, [])
            dates_str = ", ".join([v[0] for v in variants])
            subtitle = f"{len(variants)} вариант(ов): {dates_str}"
            icon = self._get_exercise_icon(ex_name)
            card = ExerciseCardWidget(ex_name, subtitle, icon)
            item = QListWidgetItem()
            item.setSizeHint(card.sizeHint())
            self.ex_list_widget.addItem(item)
            self.ex_list_widget.setItemWidget(item, card)

    def _get_exercise_icon(self, name):
        if "ХОДЬБА" in name: return "🚶"
        if "ПОВОРОТ" in name: return "🔄"
        if "ПРИСЕДАНИЯ" in name: return "🏋️"
        return "📊"

    def on_exercise_double_click(self, item):
        widget = self.ex_list_widget.itemWidget(item)
        if isinstance(widget, ExerciseCardWidget):
            ex_name = widget.get_name()
            variants = self.exercise_variants.get(ex_name, [])
            if variants:
                latest_folder = variants[0][1]
                self.exercise_selected.emit(self.current_patient, latest_folder)

    def get_available_dates(self):
        if not self.current_patient:
            return []
        patient_dir = os.path.join(PATIENTS_DIR, self.current_patient)
        dates = set()
        for folder in os.listdir(patient_dir):
            folder_path = os.path.join(patient_dir, folder)
            if not os.path.isdir(folder_path) or not os.path.exists(os.path.join(folder_path, 'angles.png')):
                continue
            date_match = re.search(r'(\d{4}-\d{2}-\d{2})', folder)
            if date_match:
                try:
                    dt = datetime.strptime(date_match.group(1), '%Y-%m-%d')
                    dates.add(dt.strftime('%d.%m.%Y'))
                except:
                    pass
        return sorted(dates, reverse=True)

    def delete_reports_by_date(self, date_str):
        if not self.current_patient:
            return False, 0, []
        patient_dir = os.path.join(PATIENTS_DIR, self.current_patient)
        deleted_count = 0
        deleted_folders = []
        try:
            target_date = datetime.strptime(date_str, '%d.%m.%Y')
            target_date_str = target_date.strftime('%Y-%m-%d')
        except:
            return False, 0, []
        for folder in os.listdir(patient_dir):
            folder_path = os.path.join(patient_dir, folder)
            if not os.path.isdir(folder_path):
                continue
            date_match = re.search(r'(\d{4}-\d{2}-\d{2})', folder)
            if date_match and date_match.group(1) == target_date_str:
                try:
                    shutil.rmtree(folder_path, ignore_errors=True)
                    deleted_count += 1
                    deleted_folders.append(folder)
                    gc.collect()
                except Exception:
                    pass
        if deleted_count > 0:
            create_full_report(self.current_patient, patient_dir)
            return True, deleted_count, deleted_folders
        return False, 0, []

    def delete_report_by_date(self):
        if not self.current_patient:
            QMessageBox.warning(self, "Ошибка", "Сначала выберите пациента")
            return
        available_dates = self.get_available_dates()
        if not available_dates:
            QMessageBox.information(self, "Нет отчётов", "Нет доступных отчётов для удаления")
            return
        dlg = DeleteReportDialog(self.current_patient, available_dates, self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            selected_date = dlg.get_selected_date()
            success, count, folders = self.delete_reports_by_date(selected_date)
            if success:
                QMessageBox.information(self, "Удаление завершено",
                                        f"Удалено упражнений: {count}\nДата: {selected_date}")
                self.load_patient_info()
                self.load_exercises()
            else:
                QMessageBox.warning(self, "Ошибка", "Не найдено упражнений для удаления за указанную дату")

    def delete_patient(self):
        if not self.current_patient:
            return
        dlg = DeleteConfirmDialog(self.current_patient, self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            patient_dir = os.path.join(PATIENTS_DIR, self.current_patient)
            try:
                shutil.rmtree(patient_dir, ignore_errors=True)
                gc.collect()
                QMessageBox.information(self, "Удалено", f"Пациент '{self.current_patient}' успешно удалён.")
                self.back_to_list.emit()
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось удалить пациента:\n{str(e)}")


# ============================================================
# ExerciseViewPage
# ============================================================
class ExerciseViewPage(QWidget):
    back_to_exercises = pyqtSignal()

    def __init__(self, parent=None):
        super().__init__(parent)
        layout = QHBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(20)

        self.left_panel = QGroupBox("Данные упражнения")
        left_layout = QVBoxLayout(self.left_panel)
        self.name_label = QLabel()
        self.name_label.setStyleSheet("font-size: 19px; font-weight: bold; color: #5a9eff;")
        self.count_label = QLabel("Выполнений: 0")
        self.count_label.setStyleSheet("font-size: 14px; color: #ccc;")
        left_layout.addWidget(self.name_label)
        left_layout.addWidget(self.count_label)

        spacer = QWidget()
        spacer.setFixedHeight(12)
        left_layout.addWidget(spacer)

        self.ai_widget = AIAnalysisWidget(self)
        left_layout.addWidget(self.ai_widget)
        left_layout.addStretch()

        self.center_panel = QGroupBox("График упражнения")
        self.center_panel.setFixedWidth(960)
        center_layout = QVBoxLayout(self.center_panel)
        self.center_stack = QStackedWidget()
        center_layout.addWidget(self.center_stack)

        self.amplitude_page = QWidget()
        amp_layout = QVBoxLayout(self.amplitude_page)
        self.figure = Figure(figsize=(14, 9), dpi=140)
        self.canvas = FigureCanvas(self.figure)
        self.canvas.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        amp_layout.addWidget(self.canvas)
        self.center_stack.addWidget(self.amplitude_page)

        self.images_page = QWidget()
        images_main_layout = QVBoxLayout(self.images_page)
        images_main_layout.setContentsMargins(0, 0, 0, 0)
        images_main_layout.setSpacing(12)

        self.images_scroll = QScrollArea()
        self.images_scroll.setWidgetResizable(True)
        self.images_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.images_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)

        self.images_content = QWidget()
        self.images_content.setLayout(QVBoxLayout())
        self.images_content.layout().setSpacing(25)

        self.images_scroll.setWidget(self.images_content)
        images_main_layout.addWidget(self.images_scroll)
        self.center_stack.addWidget(self.images_page)

        self.right_panel = QGroupBox("Выбор графика / сеанса")
        right_layout = QVBoxLayout(self.right_panel)
        self.graph_list = QListWidget()
        self.graph_list.setStyleSheet("""
            QListWidget {
                font-size: 15px;
                background: transparent;
                border: none;
            }
        """)
        self.graph_list.itemClicked.connect(self.on_graph_selected)
        right_layout.addWidget(self.graph_list)

        splitter = QSplitter(Qt.Orientation.Horizontal)
        splitter.addWidget(self.left_panel)
        splitter.addWidget(self.center_panel)
        splitter.addWidget(self.right_panel)
        splitter.setSizes([340, 960, 300])
        layout.addWidget(splitter)

        self.current_patient = None
        self.current_ex_name = None
        self.all_variants = []

    def set_exercise(self, patient_name, exercise_folder, exercise_path):
        self.current_patient = patient_name
        folder_name = os.path.basename(exercise_path)
        self.current_ex_name = folder_name.split('_', 1)[0] if '_' in folder_name else folder_name

        self.name_label.setText(self.current_ex_name)

        patient_dir = os.path.join(PATIENTS_DIR, patient_name)
        count = sum(1 for f in os.listdir(patient_dir) if f.startswith(self.current_ex_name + "_") and os.path.exists(
            os.path.join(patient_dir, f, 'max_force.txt')))
        self.count_label.setText(f"Выполнений: {count}")

        self.ai_widget.set_exercise(self.current_ex_name)

        self.load_date_variants(patient_name)
        self.load_amplitude_graph(patient_name)

    def load_date_variants(self, patient_name):
        self.graph_list.clear()
        self.all_variants = []
        patient_dir = os.path.join(PATIENTS_DIR, patient_name)

        amplitude_item = QListWidgetItem("📈 Динамика амплитуды (все даты)")
        amplitude_item.setForeground(QColor("#5a9eff"))
        amplitude_item.setData(Qt.ItemDataRole.UserRole, "AMPLITUDE")
        self.graph_list.addItem(amplitude_item)

        separator = QListWidgetItem("─" * 70)
        separator.setFlags(Qt.ItemFlag.NoItemFlags)
        self.graph_list.addItem(separator)

        folder_list = []
        for folder in os.listdir(patient_dir):
            if not folder.startswith(self.current_ex_name + "_") or not os.path.isdir(
                    os.path.join(patient_dir, folder)) or not os.path.exists(
                    os.path.join(patient_dir, folder, 'angles.png')):
                continue
            date_match = re.search(r'(\d{4}-\d{2}-\d{2})', folder)
            if date_match:
                try:
                    dt = datetime.strptime(date_match.group(1), '%Y-%m-%d')
                    folder_list.append((dt, folder))
                except:
                    folder_list.append((datetime.min, folder))

        folder_list.sort(key=lambda x: x[0], reverse=True)

        date_groups = {}
        for dt, folder in folder_list:
            formatted_date = dt.strftime('%d.%m.%Y')
            if formatted_date not in date_groups:
                date_groups[formatted_date] = []
            time_str = ""
            time_match = re.search(r'(\d{2}-\d{2}-\d{2})', folder)
            if time_match:
                time_str = time_match.group(1).replace('-', ':')[:5]
            display_time = f"{formatted_date} {time_str}".strip() if time_str else formatted_date
            date_groups[formatted_date].append((display_time, folder))

        first = True
        for formatted_date in date_groups.keys():
            if not first:
                sep = QListWidgetItem("─" * 70)
                sep.setFlags(Qt.ItemFlag.NoItemFlags)
                self.graph_list.addItem(sep)
            first = False

            date_header = QListWidgetItem(f"📅 {formatted_date}")
            date_header.setForeground(QColor("#88aaff"))
            date_header.setFont(QFont("Segoe UI", 11, QFont.Weight.Bold))
            date_header.setFlags(Qt.ItemFlag.NoItemFlags)
            self.graph_list.addItem(date_header)

            for display_time, folder in date_groups[formatted_date]:
                item = QListWidgetItem(f"   {display_time}")
                item.setData(Qt.ItemDataRole.UserRole, folder)
                self.graph_list.addItem(item)

    def on_graph_selected(self, item):
        if not item:
            return
        folder = item.data(Qt.ItemDataRole.UserRole)
        if folder is None:
            return
        if folder == "AMPLITUDE":
            self.load_amplitude_graph(self.current_patient)
        else:
            folder_path = os.path.join(PATIENTS_DIR, self.current_patient, folder)
            self.load_specific_session_graphs(folder_path)

    def load_amplitude_graph(self, patient_name):
        self.center_stack.setCurrentWidget(self.amplitude_page)
        self.figure.clear()
        ax = self.figure.add_subplot(111)

        patient_dir = os.path.join(PATIENTS_DIR, patient_name)
        date_max = {}

        for folder in os.listdir(patient_dir):
            if not folder.startswith(self.current_ex_name + "_"):
                continue
            folder_path = os.path.join(patient_dir, folder)
            max_force_path = os.path.join(folder_path, 'max_force.txt')
            if not os.path.exists(max_force_path):
                continue
            with open(max_force_path, 'r') as f:
                max_force = float(f.read().strip())
            date_str = ""
            parts = folder.split('_')
            for part in parts:
                if re.match(r'\d{4}-\d{2}-\d{2}', part):
                    try:
                        dt = datetime.strptime(part, '%Y-%m-%d')
                        date_str = dt.strftime('%d.%m.%Y')
                    except:
                        date_str = part
                    break
            if date_str:
                if date_str in date_max:
                    date_max[date_str] = max(date_max[date_str], max_force)
                else:
                    date_max[date_str] = max_force

        dates = list(date_max.keys())
        forces = list(date_max.values())

        if dates:
            ax.plot(dates, forces, marker='o', linestyle='-', color='#5a9eff', markersize=9)
            ax.set_xlabel('Дата')
            ax.set_ylabel('Максимальная сила (Н)')
            ax.set_title('Динамика амплитуды')
            ax.grid(True)
            plt.setp(ax.get_xticklabels(), rotation=45)
        else:
            ax.text(0.5, 0.5, 'Нет данных', ha='center', va='center')

        self.canvas.draw()

    def load_specific_session_graphs(self, folder_path):
        """Загружает отдельные графики для каждого канала"""
        graphs_info_path = os.path.join(folder_path, 'graphs_info.json')

        if not os.path.exists(graphs_info_path):
            # Fallback: старый формат
            angles_path = os.path.join(folder_path, 'angles.png')
            forces_path = os.path.join(folder_path, 'forces.png')
            if os.path.exists(angles_path) and os.path.exists(forces_path):
                self.center_stack.setCurrentWidget(self.images_page)
                layout = self.images_content.layout()
                for i in reversed(range(layout.count())):
                    widget = layout.itemAt(i).widget()
                    if widget:
                        widget.deleteLater()
                self._add_graph_to_layout("Углы", angles_path, layout)
                self._add_graph_to_layout("Силы", forces_path, layout)
                spacer = QWidget()
                spacer.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
                layout.addWidget(spacer)
            return

        with open(graphs_info_path, 'r', encoding='utf-8') as f:
            graphs_info = json.load(f)

        self.center_stack.setCurrentWidget(self.images_page)

        layout = self.images_content.layout()
        for i in reversed(range(layout.count())):
            widget = layout.itemAt(i).widget()
            if widget:
                widget.deleteLater()

        # Отображаем все графики (углы и силы)
        if 'angles' in graphs_info:
            for label, path in graphs_info['angles']:
                if os.path.exists(path):
                    self._add_graph_to_layout(label, path, layout)

        if 'forces' in graphs_info:
            for label, path in graphs_info['forces']:
                if os.path.exists(path):
                    self._add_graph_to_layout(label, path, layout)

        spacer = QWidget()
        spacer.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        layout.addWidget(spacer)

    def _add_graph_to_layout(self, label_text, img_path, layout):
        """Добавляет один график в layout"""
        if not os.path.exists(img_path):
            return

        # Заголовок
        caption = QLabel(label_text)
        caption.setStyleSheet("font-size: 14px; font-weight: 600; color: #cccccc; margin-top: 10px;")
        caption.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(caption)

        # Изображение
        pix = QPixmap(img_path)
        if not pix.isNull():
            max_width = 900
            if pix.width() > max_width:
                pix = pix.scaledToWidth(max_width, Qt.TransformationMode.SmoothTransformation)
            img_label = QLabel()
            img_label.setPixmap(pix)
            img_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            layout.addWidget(img_label)


# ============================================================
# WelcomePage
# ============================================================
class WelcomePage(QWidget):
    start_requested = pyqtSignal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
        self.setStyleSheet("background-color: #0a1628;")

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 140, 0, 80)
        layout.setSpacing(32)
        layout.setAlignment(Qt.AlignmentFlag.AlignCenter)

        self.logo_container = QLabel()
        self.logo_container.setFixedSize(260, 260)
        self.logo_container.setStyleSheet("""
            QLabel {
                background-color: #1e90ff;
                border-radius: 130px;
            }
        """)

        icon = qta.icon('fa5s.robot', color='white', scale_factor=2.7)
        self.logo = QLabel(self.logo_container)
        pix = icon.pixmap(125, 125)
        self.logo.setPixmap(pix)
        self.logo.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.logo.setGeometry(67, 67, 125, 125)

        self.glow_effect = QGraphicsDropShadowEffect()
        self.glow_effect.setBlurRadius(110)
        self.glow_effect.setXOffset(0)
        self.glow_effect.setYOffset(0)
        self.glow_effect.setColor(QColor(30, 144, 255, 240))
        self.logo_container.setGraphicsEffect(self.glow_effect)

        layout.addWidget(self.logo_container, alignment=Qt.AlignmentFlag.AlignCenter)

        self.title = QLabel("SWSU ROBOTICS")
        self.title.setFont(QFont("Segoe UI", 68, QFont.Weight.Bold))
        self.title.setStyleSheet("color: #ffffff; letter-spacing: -3px;")
        self.title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.title)

        subtitle = QLabel("MECHATRONICS RESEARCH & ENGINEERING KIT")
        subtitle.setFont(QFont("Segoe UI", 19, QFont.Weight.Medium))
        subtitle.setStyleSheet("color: #60c0ff;")
        subtitle.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(subtitle)

        version_btn = QPushButton("Версия 1.0.0")
        version_btn.setFixedSize(170, 40)
        version_btn.setStyleSheet(
            "QPushButton { background-color: #1e90ff; color: white; border-radius: 20px; font-size: 15px; font-weight: 600; }")
        layout.addWidget(version_btn, alignment=Qt.AlignmentFlag.AlignCenter)

        line = QFrame()
        line.setFixedWidth(460)
        line.setFixedHeight(2)
        line.setStyleSheet("background-color: #1e90ff; border: none;")
        layout.addWidget(line, alignment=Qt.AlignmentFlag.AlignCenter)

        self.start_btn = QPushButton("НАЧАТЬ РАБОТУ")
        self.start_btn.setFixedHeight(74)
        self.start_btn.setFixedWidth(460)
        self.start_btn.setFont(QFont("Segoe UI", 24, QFont.Weight.Bold))
        self.start_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.start_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #1e90ff, stop:1 #0b7be0);
                color: white;
                border-radius: 37px;
                border: none;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #40b0ff, stop:1 #1e90ff);
            }
            QPushButton:pressed {
                background: #0a5eb8;
            }
        """)
        self.start_btn.clicked.connect(self.start_requested.emit)
        layout.addWidget(self.start_btn, alignment=Qt.AlignmentFlag.AlignCenter)

        footer = QLabel("© 2025 SWSU Robotics Department. All rights reserved.")
        footer.setStyleSheet("color: #5577aa; font-size: 13px;")
        footer.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addStretch()
        layout.addWidget(footer, alignment=Qt.AlignmentFlag.AlignCenter)

        self.animate_logo()

    def animate_logo(self):
        anim = QPropertyAnimation(self.logo_container, b"pos")
        anim.setDuration(1200)
        anim.setStartValue(QPoint(self.logo_container.x(), self.logo_container.y() - 160))
        anim.setEndValue(self.logo_container.pos())
        anim.setEasingCurve(QEasingCurve.Type.OutBack)
        anim.start()
        self.glow_timer = QTimer(self)
        self.glow_timer.timeout.connect(self.pulse_glow)
        self.glow_value = 70
        self.glow_dir = 1
        self.glow_timer.start(25)

    def pulse_glow(self):
        self.glow_value += 4 * self.glow_dir
        if self.glow_value > 125 or self.glow_value < 55:
            self.glow_dir *= -1
        self.glow_effect.setBlurRadius(self.glow_value)
        self.logo_container.setGraphicsEffect(self.glow_effect)


# ============================================================
# MainWindow
# ============================================================
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("SWSU ROBOTICS")
        self.setMinimumSize(1380, 900)
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint)
        self.showFullScreen()

        central = QWidget()
        self.setCentralWidget(central)
        main_layout = QVBoxLayout(central)
        main_layout.setContentsMargins(0, 0, 0, 0)

        self.stacked = QStackedWidget()
        main_layout.addWidget(self.stacked)

        self.welcome_page = WelcomePage()
        self.welcome_page.start_requested.connect(self.show_patients_list)

        self.patients_page = PatientsListPage(self)
        self.patient_exercises_page = PatientWithExercisesPage(self)
        self.exercise_view_page = ExerciseViewPage()

        self.stacked.addWidget(self.welcome_page)
        self.stacked.addWidget(self.patients_page)
        self.stacked.addWidget(self.patient_exercises_page)
        self.stacked.addWidget(self.exercise_view_page)

        self.patients_page.patient_selected.connect(self.show_patient_exercises)
        self.patient_exercises_page.back_to_list.connect(self.show_patients_list)
        self.patient_exercises_page.exercise_selected.connect(self.show_exercise_view)

        self.progress = QProgressBar()
        self.progress.setVisible(False)
        main_layout.addWidget(self.progress)

        self.thread = None
        self.esc_shortcut = QShortcut(QKeySequence("Esc"), self)
        self.esc_shortcut.activated.connect(self.on_esc_pressed)

        self.stacked.setCurrentWidget(self.welcome_page)

    def show_patients_list(self):
        current_patient = getattr(self.patient_exercises_page, 'current_patient', None)
        if current_patient:
            self.patients_page.set_selected_patient(current_patient)
        else:
            self.patients_page.refresh()
        self.stacked.setCurrentWidget(self.patients_page)

    def show_patient_exercises(self, patient_name):
        self.patient_exercises_page.set_patient(patient_name)
        self.stacked.setCurrentWidget(self.patient_exercises_page)

    def show_exercise_view(self, patient_name, exercise_folder):
        exercise_path = os.path.join(PATIENTS_DIR, patient_name, exercise_folder)
        self.exercise_view_page.set_exercise(patient_name, exercise_folder, exercise_path)
        self.stacked.setCurrentWidget(self.exercise_view_page)

    def add_new_patient(self):
        dlg = PatientInfoDialog()
        if dlg.exec() == QDialog.DialogCode.Accepted:
            name, birth, complaint, height, weight, upper, middle, lower = dlg.get_info()
            if name is None:
                return
            safe_name = re.sub(r'[^a-zA-Zа-яА-ЯёЁ0-9]', '_', name.strip())
            safe_birth = re.sub(r'\D', '', birth.strip())[:8]
            folder_name = f"{safe_name}_{safe_birth}" if safe_birth else safe_name
            original_folder = folder_name
            counter = 1
            while os.path.exists(os.path.join(PATIENTS_DIR, folder_name)):
                folder_name = f"{original_folder}_{counter}"
                counter += 1
            patient_dir = os.path.join(PATIENTS_DIR, folder_name)
            os.makedirs(patient_dir, exist_ok=True)
            info = {
                'name': name,
                'birth_date': birth,
                'complaint': complaint,
                'height_cm': height,
                'weight_kg': weight,
                'upper_link_cm': upper,
                'middle_link_cm': middle,
                'lower_link_cm': lower,
                'created': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            with open(os.path.join(patient_dir, 'info.txt'), 'w', encoding='utf-8') as f:
                json.dump(info, f, ensure_ascii=False, indent=2)
            QMessageBox.information(self, "Готово", f"Пациент «{name}» успешно создан.")
            self.patients_page.refresh()

    def edit_patient(self, patient_folder):
        patient_dir = os.path.join(PATIENTS_DIR, patient_folder)
        info_path = os.path.join(patient_dir, 'info.txt')
        name = ""
        birth = ""
        complaint = ""
        height = ""
        weight = ""
        upper = ""
        middle = ""
        lower = ""
        if os.path.exists(info_path):
            try:
                with open(info_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    name = data.get('name', '')
                    birth = data.get('birth_date', '')
                    complaint = data.get('complaint', '')
                    height = data.get('height_cm', '')
                    weight = data.get('weight_kg', '')
                    upper = data.get('upper_link_cm', '')
                    middle = data.get('middle_link_cm', '')
                    lower = data.get('lower_link_cm', '')
            except:
                pass
        dlg = PatientInfoDialog(name, birth, complaint, height, weight, upper, middle, lower)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            new_name, new_birth, new_complaint, new_height, new_weight, new_upper, new_middle, new_lower = dlg.get_info()
            if new_name is None:
                return
            info = {
                'name': new_name,
                'birth_date': new_birth,
                'complaint': new_complaint,
                'height_cm': new_height,
                'weight_kg': new_weight,
                'upper_link_cm': new_upper,
                'middle_link_cm': new_middle,
                'lower_link_cm': new_lower,
                'edited': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            with open(info_path, 'w', encoding='utf-8') as f:
                json.dump(info, f, ensure_ascii=False, indent=2)
            current_widget = self.stacked.currentWidget()
            if isinstance(current_widget,
                          PatientWithExercisesPage) and current_widget.current_patient == patient_folder:
                current_widget.set_patient(patient_folder)
            self.patients_page.refresh()
            QMessageBox.information(self, "Готово", f"Данные пациента «{new_name}» обновлены.")

    def add_report_to_existing(self, patient_folder):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Выберите файл с данными (.docx)", "",
            "Документы Word (*.docx);;Все файлы (*.*)"
        )
        if not file_path:
            return
        self.progress.setVisible(True)
        self.progress.setRange(0, 0)
        self.setEnabled(False)
        self.thread = ProcessingThread(file_path, patient_folder)
        self.thread.finished_signal.connect(self.on_processing_finished)
        self.thread.start()

    def on_processing_finished(self, success, msg, patient_dir):
        self.progress.setVisible(False)
        self.setEnabled(True)
        if success:
            patient_name = patient_dir.split(os.sep)[-1]
            current_widget = self.stacked.currentWidget()
            if isinstance(current_widget, PatientWithExercisesPage) and current_widget.current_patient == patient_name:
                current_widget.set_patient(patient_name)
            self.patients_page.refresh()
            QMessageBox.information(self, "Успех", msg)
        else:
            QMessageBox.critical(self, "Ошибка", msg)

    def on_esc_pressed(self):
        current = self.stacked.currentWidget()
        if current == self.exercise_view_page:
            self.show_patient_exercises_page()
        elif current == self.patient_exercises_page:
            self.show_patients_list()

    def show_patient_exercises_page(self):
        self.stacked.setCurrentWidget(self.patient_exercises_page)

    def closeEvent(self, event):
        event.accept()


# ============================================================
# ЗАПУСК
# ============================================================
if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyle("Fusion")

    palette = QPalette()
    palette.setColor(QPalette.ColorRole.Window, QColor(10, 20, 40))
    palette.setColor(QPalette.ColorRole.WindowText, QColor(255, 255, 255))
    app.setPalette(palette)

    window = MainWindow()
    sys.exit(app.exec())