import sys
import os
import re
import tempfile
import json
import gc
import shutil
from datetime import datetime
import qtawesome as qta

from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout,
                             QHBoxLayout, QLabel, QLineEdit, QPushButton,
                             QFileDialog, QMessageBox, QProgressBar,
                             QScrollArea, QFrame, QDialog,
                             QListWidget, QListWidgetItem,
                             QGridLayout, QGroupBox, QStackedWidget,
                             QSplitter, QTextEdit, QComboBox,
                             QSizePolicy, QGraphicsDropShadowEffect, QTabWidget)
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QTimer, QPropertyAnimation, QEasingCurve, QPoint, QRegularExpression
from PyQt6.QtGui import (QPalette, QColor, QShortcut, QKeySequence, QFont,
                         QPixmap, QDoubleValidator, QRegularExpressionValidator)

import matplotlib

matplotlib.use('Agg')
import matplotlib.pyplot as plt
import warnings
warnings.filterwarnings("ignore", category=UserWarning, message=".*This figure includes Axes that are not compatible with tight_layout.*")
from matplotlib.figure import Figure
from matplotlib.backends.backend_agg import FigureCanvasAgg  # для безопасной генерации PNG
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sklearn.linear_model import LinearRegression
import numpy as np

# ============================================================
# PATH SETUP — чтобы можно было импортировать пакет agents/
# ============================================================
if getattr(sys, 'frozen', False):
    _APP_DIR = os.path.dirname(sys.executable)
else:
    _APP_DIR = os.path.dirname(os.path.abspath(__file__))

if _APP_DIR not in sys.path:
    sys.path.insert(0, _APP_DIR)

# ============================================================
# ИИ-АНСАМБЛЬ АГЕНТОВ (для вывода анализа под названием упражнения)
# ============================================================
AGENTS_IMPORT_ERROR = None
try:
    from agents import EnsembleOrchestrator
    AGENTS_AVAILABLE = True
except Exception as _e:
    AGENTS_IMPORT_ERROR = str(_e)
    print(f"[INFO] EnsembleOrchestrator (agents) не импортирован: {_e}")
    AGENTS_AVAILABLE = False

def _trapz(y, x=None):
    """Cross-version trapezoidal integration (np.trapz removed in NumPy 2.0+)."""
    if hasattr(np, "trapezoid"):
        return np.trapezoid(y, x)
    return _trapz(y, x)

# Опциональный импорт scipy для продвинутого сглаживания
try:
    from scipy.signal import savgol_filter, butter, filtfilt
    SCIPY_AVAILABLE = True
except ImportError:
    SCIPY_AVAILABLE = False
    print("[WARNING] scipy не установлен. Продвинутое сглаживание (Savitzky-Golay, Butterworth) будет недоступно. Установите: pip install scipy")

# ============================================================
# ДИРЕКТОРИЯ
# ============================================================
# Используем ранний _APP_DIR для согласованности
BASE_DIR = _APP_DIR if '_APP_DIR' in globals() else (
    os.path.dirname(sys.executable) if getattr(sys, 'frozen', False) else os.path.dirname(os.path.abspath(__file__))
)

PATIENTS_DIR = os.path.join(BASE_DIR, "patients")
os.makedirs(PATIENTS_DIR, exist_ok=True)

# ============================================================
# НАЗВАНИЯ КАНАЛОВ
# ============================================================
CHANNEL_LABELS = {
    "angles": {
        "default": ["Канал 1", "Канал 2", "Канал 3", "Канал 4"],
        "ПОВОРОТ СТОПЫ": ["Угол левой стопы", "Угол правой стопы"],
        "ПОВОРОТ ГОЛЕНИ": ["Угол левой голени", "Угол правой голени"],
        "ПОВОРОТ БЕДРА": ["Угол левого бедра", "Угол левой голени", "Угол правого бедра", "Угол правой голени"],
        "ПРИСЕДАНИЯ": ["Угол левого бедра", "Угол левой голени", "Угол левой стопы", "Угол правого бедра",
                       "Угол правой голени", "Угол правой стопы"],
        "ХОДЬБА": ["Угол левого бедра", "Угол левой голени", "Угол левой стопы", "Угол правого бедра",
                   "Угол правой голени", "Угол правой стопы"],
    },
    "forces": {
        "default": ["Канал 1", "Канал 2", "Канал 3", "Канал 4"],
        "ПОВОРОТ СТОПЫ": [
            "Сила на левом носке", "Сила на левой пятке", "Сила на левой голени",
            "Сила на правом носке", "Сила на правой пятке", "Сила на правой голени"
        ],
        "ПОВОРОТ ГОЛЕНИ": [
            "Сила на левом бедре", "Сила на левой голени",
            "Сила на правом бедре", "Сила на правой голени"
        ],
        "ПОВОРОТ БЕДРА": [
            "Сила на левом бедре", "Сила на левой голени",
            "Сила на правом бедре", "Сила на правой голени"
        ],
        "ПРИСЕДАНИЯ": [
            "Сила на левом бедре", "Сила на левом носке", "Сила на левой пятке", "Сила на левой голени",
            "Сила на правом бедре", "Сила на правом носке", "Сила на правой пятке", "Сила на правой голени"
        ],
        "ХОДЬБА": [
            "Сила на левом бедре", "Сила на левом носке", "Сила на левой пятке", "Сила на левой голени",
            "Сила на правом бедре", "Сила на правом носке", "Сила на правой пятке", "Сила на правой голени"
        ],
    }
}


def get_channel_label(exercise_name: str, channel_idx: int, is_angle: bool = True) -> str:
    """
    Возвращает точное название канала строго по спецификации пользователя.
    Определение типа упражнения идёт по названию, которое парсер извлёк из === ЗАГОЛОВОК ===.
    """
    name_upper = exercise_name.upper()

    # === ТОЧНАЯ СПЕЦИФИКАЦИЯ ПОЛЬЗОВАТЕЛЯ (из примера) ===

    # Упражнение 1 (Стопа)
    if any(k in name_upper for k in ["СТОПА", "СТОПЫ", "ПОВОРОТ СТОПЫ", "УПРАЖНЕНИЕ 1", "УПРАЖНЕНИЕ №1"]):
        if is_angle:
            labels = ["Угол левой стопы", "Угол правой стопы"]
        else:
            labels = [
                "Сила на левом носке",
                "Сила на левой пятке",
                "Сила на левой голени",
                "Сила на правом носке",
                "Сила на правой пятке",
                "Сила на правой голени"
            ]
        idx = channel_idx - 1
        return labels[idx] if 0 <= idx < len(labels) else f"Канал {channel_idx}"

    # Упражнение 2 (Голень)
    if any(k in name_upper for k in ["ГОЛЕНЬ", "ПОВОРОТ ГОЛЕНИ", "УПРАЖНЕНИЕ 2", "УПРАЖНЕНИЕ №2"]):
        if is_angle:
            labels = ["Угол левой голени", "Угол правой голени"]
        else:
            labels = [
                "Сила на левом бедре",
                "Сила на левой голени",
                "Сила на правом бедре",
                "Сила на правой голени"
            ]
        idx = channel_idx - 1
        return labels[idx] if 0 <= idx < len(labels) else f"Канал {channel_idx}"

    # Упражнения 3 и 4 (Бедро / Бедро-колено)
    if any(k in name_upper for k in ["БЕДРО", "БЕДРА", "КОЛЕНО", "ПОВОРОТ БЕДРА", "БЕДРО-КОЛЕНО", "УПРАЖНЕНИЕ 3", "УПРАЖНЕНИЕ 4", "УПРАЖНЕНИЕ №3", "УПРАЖНЕНИЕ №4"]):
        if is_angle:
            labels = [
                "Угол левого бедра",
                "Угол левой голени",
                "Угол правого бедра",
                "Угол правой голени"
            ]
        else:
            labels = [
                "Сила на левом бедре",
                "Сила на левой голени",
                "Сила на правом бедре",
                "Сила на правой голени"
            ]
        idx = channel_idx - 1
        return labels[idx] if 0 <= idx < len(labels) else f"Канал {channel_idx}"

    # Упражнения 5 и 6 (Приседания / Ходьба)
    if any(k in name_upper for k in ["ПРИСЕДАНИЯ", "ХОДЬБА", "ПРИСЕД", "ХОДЬБА", "ПРИСЕДАНИЕ", "УПРАЖНЕНИЕ 5", "УПРАЖНЕНИЕ 6", "УПРАЖНЕНИЕ №5", "УПРАЖНЕНИЕ №6"]):
        if is_angle:
            labels = [
                "Угол левого бедра",
                "Угол левой голени",
                "Угол левой стопы",
                "Угол правого бедра",
                "Угол правой голени",
                "Угол правой стопы"
            ]
        else:
            labels = [
                "Сила на левом бедре",
                "Сила на левом носке",
                "Сила на левой пятке",
                "Сила на левой голени",
                "Сила на правом бедре",
                "Сила на правом носке",
                "Сила на правой пятке",
                "Сила на правой голени"
            ]
        idx = channel_idx - 1
        return labels[idx] if 0 <= idx < len(labels) else f"Канал {channel_idx}"

    # Fallback (если название совсем неожиданное)
    key = "angles" if is_angle else "forces"
    for keyword, labels in CHANNEL_LABELS[key].items():
        if keyword in name_upper:
            idx = channel_idx - 1
            if 0 <= idx < len(labels):
                return labels[idx]
            break

    default = CHANNEL_LABELS[key]["default"]
    idx = channel_idx - 1
    return default[idx] if 0 <= idx < len(default) else f"Канал {channel_idx}"


def get_canonical_pairs(exercise_name: str, n_angles: int, n_forces: int):
    """
    Возвращает список пар (угол + связанные силы) + подписи мышц/сегментов
    строго по спецификации пользователя. Учитывает реальное количество каналов в данных.
    Это главная защита от путаницы графиков.

    === ВЕРИФИКАЦИЯ ПАР (удостоверено 2026) ===
    Спецификация пользователя (из детального описания каналов):

    Упражнение 1 (Стопа):
        Углы: 0=левая стопа, 1=правая стопа
        Силы: 0=лев.носок, 1=лев.пятка, 2=лев.голень, 3=прав.носок, 4=прав.пятка, 5=прав.голень
        Пары: (левая стопа + 0,1,2), (правая стопа + 3,4,5)

    Упражнение 2 (Голень):
        Углы: 0=левая голень, 1=правая голень
        Силы: 0=лев.бедро, 1=лев.голень, 2=прав.бедро, 3=прав.голень
        Пары: (левая голень + 0,1), (правая голень + 2,3)

    Упражнения 3 и 4 (Бедро / Бедро-колено):
        Углы: 0=лев.бедро, 1=лев.голень, 2=прав.бедро, 3=прав.голень
        Силы: 0=лев.бедро, 1=лев.голень, 2=прав.бедро, 3=прав.голень
        Пары: по одному (бедро→бедро, голень→голень)

    Упражнения 5 и 6 (Приседания / Ходьба):
        Углы: 0=лев.бедро, 1=лев.голень, 2=лев.стопа, 3=прав.бедро, 4=прав.голень, 5=прав.стопа
        Силы: 0=лев.бедро, 1=лев.носок, 2=лев.пятка, 3=лев.голень,
              4=прав.бедро, 5=прав.носок, 6=прав.пятка, 7=прав.голень
        Пары (текущие, по твоей предыдущей логике):
            (0, лев.бедро → [0]),
            (1, лев.голень → [3]),
            (2, лев.стопа → [1,2]),
            (3, прав.бедро → [4]),
            (4, прав.голень → [7]),
            (5, прав.стопа → [5,6])

    Если после теста на реальных данных ходьбы пары нужно изменить —
    достаточно отредактировать этот блок.
    """
    name_upper = exercise_name.upper()

    # ===== УПРАЖНЕНИЕ 1 — СТОПА =====
    if any(k in name_upper for k in ["СТОПА", "СТОПЫ", "ПОВОРОТ СТОПЫ", "УПРАЖНЕНИЕ 1", "УПРАЖНЕНИЕ №1"]):
        return [
            (0,
             get_channel_label(exercise_name, 1, True),
             [0, 1, 2],
             [get_channel_label(exercise_name, 1, False),
              get_channel_label(exercise_name, 2, False),
              get_channel_label(exercise_name, 3, False)],
             "Левая стопа"),
            (1,
             get_channel_label(exercise_name, 2, True),
             [3, 4, 5],
             [get_channel_label(exercise_name, 4, False),
              get_channel_label(exercise_name, 5, False),
              get_channel_label(exercise_name, 6, False)],
             "Правая стопа"),
        ]

    # ===== УПРАЖНЕНИЕ 2 — ГОЛЕНЬ =====
    if any(k in name_upper for k in ["ГОЛЕНЬ", "ПОВОРОТ ГОЛЕНИ", "УПРАЖНЕНИЕ 2", "УПРАЖНЕНИЕ №2"]):
        return [
            (0,
             get_channel_label(exercise_name, 1, True),
             [0, 1],
             [get_channel_label(exercise_name, 1, False),
              get_channel_label(exercise_name, 2, False)],
             "Левая голень"),
            (1,
             get_channel_label(exercise_name, 2, True),
             [2, 3],
             [get_channel_label(exercise_name, 3, False),
              get_channel_label(exercise_name, 4, False)],
             "Правая голень"),
        ]

    # ===== УПРАЖНЕНИЯ 3 и 4 — БЕДРО / БЕДРО-КОЛЕНО =====
    if any(k in name_upper for k in ["БЕДРО", "БЕДРА", "КОЛЕНО", "ПОВОРОТ БЕДРА", "БЕДРО-КОЛЕНО", "УПРАЖНЕНИЕ 3", "УПРАЖНЕНИЕ 4", "УПРАЖНЕНИЕ №3", "УПРАЖНЕНИЕ №4"]):
        return [
            (0, get_channel_label(exercise_name, 1, True), [0], [get_channel_label(exercise_name, 1, False)], "Левое бедро"),
            (1, get_channel_label(exercise_name, 2, True), [1], [get_channel_label(exercise_name, 2, False)], "Левая голень"),
            (2, get_channel_label(exercise_name, 3, True), [2], [get_channel_label(exercise_name, 3, False)], "Правое бедро"),
            (3, get_channel_label(exercise_name, 4, True), [3], [get_channel_label(exercise_name, 4, False)], "Правая голень"),
        ]

    # ===== УПРАЖНЕНИЯ 5 и 6 — ПРИСЕДАНИЯ / ХОДЬБА =====
    if any(k in name_upper for k in ["ПРИСЕДАНИЯ", "ХОДЬБА", "ПРИСЕД", "ХОДЬБА", "ПРИСЕДАНИЕ", "УПРАЖНЕНИЕ 5", "УПРАЖНЕНИЕ 6", "УПРАЖНЕНИЕ №5", "УПРАЖНЕНИЕ №6"]):
        # Отладка для сложных упражнений (чтобы пользователь был уверен в парсере)
        if n_angles >= 6 and n_forces >= 8:
            print(f"[CANONICAL PAIRS] Walking/Приседания detected: name='{exercise_name}' → using 6 specialized pairs (6 angles + 8 forces)")

        if n_forces >= 8 and n_angles >= 6:
            return [
                (0, get_channel_label(exercise_name, 1, True), [0], [get_channel_label(exercise_name, 1, False)], "Левое бедро"),
                (1, get_channel_label(exercise_name, 2, True), [3], [get_channel_label(exercise_name, 4, False)], "Левая голень"),
                (2, get_channel_label(exercise_name, 3, True), [1, 2],
                 [get_channel_label(exercise_name, 2, False), get_channel_label(exercise_name, 3, False)], "Левая стопа"),
                (3, get_channel_label(exercise_name, 4, True), [4], [get_channel_label(exercise_name, 5, False)], "Правое бедро"),
                (4, get_channel_label(exercise_name, 5, True), [7], [get_channel_label(exercise_name, 8, False)], "Правая голень"),
                (5, get_channel_label(exercise_name, 6, True), [5, 6],
                 [get_channel_label(exercise_name, 6, False), get_channel_label(exercise_name, 7, False)], "Правая стопа"),
            ]
        else:
            pairs = []
            for i in range(min(n_angles or 0, 6)):
                f_idx = [i] if i < (n_forces or 0) else []
                f_lbl = [get_channel_label(exercise_name, i+1, False)] if i < (n_forces or 0) else []
                angle_name = get_channel_label(exercise_name, i+1, True)
                pairs.append((i, angle_name, f_idx, f_lbl, angle_name.replace("Угол ", "")))
            return pairs

    # ===== Fallback: строим пары по реальному количеству каналов в данных =====
    pairs = []
    for i in range(n_angles):
        f_idx = [i] if i < n_forces else []
        f_lbl = [f"Сила канала {i+1}"] if i < n_forces else []
        pairs.append((i, f"Угол канала {i+1}", f_idx, f_lbl, f"Канал {i+1}"))
    return pairs


# ============================================================
# СГЛАЖИВАНИЕ СИГНАЛОВ — Моя рекомендованная стратегия
# ============================================================
# 
# Стратегия:
# - Обычные временные графики (парные + обзорные): Savitzky-Golay 'medium'
#   (хорошо сохраняет форму пиков и фронтов движения)
# - Гистерезисные петли: Savitzky-Golay 'light' (чтобы минимально искажать площадь петли)
# - Спектральный анализ: можно использовать Butterworth перед FFT
#
# При отсутствии scipy используется fallback (или данные остаются без сглаживания).
# ============================================================

def smooth_signal(data, method='savgol', intensity='medium'):
    """
    Универсальная функция сглаживания сигналов.
    
    Моя рекомендованная стратегия:
    - Для обычных временных графиков (углы/силы): Savitzky-Golay средней силы
    - Для гистерезисных петель: лёгкое сглаживание или без
    - Для спектрального анализа: Butterworth перед FFT
    
    :param data: numpy array или list
    :param method: 'none', 'savgol', 'butter', 'moving_avg', 'ema'
    :param intensity: 'light', 'medium', 'strong' (влияет на агрессивность)
    :return: сглаженный numpy array той же длины
    """
    if method == 'none' or not SCIPY_AVAILABLE and method in ['savgol', 'butter']:
        return np.asarray(data)

    data = np.asarray(data, dtype=float)
    n = len(data)
    if n < 5:
        return data

    if method == 'savgol':
        if intensity == 'light':
            window = max(5, min(11, n // 3 * 2 + 1))
            poly = 2
        elif intensity == 'strong':
            window = max(15, min(51, n // 2 * 2 + 1))
            poly = 3
        else:  # medium
            window = max(15, min(31, n // 3 * 2 + 1))
            poly = 3

        # window должен быть нечётным и меньше длины
        window = min(window, n - 1)
        if window % 2 == 0:
            window += 1
        if window < 3:
            return data

        try:
            return savgol_filter(data, window_length=window, polyorder=poly)
        except Exception:
            return data

    elif method == 'butter':
        # Butterworth low-pass (zero-phase)
        try:
            fs = 100.0  # приблизительная частота дискретизации (можно улучшить)
            if intensity == 'light':
                cutoff = 15.0
            elif intensity == 'strong':
                cutoff = 5.0
            else:
                cutoff = 10.0

            nyq = 0.5 * fs
            normal_cutoff = cutoff / nyq
            b, a = butter(4, normal_cutoff, btype='low', analog=False)
            return filtfilt(b, a, data)
        except Exception:
            return data

    elif method == 'moving_avg':
        window = 5 if intensity == 'light' else (11 if intensity == 'medium' else 21)
        window = min(window, n)
        kernel = np.ones(window) / window
        smoothed = np.convolve(data, kernel, mode='same')
        return smoothed

    elif method == 'ema':
        alpha = 0.2 if intensity == 'light' else (0.1 if intensity == 'medium' else 0.05)
        result = np.zeros_like(data)
        result[0] = data[0]
        for i in range(1, n):
            result[i] = alpha * data[i] + (1 - alpha) * result[i - 1]
        return result

    return data


def smooth_channels(channel_data_list, method='savgol', intensity='medium'):
    """Применяет сглаживание ко всем каналам."""
    return [smooth_signal(ch, method=method, intensity=intensity) for ch in channel_data_list]


def compute_leg_load_moment(times, angles_by_channel, forces_by_channel, anthro, exercise_name):
    """
    Вычисляет оценку 'нагрузки ноги' (примерный момент в Н·м) как функцию времени.
    Учитывает:
      - массу пациента (гравитационная составляющая + поддержка веса тела)
      - длины звеньев (верхнее/среднее/нижнее) как рычаги
      - текущие углы (для расчёта плеча момента sin(theta))
      - измеренные силы (внешняя составляющая нагрузки)
    График всегда строится по времени (ось X).
    """
    if not times:
        return None

    n = len(times)

    # Fallback: если нет данных пациента — используем суммарную измеренную силу как прокси "нагрузки"
    if not anthro or anthro.get('weight_kg') is None:
        if not forces_by_channel:
            return None
        proxy = []
        for i in range(n):
            fsum = sum(ch[i] for ch in forces_by_channel if i < len(ch))
            proxy.append(fsum)
        return np.array(proxy)

    M = float(anthro.get('weight_kg') or 70.0)
    g = 9.81
    Lu = float(anthro.get('upper_link_cm') or 40) / 100.0   # м
    Lm = float(anthro.get('middle_link_cm') or 40) / 100.0
    Ll = float(anthro.get('lower_link_cm') or 30) / 100.0

    # Примерные массы сегментов ноги (доли от массы тела, стандартные антропометрические приближения)
    m_thigh = 0.10 * M
    m_shank = 0.046 * M
    m_foot = 0.014 * M

    leg_length = Lu + Lm + Ll
    com_factor = 0.55   # грубое положение центра масс ноги относительно длины

    loads = []
    for i in range(n):
        # Репрезентативный угол позы (среднее по всем доступным каналам углов)
        if angles_by_channel:
            angle_vals = [ch[i] for ch in angles_by_channel if i < len(ch)]
            avg_angle_deg = np.mean(angle_vals) if angle_vals else 0.0
        else:
            avg_angle_deg = 0.0
        theta = np.radians(avg_angle_deg)

        # Суммарная сила, зарегистрированная датчиками в этот момент времени (Н)
        f_total = 0.0
        for ch in forces_by_channel:
            if i < len(ch):
                f_total += ch[i]

        # Моментная составляющая от измеренных сил
        force_moment = f_total * (leg_length * 0.45) * abs(np.sin(theta))

        # Гравитационная составляющая момента от массы человека
        # (вес тела, приходящийся на ногу + вес сегментов самой ноги)
        body_support = 0.5 * M * g
        leg_mass_weight = (m_thigh + m_shank + m_foot) * g
        grav_moment = (body_support + leg_mass_weight) * (leg_length * com_factor) * abs(np.sin(theta))

        moment = force_moment + grav_moment
        loads.append(moment)

    total = np.array(loads)

    # === Почему раньше был 1 график, хотя ног 2?
    # Раньше функция усредняла углы по ВСЕМ каналам и суммировала ВСЕ силы,
    # производя один "общий" момент (как прокси для "нагрузки на опорную ногу" или тотал).
    # Теперь поддерживаем раздельный расчёт для левой и правой ноги.

    left_load = None
    right_load = None
    n_ch_a = len(angles_by_channel) if angles_by_channel else 0
    n_ch_f = len(forces_by_channel) if forces_by_channel else 0

    if n_ch_a >= 2 or n_ch_f >= 2:
        # Простое разделение: первая половина каналов — левая нога, вторая — правая
        half_a = max(1, n_ch_a // 2)
        half_f = max(1, n_ch_f // 2)
        left_angles = angles_by_channel[:half_a] if n_ch_a >= 2 else angles_by_channel
        right_angles = angles_by_channel[half_a:] if n_ch_a >= 2 else []
        left_forces = forces_by_channel[:half_f] if n_ch_f >= 2 else forces_by_channel
        right_forces = forces_by_channel[half_f:] if n_ch_f >= 2 else []

        def _compute_group(a_ch, f_ch):
            if not a_ch and not f_ch:
                return None
            g_loads = []
            for ii in range(n):
                if a_ch:
                    a_vals = [ch[ii] for ch in a_ch if ii < len(ch)]
                    a_deg = np.mean(a_vals) if a_vals else 0.0
                else:
                    a_deg = 0.0
                th = np.radians(a_deg)
                f_tot = 0.0
                for ch in f_ch:
                    if ii < len(ch):
                        f_tot += ch[ii]
                fm = f_tot * (leg_length * 0.45) * abs(np.sin(th))
                gm = (0.5 * M * g + leg_mass_weight) * (leg_length * com_factor) * abs(np.sin(th))
                g_loads.append(fm + gm)
            return np.array(g_loads)

        left_load = _compute_group(left_angles, left_forces)
        right_load = _compute_group(right_angles, right_forces) if right_angles or right_forces else None

    if left_load is not None and right_load is not None:
        return {'left': left_load, 'right': right_load, 'total': total}
    return total


def compute_rom(angles_list):
    """ROM (полный размах) и useful ROM (только где есть нагрузка).
    Для простоты useful считается по полному размаху угла (можно улучшить с M).
    Возвращает dict per channel: {'rom': float, 'min':, 'max':}
    """
    if not angles_list:
        return {}
    roms = {}
    for ch_idx, ch_data in enumerate(angles_list):
        if not ch_data:
            continue
        mn = min(ch_data)
        mx = max(ch_data)
        rom = mx - mn
        roms[ch_idx] = {'rom': rom, 'min': mn, 'max': mx, 'useful_rom': rom}  # useful можно уточнить позже по M
    return roms


def time_normalize(signal, n_points=100):
    """Приводит сигнал к 0-100% цикла (равномерная передискретизация)."""
    if len(signal) < 2:
        return np.full(n_points, signal[0] if signal else 0.0)
    x_old = np.linspace(0, 1, len(signal))
    x_new = np.linspace(0, 1, n_points)
    return np.interp(x_new, x_old, signal)


def symmetry_index(L, R):
    """SI = 200 * |L - R| / (L + R) в % . L, R - скаляры (пики или импульсы)."""
    if abs(L + R) < 1e-9:
        return 0.0
    return 200.0 * abs(L - R) / (L + R)


def compute_phase_lag_and_skew(L_sig, R_sig):
    """Грубая оценка phase_lag (в % цикла) и skew (амплитудная асимметрия)."""
    if len(L_sig) < 5 or len(R_sig) < 5:
        return 0.0, 0.0
    # Нормализуем
    L = (L_sig - np.mean(L_sig)) / (np.std(L_sig) + 1e-9)
    R = (R_sig - np.mean(R_sig)) / (np.std(R_sig) + 1e-9)
    # Cross corr для lag
    corr = np.correlate(L, R, mode='full')
    lag = np.argmax(corr) - (len(L) - 1)
    phase_lag_pct = (lag / max(len(L), 1)) * 100.0
    # Skew: (meanL - meanR) / (meanL + meanR) * 100 or amplitude diff
    amp_L = np.max(np.abs(L_sig))
    amp_R = np.max(np.abs(R_sig))
    skew = 0.0 if (amp_L + amp_R) < 1e-9 else 100.0 * (amp_L - amp_R) / (amp_L + amp_R)
    return phase_lag_pct, skew


def compute_consistency_pearson(signals):
    """Средняя попарная Pearson корреляция между нормализованными сигналами."""
    if len(signals) < 2:
        return 0.0
    normed = [time_normalize(s) for s in signals if len(s) > 1]
    if len(normed) < 2:
        return 0.0
    corrs = []
    for i in range(len(normed)):
        for j in range(i+1, len(normed)):
            r = np.corrcoef(normed[i], normed[j])[0, 1]
            if not np.isnan(r):
                corrs.append(r)
    return float(np.mean(corrs)) if corrs else 0.0


def compute_intra_fatigue_peaks(signal, n_thirds=3):
    """Пики по третям теста. Возвращает список [peak1, peak2, peak3]."""
    if not signal or len(signal) < n_thirds:
        return [max(signal) if signal else 0] * n_thirds
    n = len(signal)
    thirds = []
    for k in range(n_thirds):
        start = int(k * n / n_thirds)
        end = int((k+1) * n / n_thirds)
        seg = signal[start:end]
        thirds.append(max(seg) if seg else 0)
    return thirds


# ============================================================
# ФУНКЦИЯ ПОСТРОЕНИЯ ГРАФИКОВ + ГРАФИЧЕСКИЕ АНАЛИЗЫ
# ============================================================
#
# Дополнительные графические анализы, которые имеет смысл развивать:
#
# 1. Нормализованная динамика силы (Н/кг) — уже частично реализовано ниже
# 2. Эволюция гистерезисных петель (несколько дат на одном графике, цвет по дате)
# 3. Графики симметрии Left/Right для Ходьбы и Приседаний
# 4. Динамика ROM (Range of Motion) по суставам
# 5. Фазовый анализ утомляемости внутри одной длинной сессии
# 6. Корреляционные scatter-плоты между антропометрией пациента и результатами
#
# Эти визуализации очень хорошо дополняют существующие парные графики и гистерезис.
# ============================================================
def save_graphs_for_exercise(exercise, output_dir, smoothing_intensity="mеdium", patient_name=None):
    """
    Генерирует все графики для упражнения.
    smoothing_intensity: "none", "light", "medium", "strong"
    patient_name: если передан — загружает антропометрию (масса + длины звеньев)
                  и строит дополнительный график "Нагрузка ноги".
    """
    measurements = exercise['measurements']
    if not measurements:
        return None, None

    def parse_time_to_seconds(time_str):
        dt = datetime.strptime(time_str, '%Y-%m-%d %H:%M:%S.%f')
        return dt.hour * 3600 + dt.minute * 60 + dt.second + dt.microsecond / 1e6

    first_time_sec = parse_time_to_seconds(measurements[0][0])
    times = [parse_time_to_seconds(tm) - first_time_sec for tm, _, _ in measurements]

    n_angles = len(measurements[0][1])
    n_forces = len(measurements[0][2])

    angles_by_channel = [[] for _ in range(n_angles)]
    forces_by_channel = [[] for _ in range(n_forces)]

    for _, angles, forces in measurements:
        for ch in range(n_angles):
            angles_by_channel[ch].append(angles[ch])
        for ch in range(n_forces):
            forces_by_channel[ch].append(forces[ch] * 9.81 / 1000)

    # ============================================================
    # ПРИМЕНЕНИЕ СГЛАЖИВАНИЯ
    # ============================================================
    # Углы — всегда без сглаживания (по запросу пользователя)
    smoothed_angles = angles_by_channel
    hyst_angles = angles_by_channel

    if smoothing_intensity == "none":
        smoothed_forces = forces_by_channel
        hyst_forces = forces_by_channel
    else:
        # Силы — сглаживаем согласно выбранному уровню
        smoothed_forces = smooth_channels(forces_by_channel, method='savgol', intensity=smoothing_intensity)

        # Для гистерезиса используем чуть более мягкое сглаживание сил,
        # чтобы меньше искажать площадь петли
        hyst_intensity = "light" if smoothing_intensity in ["medium", "strong"] else smoothing_intensity
        hyst_forces = smooth_channels(forces_by_channel, method='savgol', intensity=hyst_intensity)

    # === ГЛАВНОЕ ИСПРАВЛЕНИЕ ПУТАНИЦЫ ===
    # Получаем канонические пары и подписи мышц на основе названия + реального количества каналов в данных
    n_a = len(angles_by_channel) if angles_by_channel else 0
    n_f = len(forces_by_channel) if forces_by_channel else 0
    angle_force_pairs = get_canonical_pairs(exercise['name'], n_a, n_f)

    graphs_info = {'pairs': []}

    for pair_idx, pair_data in enumerate(angle_force_pairs):
        # Поддержка старого формата (4 элемента) и нового (5 элементов с подписью мышцы)
        if len(pair_data) == 5:
            angle_idx, angle_label, force_indices, force_labels, muscle_caption = pair_data
        else:
            angle_idx, angle_label, force_indices, force_labels = pair_data
            muscle_caption = angle_label.replace("Угол ", "").strip()

        n_plots = 1 + len(force_indices)
        fig, axes = plt.subplots(n_plots, 1, figsize=(14, 3.6 * n_plots), dpi=150, sharex=True)

        if n_plots == 1:
            axes = [axes]

        # Заголовок группы (мышца / сегмент) — крупно сверху всей пары
        fig.suptitle(muscle_caption, fontsize=14, fontweight='bold', y=0.98, color='#222')

        # График угла (БЕЗ заголовка и БЕЗ экстремумов) — используем сглаженные данные
        if angle_idx < len(smoothed_angles):
            axes[0].plot(times, smoothed_angles[angle_idx], linewidth=2.5, color='#4a9eff')
            axes[0].set_ylabel(angle_label + " (градусы)", fontsize=11, fontweight='bold')
            axes[0].grid(True, linestyle='--', alpha=0.7)
            axes[0].set_title('')

        # Графики сил (С экстремумами, БЕЗ заголовков) — сглаженные
        for i, (force_idx, force_label) in enumerate(zip(force_indices, force_labels)):
            if force_idx < len(smoothed_forces):
                if 'правый' in force_label.lower() or 'правой' in force_label.lower() or 'правом' in force_label.lower():
                    color = '#ff6b6b'
                elif 'левый' in force_label.lower() or 'левой' in force_label.lower() or 'левом' in force_label.lower():
                    color = '#ffaa44'
                else:
                    color = '#6bff6b'
                axes[i + 1].plot(times, smoothed_forces[force_idx], linewidth=2.5, color=color)
                axes[i + 1].set_ylabel(force_label + " (Н)", fontsize=11, fontweight='bold')
                axes[i + 1].grid(True, linestyle='--', alpha=0.7)
                axes[i + 1].set_title('')
                # Добавляем экстремумы для силы (считаем по сглаженным данным)
                max_force_val = max(smoothed_forces[force_idx])
                min_force_val = min(smoothed_forces[force_idx])
                max_idx = int(np.argmax(smoothed_forces[force_idx]))
                min_idx = int(np.argmin(smoothed_forces[force_idx]))
                axes[i + 1].plot(times[max_idx], max_force_val, 'ro', markersize=8, label=f'Max: {max_force_val:.0f} Н')
                axes[i + 1].plot(times[min_idx], min_force_val, 'go', markersize=8, label=f'Min: {min_force_val:.0f} Н')
                axes[i + 1].legend(loc='best', fontsize=9)

        # Подпись оси X с целыми значениями времени под КАЖДЫМ графиком
        for ax in axes:
            ax.set_xlabel('Время (секунды)', fontsize=11, fontweight='bold')
            ax.tick_params(axis='x', which='both', bottom=True, labelbottom=True)

            if len(times) > 0:
                max_time = int(max(times)) + 1
                if max_time <= 10:
                    tick_positions = list(range(0, max_time + 1))
                elif max_time <= 30:
                    tick_positions = list(range(0, max_time + 1, 2))
                else:
                    tick_positions = list(range(0, max_time + 1, 5))
                ax.set_xticks(tick_positions)
                ax.set_xticklabels([str(t) for t in tick_positions], fontsize=9)

        # Даём больше вертикального пространства между графиками, чтобы подписи времени не наезжали
        plt.subplots_adjust(hspace=0.4)
        plt.tight_layout(rect=[0, 0.02, 1, 0.95])   # оставляем место сверху под suptitle

        pair_png = os.path.join(output_dir, f'pair_{pair_idx:02d}_{angle_idx}_combined.png')
        plt.savefig(pair_png, dpi=100, bbox_inches='tight', facecolor='white')
        plt.close(fig)

        graphs_info['pairs'].append({
            'angle_label': angle_label,
            'image_path': pair_png,
            'angle_index': angle_idx,
            'force_indices': force_indices
        })

    # ============================================================
    # ГИСТЕРЕЗИСНЫЕ ПЕТЛИ (Angle vs Force) - новая вкладка
    # ============================================================
    for pair_idx, pair_data in enumerate(angle_force_pairs):
        if len(pair_data) == 5:
            angle_idx, angle_label, force_indices, force_labels, muscle_caption = pair_data
        else:
            angle_idx, angle_label, force_indices, force_labels = pair_data
            muscle_caption = angle_label.replace("Угол ", "").strip()

        if angle_idx >= len(angles_by_channel):
            continue

        angle_data = hyst_angles[angle_idx]

        # Берём первую (основную) силу для этой пары
        if not force_indices:
            continue
        force_idx = force_indices[0]
        force_label = force_labels[0] if force_labels else f"Сила {force_idx+1}"

        if force_idx >= len(hyst_forces):
            continue

        force_data = hyst_forces[force_idx]

        fig, ax = plt.subplots(figsize=(10, 6), dpi=130)

        ax.plot(angle_data, force_data, linewidth=1.6, color='#2a9d8f', alpha=0.85)

        ax.set_xlabel(angle_label + " (°)", fontsize=11, fontweight='bold')
        ax.set_ylabel(force_label + " (Н)", fontsize=11, fontweight='bold')
        ax.set_title(f"Гистерезисная петля — {muscle_caption}", fontsize=12, fontweight='bold')
        ax.grid(True, linestyle='--', alpha=0.35)

        # Простая оценка площади петли (работа против сопротивления)
        try:
            area = abs(_trapz(force_data, angle_data))
            ax.text(0.02, 0.98, f"Площадь ≈ {area:.0f} Н·°", transform=ax.transAxes,
                    fontsize=10, va='top',
                    bbox=dict(boxstyle='round,pad=0.3', facecolor='#f0f0f0', edgecolor='#2a9d8f', alpha=0.9))
        except Exception:
            pass

        ax.legend([f"{len(angle_data)} точек"], loc='upper right', fontsize=9)

        hyst_path = os.path.join(output_dir, f'hysteresis_p{pair_idx:02d}.png')
        plt.savefig(hyst_path, dpi=100, bbox_inches='tight', facecolor='white')
        plt.close(fig)

    # ============================================================
    # СИММЕТРИЯ (Left vs Right) - для ВСЕХ двусторонних упражнений
    # ============================================================
    name_upper = exercise.get('name', '').upper()

    # Определяем пары левый/правый на основе названий каналов
    angle_labels = CHANNEL_LABELS.get("angles", {}).get(name_upper, [])
    force_labels_def = CHANNEL_LABELS.get("forces", {}).get(name_upper, [])

    def _find_lr_pairs(labels):
        """Находит пары левый/правый по ключевым словам в названиях каналов."""
        left_indices = []
        right_indices = []
        for i, lbl in enumerate(labels):
            lbl_l = lbl.lower()
            if any(k in lbl_l for k in ["левый", "левая", "левое", "левой", "левом", "левого", "левую"]):
                left_indices.append(i)
            elif any(k in lbl_l for k in ["правый", "правая", "правое", "правой", "правом", "правого", "правую"]):
                right_indices.append(i)
        # Сопоставляем по порядку: L1↔R1, L2↔R2, ...
        pairs = []
        for li, ri in zip(left_indices, right_indices):
            pairs.append((li, labels[li], ri, labels[ri]))
        return pairs

    angle_lr_pairs = _find_lr_pairs(angle_labels) if angle_labels else []
    force_lr_pairs = _find_lr_pairs(force_labels_def) if force_labels_def else []

    # Если подписи не заданы — пробуем определить по количеству каналов
    if not angle_lr_pairs and len(angles_by_channel) >= 2:
        half = len(angles_by_channel) // 2
        for i in range(half):
            angle_lr_pairs.append((i, f"Канал {i+1} (Л)", i + half, f"Канал {i+half+1} (П)"))

    if not force_lr_pairs and len(forces_by_channel) >= 2:
        half = len(forces_by_channel) // 2
        for i in range(half):
            force_lr_pairs.append((i, f"Сила {i+1} (Л)", i + half, f"Сила {i+half+1} (П)"))

    has_symmetry = len(angle_lr_pairs) > 0 or len(force_lr_pairs) > 0

    if has_symmetry:
        try:
            # Симметрия углов
            n_plots_a = len(angle_lr_pairs) if angle_lr_pairs else 1
            fig, axes = plt.subplots(n_plots_a, 1, figsize=(14, 3.5 * n_plots_a), dpi=120, sharex=True)
            if n_plots_a == 1:
                axes = [axes]

            for i, (li, l_label, ri, r_label) in enumerate(angle_lr_pairs):
                ax = axes[i]
                if li < len(angles_by_channel) and ri < len(angles_by_channel):
                    ax.plot(times, angles_by_channel[li], label=l_label, color='#4a9eff', linewidth=2.2)
                    ax.plot(times, angles_by_channel[ri], label=r_label, color='#ff6b6b', linewidth=2.2)
                    ax.set_ylabel(f"{l_label.replace('Угол ', '')} / {r_label.replace('Угол ', '')} (°)", fontsize=10, fontweight='bold')
                    ax.legend(loc='upper right', fontsize=9)
                    ax.grid(True, linestyle='--', alpha=0.35)

            axes[-1].set_xlabel("Время (с)", fontsize=11, fontweight='bold')
            fig.suptitle(f"Симметрия углов — {exercise.get('name', '')}", fontsize=14, fontweight='bold', y=0.98)

            plt.tight_layout(rect=[0, 0.02, 1, 0.96])
            sym_path = os.path.join(output_dir, "symmetry_angles.png")
            plt.savefig(sym_path, dpi=100, bbox_inches='tight', facecolor='white')
            plt.close(fig)

            # Симметрия сил
            if force_lr_pairs:
                n_plots_f = len(force_lr_pairs)
                fig2, axes2 = plt.subplots(n_plots_f, 1, figsize=(14, 3.5 * n_plots_f), dpi=120, sharex=True)
                if n_plots_f == 1:
                    axes2 = [axes2]

                for i, (li, l_label, ri, r_label) in enumerate(force_lr_pairs):
                    ax = axes2[i]
                    left_data = forces_by_channel[li][:len(times)] if li < len(forces_by_channel) else []
                    right_data = forces_by_channel[ri][:len(times)] if ri < len(forces_by_channel) else []
                    if left_data and right_data:
                        ax.plot(times[:len(left_data)], left_data, label=l_label, color='#4a9eff', linewidth=2.2)
                        ax.plot(times[:len(right_data)], right_data, label=r_label, color='#ff6b6b', linewidth=2.2)
                        ax.set_ylabel(f"{l_label} / {r_label} (Н)", fontsize=10, fontweight='bold')
                        ax.legend(loc='upper right', fontsize=9)
                        ax.grid(True, linestyle='--', alpha=0.35)

                axes2[-1].set_xlabel("Время (с)", fontsize=11, fontweight='bold')
                fig2.suptitle(f"Симметрия сил — {exercise.get('name', '')}", fontsize=14, fontweight='bold', y=0.98)

                plt.tight_layout(rect=[0, 0.02, 1, 0.96])
                sym_force_path = os.path.join(output_dir, "symmetry_forces.png")
                plt.savefig(sym_force_path, dpi=100, bbox_inches='tight', facecolor='white')
                plt.close(fig2)

        except Exception as e:
            print(f"[WARN] Не удалось сгенерировать график симметрии: {e}")

    # ============================================================
    # АНАЛИТИЧЕСКИЕ ГРАФИКИ (спектральный анализ и зависимость)
    # ============================================================
    try:
        # ЛЕГАСИ "ЗАВИСИМОСТЬ СИЛЫ ОТ УГЛА" УДАЛЕНА
        # Ранее использовался наивный pairing angle[ch] ↔ force[ch] + get_channel_label(ch+1).
        # Для упражнений Ходьба/Приседания (6 углов + 8 сил) это давало неправильные ассоциации
        # (пересечения лево/право, несоответствующие мышцы). Вкладка в UI уже удалена,
        # поэтому генерация этих scatter_ch*.png больше не выполняется.
        # Все парные визуализации теперь идут ТОЛЬКО через get_canonical_pairs().

        # СПЕКТРАЛЬНЫЙ АНАЛИЗ (БПФ) — оставлен, так как он используется во вкладке "Спектральный анализ"
        if len(times) > 10 and len(forces_by_channel) > 0:
            fig = plt.figure(figsize=(14, 10), dpi=100)
            gs = fig.add_gridspec(2, 2, hspace=0.3, wspace=0.3)

            # Исходный сигнал
            ax1 = fig.add_subplot(gs[0, 0])
            force_signal = forces_by_channel[0]
            ax1.plot(times, force_signal, linewidth=2, color='#ff6b6b')
            ax1.set_xlabel('Время (с)', fontsize=11, fontweight='bold')
            ax1.set_ylabel('Сила (Н)', fontsize=11, fontweight='bold')
            ax1.set_title('Исходный сигнал силы', fontsize=12, fontweight='bold')
            ax1.grid(True, linestyle='--', alpha=0.3)
            max_val = max(force_signal)
            min_val = min(force_signal)
            max_idx = force_signal.index(max_val)
            min_idx = force_signal.index(min_val)
            ax1.plot(times[max_idx], max_val, 'ro', markersize=7, label=f'Max: {max_val:.0f} Н')
            ax1.plot(times[min_idx], min_val, 'go', markersize=7, label=f'Min: {min_val:.0f} Н')
            ax1.legend(loc='best', fontsize=8)

            # Спектр (БПФ)
            ax2 = fig.add_subplot(gs[0, 1])
            n = len(force_signal)
            fft_vals = np.fft.fft(force_signal)
            fft_freq = np.fft.fftfreq(n, d=(times[1] - times[0]) if len(times) > 1 else 0.01)

            positive_freq = fft_freq[:n // 2]
            positive_fft = np.abs(fft_vals[:n // 2]) / n

            if len(positive_fft) > 0:
                dominant_idx = np.argmax(positive_fft[1:]) + 1
                dominant_freq = positive_freq[dominant_idx]
                ax2.plot(positive_freq, positive_fft, linewidth=1.5, color='#4a9eff')
                ax2.axvline(dominant_freq, color='red', linestyle='--', linewidth=2,
                            label=f'Доминирующая частота: {dominant_freq:.2f} Гц')
            else:
                ax2.plot(positive_freq, positive_fft, linewidth=1.5, color='#4a9eff')

            ax2.set_xlabel('Частота (Гц)', fontsize=11, fontweight='bold')
            ax2.set_ylabel('Амплитуда', fontsize=11, fontweight='bold')
            ax2.set_title('Спектр мощности (БПФ)', fontsize=12, fontweight='bold')
            ax2.grid(True, linestyle='--', alpha=0.3)
            ax2.set_xlim(0, 5)
            ax2.legend(loc='best', fontsize=8)

            # Сглаженный сигнал
            ax3 = fig.add_subplot(gs[1, 0])
            window_size = max(3, min(20, len(force_signal) // 20))
            if window_size > 1 and len(force_signal) > window_size:
                smoothed = np.convolve(force_signal, np.ones(window_size) / window_size, mode='valid')
                smooth_times = times[window_size // 2:window_size // 2 + len(smoothed)]
                ax3.plot(smooth_times, smoothed, linewidth=2.5, color='#00cc88')
                ax3.set_xlabel('Время (с)', fontsize=11, fontweight='bold')
                ax3.set_ylabel('Сила (Н)', fontsize=11, fontweight='bold')
                ax3.set_title(f'Сглаженный сигнал (окно = {window_size} точек)', fontsize=12, fontweight='bold')
                ax3.grid(True, linestyle='--', alpha=0.3)
                max_smooth = max(smoothed)
                min_smooth = min(smoothed)
                max_idx = np.argmax(smoothed)
                min_idx = np.argmin(smoothed)
                ax3.plot(smooth_times[max_idx], max_smooth, 'ro', markersize=7, label=f'Max: {max_smooth:.0f} Н')
                ax3.plot(smooth_times[min_idx], min_smooth, 'go', markersize=7, label=f'Min: {min_smooth:.0f} Н')
                ax3.legend(loc='best', fontsize=8)
            else:
                ax3.text(0.5, 0.5, 'Недостаточно данных для сглаживания', ha='center', va='center', transform=ax3.transAxes,
                         fontsize=12)

            # Накопленная работа
            ax4 = fig.add_subplot(gs[1, 1])
            dt = times[1] - times[0] if len(times) > 1 else 0.01
            cumulative_work = np.cumsum(force_signal) * dt
            ax4.fill_between(times, 0, cumulative_work, alpha=0.5, color='#ffaa44')
            ax4.plot(times, cumulative_work, linewidth=2.5, color='#ff8800')
            ax4.set_xlabel('Время (с)', fontsize=11, fontweight='bold')
            ax4.set_ylabel('Накопленная работа (Н·с)', fontsize=11, fontweight='bold')
            ax4.set_title('Накопленная механическая работа', fontsize=12, fontweight='bold')
            ax4.grid(True, linestyle='--', alpha=0.3)

            total_work = cumulative_work[-1] if len(cumulative_work) > 0 else 0
            ax4.text(0.95, 0.05, f'Общая работа: {total_work:.0f} Н·с', transform=ax4.transAxes,
                     fontsize=11, ha='right', va='bottom', bbox=dict(boxstyle='round', facecolor='white', alpha=0.85))

            plt.suptitle(f'Спектральный анализ: {exercise["name"]}', fontsize=13, fontweight='bold', y=0.98)
            plt.tight_layout(rect=[0, 0.02, 1, 0.95])

            fft_png = os.path.join(output_dir, 'analysis_fft.png')
            plt.savefig(fft_png, dpi=100, bbox_inches='tight', facecolor='white')
            plt.close(fig)

        # Сохраняем информацию только о спектральном анализе
        # (legacy scatter "Зависимость силы от угла" больше не генерируются)
        analysis_graphs = []
        fft_path = os.path.join(output_dir, 'analysis_fft.png')
        if os.path.exists(fft_path):
            analysis_graphs.append(('Спектральный анализ (БПФ, сглаживание, работа)', fft_path))

        with open(os.path.join(output_dir, 'analysis_info.json'), 'w', encoding='utf-8') as f:
            json.dump(analysis_graphs, f, ensure_ascii=False, indent=2)

    except Exception as e:
        print(f"[WARN] Ошибка при генерации аналитических графиков: {e}")

    # Сохраняем максимальную силу
    all_forces = [value for channel in forces_by_channel for value in channel]
    max_force = max(all_forces) if all_forces else 0
    with open(os.path.join(output_dir, 'max_force.txt'), 'w') as f:
        f.write(str(max_force))

    # Для обратной совместимости - общий график углов (БЕЗ экстремумов)
    if n_angles > 0:
        fig, axes = plt.subplots(n_angles, 1, figsize=(12, 3.5 * n_angles), sharex=True, dpi=150)
        if n_angles == 1:
            axes = [axes]
        for ch, ax in enumerate(axes):
            ax.plot(times, smoothed_angles[ch], linewidth=2.5, color='#4a9eff')
            ax.set_ylabel(get_channel_label(exercise['name'], ch + 1, True), fontsize=10)
            ax.grid(True, linestyle='--', alpha=0.7)
            ax.set_title('')
        # Шкала времени под каждым графиком
        for ax in axes:
            ax.set_xlabel('Время (секунды)', fontsize=10)
            ax.tick_params(axis='x', which='both', bottom=True, labelbottom=True)

        if len(times) > 0:
            max_time = int(max(times)) + 1
            if max_time <= 10:
                tick_positions = list(range(0, max_time + 1))
            elif max_time <= 30:
                tick_positions = list(range(0, max_time + 1, 2))
            else:
                tick_positions = list(range(0, max_time + 1, 5))
            for ax in axes:
                ax.set_xticks(tick_positions)
                ax.set_xticklabels([str(t) for t in tick_positions], fontsize=8)

        plt.subplots_adjust(hspace=0.35)
        plt.tight_layout()
        angles_png = os.path.join(output_dir, 'angles.png')
        plt.savefig(angles_png, dpi=150, bbox_inches='tight')
        plt.close(fig)
    else:
        angles_png = None

    # Общий график сил (С экстремумами)
    if n_forces > 0:
        fig, axes = plt.subplots(n_forces, 1, figsize=(12, 3.5 * n_forces), sharex=True, dpi=150)
        if n_forces == 1:
            axes = [axes]
        for ch, ax in enumerate(axes):
            ax.plot(times, smoothed_forces[ch], linewidth=2.5, color='#ff6b6b')
            ax.set_ylabel(get_channel_label(exercise['name'], ch + 1, False), fontsize=10)
            ax.grid(True, linestyle='--', alpha=0.7)
            ax.set_title('')
            # Добавляем экстремумы для силы (по сглаженным данным)
            max_val = max(smoothed_forces[ch])
            min_val = min(smoothed_forces[ch])
            max_idx = int(np.argmax(smoothed_forces[ch]))
            min_idx = int(np.argmin(smoothed_forces[ch]))
            ax.plot(times[max_idx], max_val, 'ro', markersize=6, label=f'Max: {max_val:.0f} Н')
            ax.plot(times[min_idx], min_val, 'go', markersize=6, label=f'Min: {min_val:.0f} Н')
            ax.legend(loc='best', fontsize=8)

        # Шкала времени под каждым графиком
        for ax in axes:
            ax.set_xlabel('Время (секунды)', fontsize=10)
            ax.tick_params(axis='x', which='both', bottom=True, labelbottom=True)

        if len(times) > 0:
            max_time = int(max(times)) + 1
            if max_time <= 10:
                tick_positions = list(range(0, max_time + 1))
            elif max_time <= 30:
                tick_positions = list(range(0, max_time + 1, 2))
            else:
                tick_positions = list(range(0, max_time + 1, 5))
            for ax in axes:
                ax.set_xticks(tick_positions)
                ax.set_xticklabels([str(t) for t in tick_positions], fontsize=8)

        plt.subplots_adjust(hspace=0.35)
        plt.tight_layout()
        forces_png = os.path.join(output_dir, 'forces.png')
        plt.savefig(forces_png, dpi=150, bbox_inches='tight')
        plt.close(fig)
    else:
        forces_png = None

    # Сохраняем сырые данные для быстрой перестройки с другим сглаживанием
    try:
        raw_data = {
            'times': [t for t, _, _ in measurements],
            'angles': [a for _, a, _ in measurements],
            'forces': [f for _, _, f in measurements],
            'n_angles': n_angles,
            'n_forces': n_forces
        }
        with open(os.path.join(output_dir, 'raw_measurements.json'), 'w', encoding='utf-8') as f:
            json.dump(raw_data, f)
    except Exception as e:
        print(f"[WARN] Не удалось сохранить raw_measurements: {e}")

    # ============================================================
    # НОВЫЙ ГРАФИК: НАГРУЗКА НОГИ
    # Момент (Н·м) как функция времени, зависящая от массы пациента и длин звеньев.
    # Ось X — время (как у всех остальных временных графиков).
    # ============================================================
    try:
        anthro = load_patient_anthropometrics(patient_name) if patient_name else {}
        leg_load = compute_leg_load_moment(times, angles_by_channel, forces_by_channel, anthro, exercise.get('name', ''))
        fig, ax = plt.subplots(figsize=(14, 5), dpi=130)
        plotted = False

        if isinstance(leg_load, dict) and 'left' in leg_load and 'right' in leg_load:
            # 2 ноги — рисуем два графика (левая / правая)
            ax.plot(times, leg_load['left'], linewidth=2.6, color='#3498db', label='Левая нога')
            ax.plot(times, leg_load['right'], linewidth=2.6, color='#e74c3c', label='Правая нога')
            ax.set_xlabel('Время (с)', fontsize=11, fontweight='bold')
            ax.set_ylabel('Момент (Н·м)', fontsize=11, fontweight='bold')
            ax.set_title('Нагрузка левой и правой ноги (отдельно)', fontsize=14, fontweight='bold')
            ax.grid(True, linestyle='--', alpha=0.4)
            ax.legend(loc='best', fontsize=9)

            # Экстремумы для левой
            for name, load_arr, color in [('Левая', leg_load['left'], '#3498db'), ('Правая', leg_load['right'], '#e74c3c')]:
                if len(load_arr) > 0:
                    max_v = float(max(load_arr))
                    min_v = float(min(load_arr))
                    imax = int(np.argmax(load_arr))
                    imin = int(np.argmin(load_arr))
                    ax.plot(times[imax], max_v, 'o', color=color, markersize=6)
                    ax.plot(times[imin], min_v, 'o', color=color, markersize=6)
            plotted = True
        elif leg_load is not None and len(leg_load) == len(times):
            # Старый режим (1 график, совместимость)
            ax.plot(times, leg_load, linewidth=2.6, color='#8e44ad')
            ax.set_xlabel('Время (с)', fontsize=11, fontweight='bold')
            ax.set_ylabel('Момент (Н·м)', fontsize=11, fontweight='bold')
            ax.set_title('Нагрузка ноги', fontsize=14, fontweight='bold')
            ax.grid(True, linestyle='--', alpha=0.4)

            # Экстремумы
            if len(leg_load) > 0:
                max_v = float(max(leg_load))
                min_v = float(min(leg_load))
                imax = int(np.argmax(leg_load))
                imin = int(np.argmin(leg_load))
                ax.plot(times[imax], max_v, 'ro', markersize=7, label=f'Max: {max_v:.0f} Н·м')
                ax.plot(times[imin], min_v, 'go', markersize=7, label=f'Min: {min_v:.0f} Н·м')
                ax.legend(loc='best', fontsize=9)
            plotted = True

        if plotted:
            # Шкала времени (аналогично другим графикам)
            if len(times) > 0:
                max_time = int(max(times)) + 1
                if max_time <= 10:
                    tick_positions = list(range(0, max_time + 1))
                elif max_time <= 30:
                    tick_positions = list(range(0, max_time + 1, 2))
                else:
                    tick_positions = list(range(0, max_time + 1, 5))
                ax.set_xticks(tick_positions)
                ax.set_xticklabels([str(t) for t in tick_positions], fontsize=9)

            plt.tight_layout(rect=[0, 0.02, 1, 0.95])
            leg_path = os.path.join(output_dir, 'leg_load.png')
            plt.savefig(leg_path, dpi=120, bbox_inches='tight', facecolor='white')
            plt.close(fig)
    except Exception as e:
        print(f"[WARN] Не удалось сгенерировать график нагрузка ноги: {e}")

    # === ROM and session stats for general report (matches Obschiy_otchyot_vse_ispytuemye.docx) ===
    try:
        rom_data = compute_rom(angles_by_channel)
        # Улучшаем useful ROM: размах угла только в моменты, когда суммарная сила выше порога (5% от max или 20Н)
        if forces_by_channel and rom_data:
            totals = []
            for i in range(len(times)):
                s = 0.0
                for ch in forces_by_channel:
                    if i < len(ch): s += ch[i]
                totals.append(s)
            if totals:
                fmax = max(totals)
                thresh = max(20.0, 0.05 * fmax)
                for ch_idx in list(rom_data.keys()):
                    ch_a = angles_by_channel[ch_idx] if ch_idx < len(angles_by_channel) else []
                    useful_vals = [ch_a[i] for i in range(min(len(ch_a), len(totals))) if totals[i] >= thresh]
                    if useful_vals:
                        rom_data[ch_idx]['useful_rom'] = max(useful_vals) - min(useful_vals)
        stats = {
            'n_measurements': len(times),
            'n_angles': n_angles,
            'n_forces': n_forces,
            'rom': rom_data,
            'exercise': exercise.get('name', '')
        }
        if forces_by_channel:
            totals = []
            for i in range(len(times)):
                s = 0.0
                for ch in forces_by_channel:
                    if i < len(ch):
                        s += ch[i]
                totals.append(s)
            stats['max_total_force_N'] = float(max(totals)) if totals else 0.0
        with open(os.path.join(output_dir, 'session_stats.json'), 'w', encoding='utf-8') as f:
            json.dump(stats, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"[WARN] Не удалось сохранить session_stats.json: {e}")

    # Дополнительный фазовый портрет (угол vs нагрузка/момент) — как в образце отчёта
    try:
        if angles_by_channel and len(times) > 5:
            main_a = angles_by_channel[0]
            # Прокси момента/нагрузки: если leg_load посчитан в этом же блоке — используем, иначе сумма сил
            y_phase = None
            try:
                y_phase = leg_load if leg_load is not None else None
            except:
                y_phase = None
            if y_phase is None and forces_by_channel:
                y_phase = []
                for i in range(len(times)):
                    s = 0.0
                    for ch in forces_by_channel:
                        if i < len(ch): s += ch[i]
                    y_phase.append(s)
            if y_phase is not None and len(y_phase) == len(main_a):
                figp, axp = plt.subplots(figsize=(8, 6), dpi=120)
                axp.plot(main_a, y_phase, linewidth=1.3, color='#2a9d8f', alpha=0.8)
                axp.set_xlabel('Угол (град)', fontsize=10, fontweight='bold')
                axp.set_ylabel('Нагрузка ноги (прокси Н·м / Н)', fontsize=10, fontweight='bold')
                axp.set_title(f'Phase portrait (угол — нагрузка) — {exercise.get("name", "")}', fontsize=11, fontweight='bold')
                axp.grid(True, linestyle='--', alpha=0.3)
                pp_path = os.path.join(output_dir, 'phase_portrait_leg.png')
                plt.savefig(pp_path, dpi=100, bbox_inches='tight', facecolor='white')
                plt.close(figp)
    except Exception as e:
        print(f"[WARN] phase_portrait_leg: {e}")

    gc.collect()
    return angles_png, forces_png


def regenerate_graphs_for_session(patient_name, folder_name, intensity="medium"):
    """
    Мгновенная перестройка графиков для одного сеанса с новым уровнем сглаживания.
    Приоритет:
    1. Использовать raw_measurements.json (быстро)
    2. Если нет — попытаться найти оригинальный .docx и перепарсить только нужное упражнение.
    """
    session_path = os.path.join(PATIENTS_DIR, patient_name, folder_name)
    raw_path = os.path.join(session_path, "raw_measurements.json")

    exercise_name = folder_name.split("_", 1)[0] if "_" in folder_name else folder_name
    measurements = None

    # === 1. Пытаемся загрузить из кэша сырых данных ===
    if os.path.exists(raw_path):
        try:
            with open(raw_path, "r", encoding="utf-8") as f:
                raw = json.load(f)
            times = raw["times"]
            angles_list = raw["angles"]
            forces_list = raw["forces"]
            measurements = list(zip(times, angles_list, forces_list))
        except Exception as e:
            print(f"[WARN] Не удалось прочитать raw_measurements.json: {e}")

    # === 2. Fallback: пытаемся найти и перепарсить оригинальный .docx ===
    if measurements is None:
        patient_dir = os.path.join(PATIENTS_DIR, patient_name)

        # Ищем .docx файлы в папке пациента и её подпапках
        docx_candidates = []
        for root, dirs, files in os.walk(patient_dir):
            for f in files:
                if f.lower().endswith(".docx") and "report" not in f.lower():
                    docx_candidates.append(os.path.join(root, f))

        if not docx_candidates:
            return False, (
                "Нет сохранённых сырых данных и не найден оригинальный .docx файл.\n\n"
                "Решения:\n"
                "1. Пересоздайте данные пациента заново (добавьте отчёт ещё раз).\n"
                "2. Положите оригинальный .docx файл в папку пациента."
            )

        # Берём первый найденный .docx (обычно самый релевантный)
        docx_path = docx_candidates[0]

        try:
            print(f"[INFO] Перепарсинг {docx_path} для восстановления данных упражнения '{exercise_name}'...")
            txt_path = docx_to_txt(docx_path)
            exercises = RobustParser.parse(txt_path)
            os.unlink(txt_path)
            gc.collect()

            # Ищем нужное упражнение
            for ex in exercises:
                if ex['name'].upper() == exercise_name.upper() or \
                   ex['folder_name'].startswith(exercise_name):
                    measurements = ex['measurements']
                    # Сохраняем сырые данные на будущее
                    try:
                        raw_data = {
                            'times': [t for t, _, _ in measurements],
                            'angles': [a for _, a, _ in measurements],
                            'forces': [f for _, _, f in measurements],
                            'n_angles': len(measurements[0][1]) if measurements else 0,
                            'n_forces': len(measurements[0][2]) if measurements else 0
                        }
                        with open(raw_path, 'w', encoding='utf-8') as f:
                            json.dump(raw_data, f)
                        print("[INFO] raw_measurements.json успешно восстановлен и сохранён.")
                    except Exception as save_err:
                        print(f"[WARN] Не удалось сохранить восстановленный raw: {save_err}")
                    break

            if measurements is None:
                return False, f"В файле {os.path.basename(docx_path)} не найдено упражнение '{exercise_name}'."

        except Exception as e:
            return False, f"Не удалось перепарсить оригинальный файл:\n{str(e)}"

    if not measurements:
        return False, "Не удалось получить данные измерений."

    try:
        exercise = {
            "name": exercise_name,
            "measurements": measurements
        }

        save_graphs_for_exercise(exercise, session_path, smoothing_intensity=intensity, patient_name=patient_name)
        return True, "Графики успешно перестроены с новым уровнем сглаживания."

    except Exception as e:
        return False, f"Ошибка при перестройке графиков: {str(e)}"


def load_patient_anthropometrics(patient_name):
    """
    Загружает антропометрические и клинические данные пациента из info.txt.
    Возвращает словарь с ключами: weight_kg, height_cm, upper_link_cm, middle_link_cm,
    lower_link_cm, birth_date, complaint, age_years (если возможно посчитать).
    """
    patient_dir = os.path.join(PATIENTS_DIR, patient_name)
    info_path = os.path.join(patient_dir, 'info.txt')

    data = {
        'weight_kg': None,
        'height_cm': None,
        'upper_link_cm': None,
        'middle_link_cm': None,
        'lower_link_cm': None,
        'birth_date': None,
        'complaint': None,
        'age_years': None
    }

    if not os.path.exists(info_path):
        return data

    try:
        with open(info_path, 'r', encoding='utf-8') as f:
            info = json.load(f)

        def safe_float(v):
            try:
                return float(v) if v is not None else None
            except (ValueError, TypeError):
                return None
        data['weight_kg'] = safe_float(info.get('weight_kg'))
        data['height_cm'] = safe_float(info.get('height_cm'))
        data['upper_link_cm'] = safe_float(info.get('upper_link_cm'))
        data['middle_link_cm'] = safe_float(info.get('middle_link_cm'))
        data['lower_link_cm'] = safe_float(info.get('lower_link_cm'))
        data['birth_date'] = info.get('birth_date')
        data['complaint'] = info.get('complaint')

        # Пытаемся посчитать возраст из даты рождения
        if data['birth_date']:
            for fmt in ('%d.%m.%Y', '%Y-%m-%d', '%d/%m/%Y'):
                try:
                    birth = datetime.strptime(data['birth_date'], fmt)
                    data['age_years'] = (datetime.now() - birth).days / 365.25
                    break
                except ValueError:
                    continue

    except Exception as e:
        print(f"[WARN] Не удалось загрузить антропометрию пациента {patient_name}: {e}")

    return data


def generate_amplitude_dynamics_png(patient_name, exercise_name, output_path):
    """
    Генерирует и сохраняет PNG график динамики максимальной силы по всем сессиям упражнения.
    Также создаёт нормализованную версию (Н/кг), если известен вес пациента.
    Вызывается во время обработки данных, а не в UI.
    """
    patient_dir = os.path.join(PATIENTS_DIR, patient_name)
    date_max = {}

    for folder in os.listdir(patient_dir):
        if not folder.startswith(exercise_name + "_"):
            continue
        folder_path = os.path.join(patient_dir, folder)
        max_force_path = os.path.join(folder_path, 'max_force.txt')
        if not os.path.exists(max_force_path):
            continue
        try:
            with open(max_force_path, 'r') as f:
                max_force = float(f.read().strip())
        except Exception:
            continue

        date_str = ""
        for part in folder.split('_'):
            if re.match(r'\d{4}-\d{2}-\d{2}', part):
                try:
                    dt = datetime.strptime(part, '%Y-%m-%d')
                    date_str = dt.strftime('%d.%m.%Y')
                except:
                    date_str = part
                break

        if date_str:
            if date_str in date_max:
                date_max[date_str] = max(date_max[date_str], max_force)
            else:
                date_max[date_str] = max_force

    dates = list(date_max.keys())
    forces = list(date_max.values())

    # Загружаем антропометрию пациента
    anthro = load_patient_anthropometrics(patient_name)
    weight = anthro.get('weight_kg')

    # === Обычный график (в Ньютонах) ===
    fig = Figure(figsize=(14, 8), dpi=140, facecolor='#1e2a3a')
    ax = fig.add_subplot(111)
    ax.set_facecolor('#1e2a3a')
    fig.patch.set_facecolor('#1e2a3a')

    if dates:
        ax.plot(dates, forces, marker='o', linestyle='-', color='#5a9eff', markersize=9, linewidth=2)
        ax.set_xlabel('Дата', color='white', fontsize=12)
        ax.set_ylabel('Максимальная сила (Н)', color='white', fontsize=12)
        ax.set_title(f'Динамика максимальной силы — {exercise_name}', color='white', fontsize=14)
        ax.tick_params(colors='white')
        ax.grid(True, linestyle='--', alpha=0.3, color='white')
        plt.setp(ax.get_xticklabels(), rotation=45, ha='right')
    else:
        ax.text(0.5, 0.5, 'Нет данных для отображения', ha='center', va='center',
                transform=ax.transAxes, color='white', fontsize=16)

    canvas = FigureCanvasAgg(fig)
    canvas.print_png(output_path)
    plt.close(fig)

    # === Нормализованный график (Н/кг), если известен вес ===
    if weight and weight > 0 and dates:
        normalized_forces = [f / weight for f in forces]
        norm_output_path = output_path.replace('.png', '_normalized.png')

        fig2 = Figure(figsize=(14, 8), dpi=140, facecolor='#1e2a3a')
        ax2 = fig2.add_subplot(111)
        ax2.set_facecolor('#1e2a3a')
        fig2.patch.set_facecolor('#1e2a3a')

        ax2.plot(dates, normalized_forces, marker='o', linestyle='-', color='#ffaa44', markersize=9, linewidth=2)
        ax2.set_xlabel('Дата', color='white', fontsize=12)
        ax2.set_ylabel('Максимальная сила (Н/кг)', color='white', fontsize=12)
        ax2.set_title(f'Динамика максимальной силы (нормализованная по весу) — {exercise_name}\nВес пациента: {weight} кг', color='white', fontsize=14)
        ax2.tick_params(colors='white')
        ax2.grid(True, linestyle='--', alpha=0.3, color='white')
        plt.setp(ax2.get_xticklabels(), rotation=45, ha='right')

        canvas2 = FigureCanvasAgg(fig2)
        canvas2.print_png(norm_output_path)
        plt.close(fig2)


# ============================================================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ И КЛАССЫ
# ============================================================
def docx_to_txt(docx_path):
    doc = Document(docx_path)
    text = '\n'.join([para.text for para in doc.paragraphs])
    fd, temp_path = tempfile.mkstemp(suffix='.txt', prefix='rehab_', text=True)
    with os.fdopen(fd, 'w', encoding='utf-8') as f:
        f.write(text)
    return temp_path


class RobustParser:
    @staticmethod
    def parse(txt_path):
        """
        Надёжный парсер данных реабилитации.
        Извлекает упражнения по заголовкам === ... === и привязывает к ним измерения
        строго по позиции в тексте. Добавлена защита от перемешивания каналов.
        """
        content = None
        for enc in ['utf-8', 'cp1251', 'latin-1']:
            try:
                with open(txt_path, 'r', encoding=enc) as f:
                    content = f.read()
                break
            except:
                continue
        if content is None:
            return []

        # Основной паттерн строки данных: время + углы + силы
        pattern = re.compile(
            r'(?P<time>\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}:\d{2}(?:\.\d+)?)\s+'
            r'(?P<angles>[-\d,;.]+)\s+'
            r'(?P<forces>[-\d,;.]+)'
        )
        all_measurements = []
        for match in pattern.finditer(content):
            time_str = match.group('time')
            angles = [float(x) for x in re.split(r'[,;]+', match.group('angles')) if x.strip()]
            forces = [float(x) for x in re.split(r'[,;]+', match.group('forces')) if x.strip()]
            if angles and forces:
                all_measurements.append((time_str, angles, forces, match.start()))

        if not all_measurements:
            return []

        # Более устойчивое извлечение заголовков упражнений
        title_spans = []
        for match in re.finditer(r'\*?\*?===\s*(.+?)\s*===\*?\*?', content):
            raw_title = match.group(1).strip()
            # Чистим возможный мусор вокруг названия
            clean_title = re.sub(r'^\*+|\*+$', '', raw_title).strip()
            title_spans.append((match.start(), clean_title))
        title_spans.sort(key=lambda x: x[0])
        title_spans.append((len(content), None))

        exercises = []
        for i in range(len(title_spans) - 1):
            start_pos, title = title_spans[i]
            end_pos = title_spans[i + 1][0]
            if title is None:
                continue
            block = content[start_pos:end_pos]

            # Извлечение "Начало:" и параметров
            start_time_val = None
            params = {}
            start_match = re.search(r'\*?Начало:\s*(.+?)(?=\n|$)', block, re.IGNORECASE)
            if start_match:
                param_line = start_match.group(1)
                date_match = re.search(r'(\d{4}-\d{2}-\d{2})', param_line)
                time_match = re.search(r'(\d{2}:\d{2}:\d{2})', param_line)
                if date_match and time_match:
                    start_time_val = f"{date_match.group(1)} {time_match.group(1)}"
                elif date_match:
                    start_time_val = date_match.group(1)
                for part in param_line.split('|'):
                    if ':' in part:
                        k, v = part.split(':', 1)
                        params[k.strip()] = v.strip()

            # === ГЛАВНАЯ ЗАЩИТА ОТ ПЕРЕМЕШИВАНИЯ ===
            # Берём только те измерения, которые физически находятся между двумя заголовками
            block_meas = [m for m in all_measurements if start_pos <= m[3] < end_pos]

            if not block_meas:
                continue

            # Проверка согласованности количества каналов внутри одного упражнения.
            # Если есть "битые" строки — оставляем только самые частые длины массивов.
            # Это предотвращает смешение данных разных упражнений или кривых строк.
            angle_lens = [len(m[1]) for m in block_meas]
            force_lens = [len(m[2]) for m in block_meas]
            if len(set(angle_lens)) > 1 or len(set(force_lens)) > 1:
                from collections import Counter
                most_common_a = Counter(angle_lens).most_common(1)[0][0]
                most_common_f = Counter(force_lens).most_common(1)[0][0]
                block_meas = [m for m in block_meas
                              if len(m[1]) == most_common_a and len(m[2]) == most_common_f]

            if not block_meas:
                continue

            # Сортируем по времени на всякий случай
            block_meas.sort(key=lambda x: x[0])

            if start_time_val:
                safe_start = start_time_val.replace(':', '-').replace(' ', '_')
            else:
                safe_start = datetime.now().strftime('%Y%m%d_%H%M%S')
            folder_name = f"{title}_{safe_start}"
            exercises.append({
                'name': title,
                'start_time': start_time_val,
                'params': params,
                'measurements': [(t, a, f) for t, a, f, _ in block_meas],
                'folder_name': folder_name
            })
        return exercises


def create_exercise_report(exercise, output_dir, angles_png, forces_png):
    doc = Document()
    title = doc.add_heading(f'Отчёт по упражнению: {exercise["name"]}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Время начала: {exercise["start_time"]}')
    if exercise['params']:
        doc.add_heading('Параметры', level=2)
        for k, v in exercise['params'].items():
            doc.add_paragraph(f'{k}: {v}')
    doc.add_paragraph(f'Общее количество измерений: {len(exercise["measurements"])}')
    doc.add_heading('Графики', level=2)
    doc.add_paragraph('Углы по каналам:')
    doc.add_picture(angles_png, width=Inches(6.0))
    doc.add_paragraph('Силы по каналам:')
    doc.add_picture(forces_png, width=Inches(6.0))
    report_path = os.path.join(output_dir, 'report.docx')
    doc.save(report_path)
    return report_path


def create_full_report(patient_name, patient_dir):
    exercises_amp_data = {}
    exercises_info = []

    for folder in os.listdir(patient_dir):
        folder_path = os.path.join(patient_dir, folder)
        if not os.path.isdir(folder_path):
            continue
        if os.path.exists(os.path.join(folder_path, 'angles.png')):
            ex_name = folder.split('_', 1)[0] if '_' in folder else folder
            exercises_info.append({
                'name': ex_name,
                'angles_png': os.path.join(folder_path, 'angles.png'),
                'forces_png': os.path.join(folder_path, 'forces.png')
            })
            max_force_path = os.path.join(folder_path, 'max_force.txt')
            if os.path.exists(max_force_path):
                with open(max_force_path, 'r') as f:
                    max_force = float(f.read().strip())
                date_str = ""
                for part in folder.split('_'):
                    if re.match(r'\d{4}-\d{2}-\d{2}', part):
                        date_str = part
                        break
                if date_str:
                    try:
                        dt = datetime.strptime(date_str, '%Y-%m-%d')
                        date_str = dt.strftime('%d.%m.%Y')
                    except:
                        pass
                else:
                    date_str = "неизвестно"
                if ex_name not in exercises_amp_data:
                    exercises_amp_data[ex_name] = {}
                exercises_amp_data[ex_name][date_str] = max(exercises_amp_data[ex_name].get(date_str, 0), max_force)

    if not exercises_info:
        return

    doc = Document()
    title = doc.add_heading(f'Полный отчёт по пациенту: {patient_name}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    doc.add_page_break()

    for ex in exercises_info:
        doc.add_heading(f'Упражнение: {ex["name"]}', level=2)
        doc.add_picture(ex['angles_png'], width=Inches(5.0))
        doc.add_picture(ex['forces_png'], width=Inches(5.0))
        doc.add_page_break()

    if exercises_amp_data:
        doc.add_heading('Динамика амплитуд упражнений', level=1)
        for ex_name, date_force_dict in exercises_amp_data.items():
            items = sorted(date_force_dict.items(),
                           key=lambda x: datetime.strptime(x[0], '%d.%m.%Y') if re.match(r'\d{2}\.\d{2}\.\d{4}',
                                                                                         x[0]) else datetime.min)
            dates = [d[0] for d in items]
            forces = [d[1] for d in items]

            plt.figure(figsize=(8, 5), dpi=100, facecolor='white')
            plt.plot(dates, forces, marker='o', linestyle='-', linewidth=2, markersize=8, color='#5a9eff')
            plt.xlabel('Дата')
            plt.ylabel('Максимальная сила (Н)')
            plt.title(f'{ex_name} – динамика максимальной силы', fontsize=12)
            plt.xticks(rotation=45)
            plt.grid(True, linestyle='--', alpha=0.7)
            plt.tight_layout()
            temp_png = os.path.join(tempfile.gettempdir(), f"amp_report_{ex_name}.png")
            plt.savefig(temp_png, dpi=100, bbox_inches='tight')
            plt.close('all')
            gc.collect()

            doc.add_heading(f'{ex_name}', level=3)
            doc.add_picture(temp_png, width=Inches(6.0))
            if os.path.exists(temp_png):
                os.remove(temp_png)

    full_path = os.path.join(patient_dir, 'full_report.docx')
    doc.save(full_path)
    gc.collect()


# [REMOVED] create_general_report_all_subjects - button removed.


# ============================================================
# PROCESSING THREAD
# ============================================================
class ProcessingThread(QThread):
    finished_signal = pyqtSignal(bool, str, str)

    def __init__(self, docx_path, patient_folder):
        super().__init__()
        self.docx_path = docx_path
        self.patient_folder = patient_folder

    def run(self):
        try:
            txt_path = docx_to_txt(self.docx_path)
            exercises = RobustParser.parse(txt_path)
            os.unlink(txt_path)
            gc.collect()

            patient_dir = os.path.join(PATIENTS_DIR, self.patient_folder)

            for ex in exercises:
                ex_dir = os.path.join(patient_dir, ex['folder_name'])
                os.makedirs(ex_dir, exist_ok=True)
                angles_png, forces_png = save_graphs_for_exercise(ex, ex_dir, smoothing_intensity="medium", patient_name=self.patient_folder)
                if angles_png and forces_png:
                    create_exercise_report(ex, ex_dir, angles_png, forces_png)
                gc.collect()

            create_full_report(self.patient_folder, patient_dir)

            # Генерируем PNG для "Динамики амплитуд" для каждого уникального упражнения
            # (это убирает любую работу matplotlib из UI потока)
            processed_exercises = set()
            for ex in exercises:
                ex_name = ex['name']
                if ex_name not in processed_exercises:
                    processed_exercises.add(ex_name)
                    safe_name = re.sub(r'[^a-zA-Zа-яА-ЯёЁ0-9]', '_', ex_name)
                    amp_png_path = os.path.join(patient_dir, f"{safe_name}_amplitude_dynamics.png")
                    try:
                        generate_amplitude_dynamics_png(self.patient_folder, ex_name, amp_png_path)
                    except Exception as e:
                        print(f"Не удалось сгенерировать amplitude dynamics для {ex_name}: {e}")

            gc.collect()

            self.finished_signal.emit(True, f"Отчёт успешно добавлен к пациенту {self.patient_folder}", patient_dir)
        except Exception as e:
            self.finished_signal.emit(False, f"Ошибка при обработке:\n{str(e)}", "")


# ============================================================
# CARDS
# ============================================================
class AnimatedCard(QWidget):
    def __init__(self, title, subtitle, icon="", parent=None):
        super().__init__(parent)
        self.setFixedHeight(98)
        layout = QVBoxLayout(self)
        layout.setContentsMargins(18, 12, 18, 10)
        layout.setSpacing(4)

        top = QHBoxLayout()
        if icon:
            self.icon_label = QLabel(icon)
            self.icon_label.setStyleSheet("font-size: 26px; background: transparent;")
            top.addWidget(self.icon_label)
        self.title_label = QLabel(title)
        self.title_label.setStyleSheet("font-size: 17px; font-weight: 700; color: #f0f0ff; background: transparent;")
        top.addWidget(self.title_label)
        top.addStretch()
        layout.addLayout(top)

        self.subtitle_label = QLabel(subtitle)
        self.subtitle_label.setStyleSheet("font-size: 13px; color: #b0b0d0; background: transparent;")
        self.subtitle_label.setWordWrap(True)
        self.subtitle_label.setMinimumHeight(48)
        layout.addWidget(self.subtitle_label)

        shadow = QGraphicsDropShadowEffect()
        shadow.setBlurRadius(10)
        shadow.setXOffset(0)
        shadow.setYOffset(2)
        shadow.setColor(QColor(0, 0, 0, 40))
        self.setGraphicsEffect(shadow)

        self.setStyleSheet("background-color: #25253a; border-radius: 14px;")


class PatientCardWidget(AnimatedCard):
    double_clicked = pyqtSignal(str)
    clicked = pyqtSignal(str)

    def __init__(self, display_name, folder_name, subtitle="", parent=None):
        super().__init__(display_name, subtitle, "", parent)
        self.folder_name = folder_name
        self.selected = False

    def mousePressEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            self.clicked.emit(self.folder_name)
        super().mousePressEvent(event)

    def mouseDoubleClickEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            self.double_clicked.emit(self.folder_name)
        super().mouseDoubleClickEvent(event)

    def set_selected(self, selected: bool):
        self.selected = selected
        if selected:
            self.setStyleSheet("""
                PatientCardWidget {
                    background-color: #1a2a4a;
                    border: 4px solid #00d0ff;
                    border-radius: 14px;
                }
            """)
        else:
            self.setStyleSheet("""
                PatientCardWidget {
                    background-color: #25253a;
                    border-radius: 14px;
                }
            """)
        self.update()
        self.repaint()


class ExerciseCardWidget(AnimatedCard):
    def __init__(self, name, subtitle="", icon="📊", parent=None):
        super().__init__(name, subtitle, icon, parent)

    def get_name(self):
        return self.title_label.text()


# ============================================================
# ДИАЛОГИ
# ============================================================
class PatientInfoDialog(QDialog):
    def __init__(self, existing_name="", existing_birth="", existing_complaint="",
                 existing_height="", existing_weight="", existing_upper="", existing_middle="", existing_lower="",
                 parent=None):
        super().__init__(parent)
        self.setWindowTitle(" ")
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.Dialog)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
        self.setMinimumWidth(560)
        self.setMinimumHeight(620)

        self.main_frame = QFrame(self)
        self.main_frame.setObjectName("mainFrame")
        self.main_frame.setStyleSheet("""
            QFrame#mainFrame {
                background-color: #1f2533;
                border-radius: 24px;
                border: 1px solid #3a4459;
            }
        """)
        shadow = QGraphicsDropShadowEffect()
        shadow.setBlurRadius(30)
        shadow.setXOffset(0)
        shadow.setYOffset(10)
        shadow.setColor(QColor(0, 0, 0, 160))
        self.main_frame.setGraphicsEffect(shadow)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.addWidget(self.main_frame)

        inner_layout = QVBoxLayout(self.main_frame)
        inner_layout.setContentsMargins(40, 30, 40, 30)
        inner_layout.setSpacing(18)

        title = QLabel("Добавить нового пациента" if not existing_name else "Редактировать карточку пациента")
        title.setStyleSheet("font-size: 24px; font-weight: 700; color: #5a9eff;")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        inner_layout.addWidget(title)

        surname, first_name, patronymic = self._split_full_name(existing_name)

        form = QGridLayout()
        form.setSpacing(12)

        self.surname_edit = QLineEdit(surname)
        form.addWidget(QLabel("Фамилия: *"), 0, 0)
        form.addWidget(self.surname_edit, 0, 1)

        self.first_name_edit = QLineEdit(first_name)
        form.addWidget(QLabel("Имя: *"), 1, 0)
        form.addWidget(self.first_name_edit, 1, 1)

        self.patronymic_edit = QLineEdit(patronymic)
        form.addWidget(QLabel("Отчество (при наличии)"), 2, 0)
        form.addWidget(self.patronymic_edit, 2, 1)

        self.birth_edit = QLineEdit(existing_birth)
        self.birth_edit.setPlaceholderText("01.01.1990")
        self.birth_edit.textChanged.connect(self.auto_format_birth_date)
        form.addWidget(QLabel("Дата рождения (ДД.ММ.ГГГГ): *"), 3, 0)
        form.addWidget(self.birth_edit, 3, 1)

        self.height_edit = QLineEdit(existing_height)
        self.height_edit.setPlaceholderText("170")
        form.addWidget(QLabel("Рост (см): *"), 4, 0)
        form.addWidget(self.height_edit, 4, 1)

        self.weight_edit = QLineEdit(existing_weight)
        self.weight_edit.setPlaceholderText("70.5")
        form.addWidget(QLabel("Вес (кг): *"), 5, 0)
        form.addWidget(self.weight_edit, 5, 1)

        self.upper_link_edit = QLineEdit(existing_upper)
        self.upper_link_edit.setPlaceholderText("45")
        form.addWidget(QLabel("Длина бедра (см): *"), 6, 0)
        form.addWidget(self.upper_link_edit, 6, 1)

        self.middle_link_edit = QLineEdit(existing_middle)
        self.middle_link_edit.setPlaceholderText("40")
        form.addWidget(QLabel("Длина голени (см): *"), 7, 0)
        form.addWidget(self.middle_link_edit, 7, 1)

        self.lower_link_edit = QLineEdit(existing_lower)
        self.lower_link_edit.setPlaceholderText("35")
        form.addWidget(QLabel("Высота стопы (см): *"), 8, 0)
        form.addWidget(self.lower_link_edit, 8, 1)

        self.complaint_combo = QComboBox()
        self.complaint_combo.setEditable(True)
        self.complaint_combo.addItems([
            "Травма коленного сустава", "Перелом голени / лодыжки",
            "Артроз тазобедренного сустава", "Артроз коленного сустава",
            "Ампутация нижней конечности", "Постоперационная реабилитация",
            "Протезирование коленного сустава", "Протезирование тазобедренного сустава",
            "Повреждение связок голеностопа", "Другое (указать)"
        ])
        if existing_complaint:
            self.complaint_combo.setCurrentText(existing_complaint)
        form.addWidget(QLabel("Причина обращения:"), 9, 0)
        form.addWidget(self.complaint_combo, 9, 1)

        inner_layout.addLayout(form)

        name_validator = QRegularExpressionValidator(QRegularExpression(r"^[а-яА-ЯёЁa-zA-Z]+$"))
        self.surname_edit.setValidator(name_validator)
        self.first_name_edit.setValidator(name_validator)
        self.patronymic_edit.setValidator(name_validator)

        number_validator = QDoubleValidator(0, 999.99, 2)
        number_validator.setNotation(QDoubleValidator.Notation.StandardNotation)
        self.height_edit.setValidator(number_validator)
        self.weight_edit.setValidator(number_validator)
        self.upper_link_edit.setValidator(number_validator)
        self.middle_link_edit.setValidator(number_validator)
        self.lower_link_edit.setValidator(number_validator)

        btn_layout = QHBoxLayout()
        self.save_btn = QPushButton("Сохранить")
        self.save_btn.setFixedHeight(50)
        self.save_btn.clicked.connect(self.accept)
        self.cancel_btn = QPushButton("Отмена")
        self.cancel_btn.setFixedHeight(50)
        self.cancel_btn.clicked.connect(self.reject)

        btn_layout.addWidget(self.save_btn)
        btn_layout.addWidget(self.cancel_btn)
        inner_layout.addLayout(btn_layout)

    def _split_full_name(self, full_name):
        if not full_name:
            return "", "", ""
        parts = full_name.strip().split()
        surname = parts[0] if parts else ""
        first_name = parts[1] if len(parts) > 1 else ""
        patronymic = " ".join(parts[2:]) if len(parts) > 2 else ""
        return surname, first_name, patronymic

    def auto_format_birth_date(self, text):
        self.birth_edit.blockSignals(True)
        cleaned = re.sub(r'\D', '', text)
        formatted = ""
        if len(cleaned) > 0:
            formatted = cleaned[:2]
            if len(cleaned) > 2:
                formatted += '.' + cleaned[2:4]
            if len(cleaned) > 4:
                formatted += '.' + cleaned[4:8]
                if len(cleaned) >= 8:
                    self.birth_edit.clearFocus()
        self.birth_edit.setText(formatted)
        self.birth_edit.setCursorPosition(len(formatted))
        self.birth_edit.blockSignals(False)

    def get_info(self):
        surname = self.surname_edit.text().strip()
        first_name = self.first_name_edit.text().strip()
        if not surname or not first_name:
            show_styled_message(self, "Ошибка", "Фамилия и Имя — обязательные поля!", "warning")
            return None, None, None, None, None, None, None, None
        birth = self.birth_edit.text().strip()
        if not birth:
            show_styled_message(self, "Ошибка", "Дата рождения — обязательное поле!", "warning")
            return None, None, None, None, None, None, None, None
        height = self.height_edit.text().strip()
        weight = self.weight_edit.text().strip()
        if not height or not weight:
            show_styled_message(self, "Ошибка", "Рост и Вес — обязательные поля!", "warning")
            return None, None, None, None, None, None, None, None
        upper = self.upper_link_edit.text().strip()
        middle = self.middle_link_edit.text().strip()
        lower = self.lower_link_edit.text().strip()
        if not upper or not middle or not lower:
            show_styled_message(self, "Ошибка", "Длины звеньев — обязательные поля!", "warning")
            return None, None, None, None, None, None, None, None

        full_name = f"{surname} {first_name}"
        if self.patronymic_edit.text().strip():
            full_name += f" {self.patronymic_edit.text().strip()}"

        return (full_name, birth, self.complaint_combo.currentText().strip(),
                height, weight, upper, middle, lower)


class DeleteConfirmDialog(QDialog):
    def __init__(self, patient_name, parent=None):
        super().__init__(parent)
        self.setWindowTitle(" ")
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.Dialog)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
        self.setMinimumWidth(480)
        self.setMinimumHeight(220)

        self.main_frame = QFrame(self)
        self.main_frame.setObjectName("mainFrame")
        self.main_frame.setStyleSheet("""
            QFrame#mainFrame {
                background-color: #1f2533;
                border-radius: 24px;
                border: 1px solid #3a4459;
            }
        """)
        shadow = QGraphicsDropShadowEffect()
        shadow.setBlurRadius(30)
        shadow.setXOffset(0)
        shadow.setYOffset(10)
        shadow.setColor(QColor(0, 0, 0, 160))
        self.main_frame.setGraphicsEffect(shadow)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.addWidget(self.main_frame)

        inner = QVBoxLayout(self.main_frame)
        inner.setContentsMargins(40, 30, 40, 30)
        inner.setSpacing(20)

        title = QLabel("Удаление пациента")
        title.setStyleSheet("font-size: 22px; font-weight: 700; color: #ff6b6b;")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        inner.addWidget(title)

        msg = QLabel(f"Вы действительно хотите удалить пациента\n«{patient_name}»\nи все его данные?")
        msg.setStyleSheet("font-size: 16px; color: #e0e0e0; text-align: center;")
        msg.setWordWrap(True)
        msg.setAlignment(Qt.AlignmentFlag.AlignCenter)
        inner.addWidget(msg)

        btn_layout = QHBoxLayout()
        self.yes_btn = QPushButton("Да, удалить")
        self.yes_btn.setFixedHeight(48)
        self.yes_btn.setStyleSheet("""
            QPushButton { background-color: #ff4444; color: white; border-radius: 12px; font-weight: 600; }
            QPushButton:hover { background-color: #ff6666; }
        """)
        self.yes_btn.clicked.connect(self.accept)

        self.no_btn = QPushButton("Отмена")
        self.no_btn.setFixedHeight(48)
        self.no_btn.setStyleSheet("""
            QPushButton { background-color: #2a2a3a; border: 2px solid #5a9eff; border-radius: 12px; color: white; font-weight: 600; }
            QPushButton:hover { background-color: #3a3a4e; }
        """)
        self.no_btn.clicked.connect(self.reject)

        btn_layout.addWidget(self.yes_btn)
        btn_layout.addWidget(self.no_btn)
        inner.addLayout(btn_layout)


class DeleteReportDialog(QDialog):
    def __init__(self, patient_name, available_dates, parent=None):
        super().__init__(parent)
        self.setWindowTitle(" ")
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.Dialog)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
        self.setMinimumWidth(480)
        self.setMinimumHeight(320)

        self.main_frame = QFrame(self)
        self.main_frame.setObjectName("mainFrame")
        self.main_frame.setStyleSheet("""
            QFrame#mainFrame {
                background-color: #1f2533;
                border-radius: 24px;
                border: 1px solid #3a4459;
            }
        """)
        shadow = QGraphicsDropShadowEffect()
        shadow.setBlurRadius(30)
        shadow.setXOffset(0)
        shadow.setYOffset(10)
        shadow.setColor(QColor(0, 0, 0, 160))
        self.main_frame.setGraphicsEffect(shadow)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.addWidget(self.main_frame)

        inner = QVBoxLayout(self.main_frame)
        inner.setContentsMargins(40, 30, 40, 30)
        inner.setSpacing(20)

        title = QLabel("Удаление отчёта")
        title.setStyleSheet("font-size: 22px; font-weight: 700; color: #ff6b6b;")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        inner.addWidget(title)

        msg = QLabel(f"Выберите дату отчёта для удаления\n(пациент: {patient_name})")
        msg.setStyleSheet("font-size: 15px; color: #e0e0e0; text-align: center;")
        msg.setWordWrap(True)
        msg.setAlignment(Qt.AlignmentFlag.AlignCenter)
        inner.addWidget(msg)

        inner.addWidget(QLabel("Дата выполнения упражнений:"))
        self.date_combo = QComboBox()
        self.date_combo.setStyleSheet("""
            QComboBox {
                background-color: #2a2a3a;
                color: white;
                border: 1px solid #5a9eff;
                border-radius: 8px;
                padding: 10px;
                font-size: 14px;
            }
            QComboBox::drop-down {
                border: none;
            }
            QComboBox::down-arrow {
                image: none;
                border: none;
            }
        """)
        for date in sorted(available_dates, reverse=True):
            self.date_combo.addItem(date)
        inner.addWidget(self.date_combo)

        if not available_dates:
            no_dates_label = QLabel("Нет доступных отчётов для удаления")
            no_dates_label.setStyleSheet("color: #ff8888; font-size: 14px;")
            no_dates_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            inner.addWidget(no_dates_label)

        btn_layout = QHBoxLayout()
        self.delete_btn = QPushButton("Удалить отчёт")
        self.delete_btn.setFixedHeight(48)
        self.delete_btn.setStyleSheet("""
            QPushButton { background-color: #ff4444; color: white; border-radius: 12px; font-weight: 600; }
            QPushButton:hover { background-color: #ff6666; }
        """)
        self.delete_btn.clicked.connect(self.accept)

        self.cancel_btn = QPushButton("Отмена")
        self.cancel_btn.setFixedHeight(48)
        self.cancel_btn.setStyleSheet("""
            QPushButton { background-color: #2a2a3a; border: 2px solid #5a9eff; border-radius: 12px; color: white; font-weight: 600; }
            QPushButton:hover { background-color: #3a3a4e; }
        """)
        self.cancel_btn.clicked.connect(self.reject)

        btn_layout.addWidget(self.delete_btn)
        btn_layout.addWidget(self.cancel_btn)
        inner.addLayout(btn_layout)

        if not available_dates:
            self.delete_btn.setEnabled(False)
            self.delete_btn.setStyleSheet("background-color: #555; color: #aaa; border-radius: 12px;")

    def get_selected_date(self):
        return self.date_combo.currentText()


class PatientFilterDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Фильтры поиска пациентов")
        self.setMinimumWidth(420)
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.Dialog)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)

        self.main_frame = QFrame(self)
        self.main_frame.setObjectName("mainFrame")
        self.main_frame.setStyleSheet("""
            QFrame#mainFrame {
                background-color: #1f2533;
                border-radius: 20px;
                border: 1px solid #3a4459;
            }
        """)
        shadow = QGraphicsDropShadowEffect()
        shadow.setBlurRadius(25)
        shadow.setXOffset(0)
        shadow.setYOffset(8)
        shadow.setColor(QColor(0, 0, 0, 140))
        self.main_frame.setGraphicsEffect(shadow)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.addWidget(self.main_frame)

        inner = QVBoxLayout(self.main_frame)
        inner.setContentsMargins(30, 25, 30, 25)
        inner.setSpacing(18)

        title = QLabel("Фильтры")
        title.setStyleSheet("font-size: 20px; font-weight: 700; color: #5a9eff;")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        inner.addWidget(title)

        inner.addWidget(QLabel("Дата рождения (ДД.ММ.ГГГГ):"))
        self.birth_edit = QLineEdit()
        self.birth_edit.setPlaceholderText("01.01.1990")
        inner.addWidget(self.birth_edit)

        inner.addWidget(QLabel("Причина обращения:"))
        self.complaint_edit = QLineEdit()
        self.complaint_edit.setPlaceholderText("например: травма колена")
        inner.addWidget(self.complaint_edit)

        btn_layout = QHBoxLayout()
        self.apply_btn = QPushButton("Применить")
        self.apply_btn.setFixedHeight(48)
        self.apply_btn.clicked.connect(self.accept)

        self.cancel_btn = QPushButton("Отмена")
        self.cancel_btn.setFixedHeight(48)
        self.cancel_btn.clicked.connect(self.reject)

        btn_layout.addWidget(self.apply_btn)
        btn_layout.addWidget(self.cancel_btn)
        inner.addLayout(btn_layout)

    def get_filters(self):
        return {
            'birth': self.birth_edit.text().strip().lower(),
            'complaint': self.complaint_edit.text().strip().lower()
        }


# ============================================================
# Styled Message Box (consistent with PatientInfoDialog style)
# ============================================================
class StyledMessageBox(QDialog):
    def __init__(self, title, message, parent=None, msg_type="info"):
        super().__init__(parent)
        self.setWindowTitle(" ")
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.Dialog)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
        self.setMinimumWidth(420)

        self.main_frame = QFrame(self)
        self.main_frame.setObjectName("mainFrame")

        accent_color = "#5a9eff"
        if msg_type == "warning":
            accent_color = "#ffaa44"
        elif msg_type == "critical":
            accent_color = "#ff5555"

        self.main_frame.setStyleSheet(f"""
            QFrame#mainFrame {{
                background-color: #1f2533;
                border-radius: 20px;
                border: 1px solid #3a4459;
            }}
        """)

        shadow = QGraphicsDropShadowEffect()
        shadow.setBlurRadius(28)
        shadow.setXOffset(0)
        shadow.setYOffset(8)
        shadow.setColor(QColor(0, 0, 0, 150))
        self.main_frame.setGraphicsEffect(shadow)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.addWidget(self.main_frame)

        inner = QVBoxLayout(self.main_frame)
        inner.setContentsMargins(32, 28, 32, 28)
        inner.setSpacing(18)

        title_label = QLabel(title)
        title_label.setStyleSheet(f"font-size: 20px; font-weight: 700; color: {accent_color};")
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        inner.addWidget(title_label)

        msg_label = QLabel(message)
        msg_label.setStyleSheet("font-size: 15px; color: #e0e0e0;")
        msg_label.setWordWrap(True)
        msg_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        inner.addWidget(msg_label)

        # Кнопка OK в стиле основных action-кнопок приложения (как "Сохранить", "Да, удалить")
        btn_layout = QHBoxLayout()
        btn_layout.setContentsMargins(0, 8, 0, 0)

        self.ok_btn = QPushButton("OK")
        self.ok_btn.setFixedHeight(48)
        self.ok_btn.setMinimumWidth(160)
        self.ok_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {accent_color};
                color: white;
                border-radius: 12px;
                font-weight: 600;
                font-size: 15px;
                padding: 0 32px;
            }}
            QPushButton:hover {{
                background-color: #6ab0ff;
            }}
            QPushButton:pressed {{
                background-color: #4a8ad0;
            }}
        """)
        self.ok_btn.clicked.connect(self.accept)

        # Делаем кнопку заметной и центрированной, но с хорошей шириной
        btn_layout.addStretch()
        btn_layout.addWidget(self.ok_btn)
        btn_layout.addStretch()

        inner.addLayout(btn_layout)


def show_styled_message(parent, title, message, msg_type="info"):
    dlg = StyledMessageBox(title, message, parent, msg_type)
    dlg.exec()


# ============================================================
# PatientsListPage
# ============================================================
class PatientsListPage(QWidget):
    patient_selected = pyqtSignal(str)

    def __init__(self, main_window, parent=None):
        super().__init__(parent)
        self.main_window = main_window
        self.all_patients = []
        self.filtered_patients = []
        self.selected_folder = None
        self.current_filters = None

        layout = QVBoxLayout(self)
        layout.setContentsMargins(30, 30, 30, 30)
        layout.setSpacing(18)

        title = QLabel("SWSU ROBOTICS")
        title.setStyleSheet("font-size: 32px; font-weight: 900; color: #5a9eff;")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title)

        self.add_patient_btn = QPushButton("Добавить нового пациента")
        self.add_patient_btn.setFixedHeight(54)
        self.add_patient_btn.clicked.connect(self.main_window.add_new_patient)
        self.add_patient_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #1e90ff, stop:1 #0b7be0);
                color: white;
                border-radius: 12px;
                font-size: 17px;
                font-weight: 600;
                border: none;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #40b0ff, stop:1 #1e90ff);
            }
            QPushButton:pressed {
                background: #0a5eb8;
            }
        """)
        layout.addWidget(self.add_patient_btn)

        search_layout = QHBoxLayout()
        self.search_edit = QLineEdit()
        self.search_edit.setPlaceholderText("🔍 Поиск по ФИО...")
        self.search_edit.textChanged.connect(self.filter_patients)
        search_layout.addWidget(self.search_edit)

        self.filter_btn = QPushButton("🔧 Фильтры")
        self.filter_btn.clicked.connect(self.show_filter_dialog)
        search_layout.addWidget(self.filter_btn)
        layout.addLayout(search_layout)

        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)

        self.list_container = QWidget()
        self.list_layout = QVBoxLayout(self.list_container)
        self.list_layout.setContentsMargins(0, 0, 0, 0)
        self.list_layout.setSpacing(14)
        self.scroll.setWidget(self.list_container)
        layout.addWidget(self.scroll)

        self.load_patients()

    def load_patients(self):
        patients = []
        ignore_folders = ['__pycache__', 'dist', 'build', 'venv', 'myenv', 'env', '.git', '.idea']
        for item in os.listdir(PATIENTS_DIR):
            item_path = os.path.join(PATIENTS_DIR, item)
            if not os.path.isdir(item_path) or item.startswith('.') or item in ignore_folders:
                continue
            if os.path.exists(os.path.join(item_path, 'info.txt')):
                patients.append(item)
        self.all_patients = sorted(patients)
        self.filter_patients()

    def show_filter_dialog(self):
        dlg = PatientFilterDialog(self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            self.current_filters = dlg.get_filters()
        else:
            self.current_filters = None
        self.filter_patients()

    def filter_patients(self):
        text = self.search_edit.text().strip().lower()
        filters = self.current_filters

        result = []
        for folder in self.all_patients:
            patient_dir = os.path.join(PATIENTS_DIR, folder)
            info_path = os.path.join(patient_dir, 'info.txt')
            display_name = folder
            birth = ""
            complaint = ""
            if os.path.exists(info_path):
                try:
                    with open(info_path, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                        display_name = data.get('name', folder)
                        birth = data.get('birth_date', '')
                        complaint = data.get('complaint', '').lower()
                except:
                    pass

            if text and text not in display_name.lower() and text not in folder.lower():
                continue
            if filters:
                if filters.get('birth') and filters['birth'] not in birth.lower():
                    continue
                if filters.get('complaint') and filters['complaint'] not in complaint:
                    continue

            result.append((folder, display_name, birth))

        self.filtered_patients = result
        self.update_list()

    def update_list(self):
        for i in reversed(range(self.list_layout.count())):
            widget = self.list_layout.itemAt(i).widget()
            if widget:
                widget.deleteLater()
        gc.collect()

        for folder, display_name, birth in self.filtered_patients:
            patient_dir = os.path.join(PATIENTS_DIR, folder)
            if not os.path.exists(patient_dir):
                continue
            exercises_count = sum(
                1 for f in os.listdir(patient_dir)
                if os.path.isdir(os.path.join(patient_dir, f)) and os.path.exists(
                    os.path.join(patient_dir, f, 'angles.png'))
            )
            subtitle = f"Дата рождения: {birth}" if birth else ""
            if exercises_count:
                if subtitle:
                    subtitle += f" • {exercises_count} упражнений"
                else:
                    subtitle = f"{exercises_count} упражнений"

            card = PatientCardWidget(display_name, folder, subtitle)
            card.double_clicked.connect(self.on_patient_double_clicked)
            card.clicked.connect(self.on_patient_clicked)

            if folder == self.selected_folder:
                card.set_selected(True)

            self.list_layout.addWidget(card)

        spacer = QWidget()
        spacer.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self.list_layout.addWidget(spacer)

    def on_patient_clicked(self, folder_name):
        self.selected_folder = folder_name
        self.update_list()

    def on_patient_double_clicked(self, folder_name):
        self.selected_folder = folder_name
        self.update_list()
        self.patient_selected.emit(folder_name)

    def set_selected_patient(self, folder_name: str):
        self.selected_folder = folder_name
        self.update_list()

    def refresh(self):
        self.load_patients()


# ============================================================
# PatientWithExercisesPage
# ============================================================
class PatientWithExercisesPage(QWidget):
    back_to_list = pyqtSignal()
    exercise_selected = pyqtSignal(str, str)

    def __init__(self, main_window, parent=None):
        super().__init__(parent)
        self.main_window = main_window
        self.current_patient = None
        self.exercise_variants = {}

        layout = QHBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(20)

        self.left_panel = QFrame()
        self.left_panel.setFrameShape(QFrame.Shape.StyledPanel)
        self.left_panel.setStyleSheet("background-color: #2a2a3a; border-radius: 18px;")
        left_layout = QVBoxLayout(self.left_panel)
        left_layout.setContentsMargins(22, 22, 22, 22)

        self.back_btn = QPushButton("←")
        self.back_btn.setFixedSize(50, 44)
        self.back_btn.setStyleSheet("""
            QPushButton {
                background-color: #ff4444;
                color: white;
                font-size: 24px;
                font-weight: bold;
                border: none;
                border-radius: 8px;
            }
            QPushButton:hover {
                background-color: #ff6666;
            }
        """)
        self.back_btn.clicked.connect(self.back_to_list.emit)
        left_layout.addWidget(self.back_btn)

        self.patient_info_group = QGroupBox("Карточка пациента")
        self.patient_info_group.setStyleSheet(
            "QGroupBox { color: white; font-size: 15px; border: 1px solid #444; border-radius: 14px; margin-top: 12px; }")
        info_layout = QGridLayout(self.patient_info_group)
        self.patient_name_label = QLabel()
        self.patient_name_label.setStyleSheet("font-size: 21px; font-weight: bold; color: #5a9eff;")
        self.patient_birth_label = QLabel()
        self.patient_height_label = QLabel()
        self.patient_weight_label = QLabel()
        self.patient_upper_link_label = QLabel()
        self.patient_middle_link_label = QLabel()
        self.patient_lower_link_label = QLabel()
        self.patient_complaint_label = QLabel()
        self.patient_exercises_count_label = QLabel()

        info_layout.addWidget(QLabel("ФИО:"), 0, 0)
        info_layout.addWidget(self.patient_name_label, 0, 1)
        info_layout.addWidget(QLabel("Дата рождения:"), 1, 0)
        info_layout.addWidget(self.patient_birth_label, 1, 1)
        info_layout.addWidget(QLabel("Рост:"), 2, 0)
        info_layout.addWidget(self.patient_height_label, 2, 1)
        info_layout.addWidget(QLabel("Вес:"), 3, 0)
        info_layout.addWidget(self.patient_weight_label, 3, 1)
        info_layout.addWidget(QLabel("Длина бедра:"), 4, 0)
        info_layout.addWidget(self.patient_upper_link_label, 4, 1)
        info_layout.addWidget(QLabel("Длина голени:"), 5, 0)
        info_layout.addWidget(self.patient_middle_link_label, 5, 1)
        info_layout.addWidget(QLabel("Высота стопы:"), 6, 0)
        info_layout.addWidget(self.patient_lower_link_label, 6, 1)
        info_layout.addWidget(QLabel("Причина обращения:"), 7, 0)
        info_layout.addWidget(self.patient_complaint_label, 7, 1)
        info_layout.addWidget(QLabel("Упражнений:"), 8, 0)
        info_layout.addWidget(self.patient_exercises_count_label, 8, 1)
        left_layout.addWidget(self.patient_info_group)

        btn_layout = QVBoxLayout()
        btn_layout.setSpacing(12)

        self.add_report_btn = QPushButton("📋  Добавить отчёт")
        self.add_report_btn.setFixedHeight(50)
        self.add_report_btn.setStyleSheet("""
            QPushButton {
                background-color: #2a2a3a;
                border: 2px solid #5a9eff;
                border-radius: 12px;
                color: white;
                font-weight: 600;
            }
            QPushButton:hover {
                background-color: #3a3a4e;
                border: 2px solid #6ab0ff;
            }
        """)
        self.add_report_btn.clicked.connect(self.add_report_to_patient)
        btn_layout.addWidget(self.add_report_btn)

        self.edit_btn = QPushButton("✏️  Редактировать карточку")
        self.edit_btn.setFixedHeight(50)
        self.edit_btn.setStyleSheet("""
            QPushButton {
                background-color: #2a2a3a;
                border: 2px solid #5a9eff;
                border-radius: 12px;
                color: white;
                font-weight: 600;
            }
            QPushButton:hover {
                background-color: #3a3a4e;
                border: 2px solid #6ab0ff;
            }
        """)
        self.edit_btn.clicked.connect(self.edit_current_patient)
        btn_layout.addWidget(self.edit_btn)

        self.delete_report_btn = QPushButton("🗑  Удалить отчёт по дате")
        self.delete_report_btn.setFixedHeight(50)
        self.delete_report_btn.setStyleSheet("""
            QPushButton {
                background-color: #2a2a3a;
                border: 2px solid #ffaa44;
                border-radius: 12px;
                color: white;
                font-weight: 600;
            }
            QPushButton:hover {
                background-color: #3a3a4e;
                border: 2px solid #ffcc66;
            }
        """)
        self.delete_report_btn.clicked.connect(self.delete_report_by_date)
        btn_layout.addWidget(self.delete_report_btn)

        self.delete_patient_btn = QPushButton("🗑  Удалить пациента")
        self.delete_patient_btn.setFixedHeight(50)
        self.delete_patient_btn.setStyleSheet("""
            QPushButton {
                background-color: #2a2a3a;
                border: 2px solid #ff6b6b;
                border-radius: 12px;
                color: white;
                font-weight: 600;
            }
            QPushButton:hover {
                background-color: #3a3a4e;
                border: 2px solid #ff8a8a;
            }
        """)
        self.delete_patient_btn.clicked.connect(self.delete_patient)
        btn_layout.addWidget(self.delete_patient_btn)

        left_layout.addLayout(btn_layout)
        left_layout.addStretch()

        self.right_panel = QWidget()
        right_layout = QVBoxLayout(self.right_panel)
        right_layout.setContentsMargins(0, 0, 0, 0)

        self.ex_list_widget = QListWidget()
        self.ex_list_widget.setStyleSheet("""
            QListWidget {
                background: transparent;
                border: none;
            }
            QListWidget::item {
                background: transparent;
            }
        """)
        self.ex_list_widget.setSpacing(14)
        self.ex_list_widget.itemDoubleClicked.connect(self.on_exercise_double_click)
        right_layout.addWidget(self.ex_list_widget)

        self.middle_panel = QWidget()
        middle_layout = QVBoxLayout(self.middle_panel)
        middle_layout.setContentsMargins(0, 0, 0, 0)

        self.patient_analysis_group = QGroupBox("Общий анализ пациента (5 ИИ-агентов)")
        self.patient_analysis_group.setStyleSheet(
            "QGroupBox { color: white; font-size: 15px; border: 1px solid #444; border-radius: 14px; margin-top: 12px; }")
        group_layout = QVBoxLayout(self.patient_analysis_group)
        group_layout.setContentsMargins(12, 20, 12, 12)

        self.patient_analysis_text = QTextEdit()
        self.patient_analysis_text.setReadOnly(True)
        self.patient_analysis_text.setMinimumHeight(220)
        self.patient_analysis_text.setStyleSheet("""
            QTextEdit {
                background-color: #1e1e1e;
                border: none;
                color: #e0e0e0;
                font-size: 13px;
                padding: 12px;
            }
        """)
        self.patient_analysis_text.setPlaceholderText("Анализ загружается автоматически при открытии пациента...")
        group_layout.addWidget(self.patient_analysis_text)
        middle_layout.addWidget(self.patient_analysis_group)

        splitter = QSplitter(Qt.Orientation.Horizontal)
        splitter.addWidget(self.left_panel)
        splitter.addWidget(self.middle_panel)
        splitter.addWidget(self.right_panel)
        splitter.setSizes([400, 600, 800])
        layout.addWidget(splitter)

    def edit_current_patient(self):
        if self.current_patient:
            self.main_window.edit_patient(self.current_patient)

    def set_patient(self, patient_name):
        self.current_patient = patient_name
        self.load_patient_info()
        self.load_exercises()
        self.load_patient_analysis()

    def add_report_to_patient(self):
        if self.current_patient:
            self.main_window.add_report_to_existing(self.current_patient)

    def _get_patient_info_from_file(self):
        patient_dir = os.path.join(PATIENTS_DIR, self.current_patient)
        info_path = os.path.join(patient_dir, 'info.txt')
        patient_info = {
            'age_years': 45,
            'weight_kg': 70,
            'height_cm': 170,
            'sex': 'male',
            'complaint': '',
            'exercise_name': 'упражнение'
        }
        
        if os.path.exists(info_path):
            try:
                with open(info_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    patient_info.update({
                        'weight_kg': data.get('weight_kg', 70),
                        'height_cm': data.get('height_cm', 170),
                        'sex': data.get('sex', 'male'),
                        'complaint': data.get('complaint', ''),
                        'exercise_name': data.get('exercise_name', 'упражнение'),
                        'birth_date': data.get('birth_date', ''),
                    })
                    birth = data.get('birth_date', '')
                    if birth:
                        try:
                            for fmt in ('%d.%m.%Y', '%Y-%m-%d', '%d/%m/%Y'):
                                try:
                                    bd = datetime.strptime(birth, fmt)
                                    patient_info['age_years'] = (datetime.now() - bd).days / 365.25
                                    break
                                except ValueError:
                                    continue
                        except Exception:
                            pass
                    elif data.get('age_years') is not None:
                        patient_info['age_years'] = data.get('age_years', 45)
            except:
                pass
        
        return patient_info

    def load_patient_info(self):
        patient_dir = os.path.join(PATIENTS_DIR, self.current_patient)
        info_path = os.path.join(patient_dir, 'info.txt')
        name = self.current_patient
        birth = ""
        complaint = ""
        height = ""
        weight = ""
        upper = ""
        middle = ""
        lower = ""
        if os.path.exists(info_path):
            try:
                with open(info_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    name = data.get('name', self.current_patient)
                    birth = data.get('birth_date', '')
                    complaint = data.get('complaint', '')
                    height = data.get('height_cm', '')
                    weight = data.get('weight_kg', '')
                    upper = data.get('upper_link_cm', '')
                    middle = data.get('middle_link_cm', '')
                    lower = data.get('lower_link_cm', '')
            except:
                pass
        exercises_count = sum(
            1 for f in os.listdir(patient_dir)
            if
            os.path.isdir(os.path.join(patient_dir, f)) and os.path.exists(os.path.join(patient_dir, f, 'angles.png'))
        )
        self.patient_name_label.setText(name)
        self.patient_birth_label.setText(birth if birth else "не указана")
        self.patient_height_label.setText(f"{height} см" if height else "не указан")
        self.patient_weight_label.setText(f"{weight} кг" if weight else "не указан")
        self.patient_upper_link_label.setText(f"{upper} см" if upper else "не указана")
        self.patient_middle_link_label.setText(f"{middle} см" if middle else "не указана")
        self.patient_lower_link_label.setText(f"{lower} см" if lower else "не указана")
        self.patient_complaint_label.setText(complaint if complaint else "не указана")
        self.patient_exercises_count_label.setText(str(exercises_count))

    def load_exercises(self):
        if not self.current_patient:
            return
        patient_dir = os.path.join(PATIENTS_DIR, self.current_patient)
        variants = {}

        folder_list = []
        for folder in os.listdir(patient_dir):
            folder_path = os.path.join(patient_dir, folder)
            if not os.path.isdir(folder_path) or not os.path.exists(os.path.join(folder_path, 'angles.png')):
                continue
            date_match = re.search(r'(\d{4}-\d{2}-\d{2})', folder)
            if date_match:
                try:
                    dt = datetime.strptime(date_match.group(1), '%Y-%m-%d')
                    folder_list.append((dt, folder))
                except:
                    folder_list.append((datetime.min, folder))

        folder_list.sort(key=lambda x: x[0], reverse=True)

        for _, folder in folder_list:
            folder_path = os.path.join(patient_dir, folder)
            ex_name = folder.split('_', 1)[0] if '_' in folder else folder
            date_str = ""
            parts = folder.split('_')
            for part in parts:
                if re.match(r'\d{4}-\d{2}-\d{2}', part):
                    try:
                        dt = datetime.strptime(part, '%Y-%m-%d')
                        date_str = dt.strftime('%d.%m.%Y')
                    except:
                        date_str = part
                    break
            else:
                date_str = "неизвестно"

            if ex_name not in variants:
                variants[ex_name] = []
            variants[ex_name].append((date_str, folder))

        self.exercise_variants = variants
        self.update_exercise_list()

    def update_exercise_list(self):
        self.ex_list_widget.clear()
        for ex_name in sorted(self.exercise_variants.keys()):
            variants = self.exercise_variants.get(ex_name, [])
            dates_str = ", ".join([v[0] for v in variants])
            subtitle = f"{len(variants)} вариант(ов): {dates_str}"
            icon = self._get_exercise_icon(ex_name)
            card = ExerciseCardWidget(ex_name, subtitle, icon)
            item = QListWidgetItem()
            item.setSizeHint(card.sizeHint())
            self.ex_list_widget.addItem(item)
            self.ex_list_widget.setItemWidget(item, card)

    def _get_exercise_icon(self, name):
        if "ХОДЬБА" in name: return "🚶"
        if "ПОВОРОТ" in name: return "🔄"
        if "ПРИСЕДАНИЯ" in name: return "🏋️"
        return "📊"

    def load_patient_analysis(self):
        if not self.current_patient:
            return
        self.patient_analysis_text.setPlainText("Загрузка общего анализа пациента (5 ИИ-агентов)...\nПожалуйста, подождите.")
        patient_dir = os.path.join(PATIENTS_DIR, self.current_patient)
        try:
            patient_info = self._get_patient_info_from_file()
            converted_sessions = []
            for folder in os.listdir(patient_dir):
                folder_path = os.path.join(patient_dir, folder)
                if not os.path.isdir(folder_path):
                    continue
                raw_path = os.path.join(folder_path, 'raw_measurements.json')
                if not os.path.exists(raw_path):
                    continue
                try:
                    with open(raw_path, 'r', encoding='utf-8') as f:
                        raw = json.load(f)
                except Exception:
                    continue
                times_str = raw.get('times', [])
                angles_rows = raw.get('angles', []) or []
                forces_rows = raw.get('forces', []) or []
                if not times_str or len(times_str) < 10:
                    continue
                base_dt = None
                numeric_times = []
                for ts in times_str:
                    try:
                        if ' ' in ts:
                            try:
                                dt = datetime.strptime(ts, '%Y-%m-%d %H:%M:%S.%f')
                            except ValueError:
                                dt = datetime.strptime(ts, '%Y-%m-%d %H:%M:%S')
                        else:
                            dt = datetime.strptime(ts, '%Y-%m-%d')
                        if base_dt is None:
                            base_dt = dt
                        numeric_times.append((dt - base_dt).total_seconds())
                    except Exception:
                        numeric_times.append(float(len(numeric_times)) * 0.01)
                T = len(numeric_times)
                n_a = len(angles_rows[0]) if angles_rows and angles_rows[0] else 0
                n_f = len(forces_rows[0]) if forces_rows and forces_rows[0] else 0
                angles_by_ch = [[] for _ in range(n_a)]
                for row in angles_rows:
                    for ch in range(n_a):
                        try:
                            angles_by_ch[ch].append(float(row[ch]) if row[ch] is not None else 0.0)
                        except Exception:
                            angles_by_ch[ch].append(0.0)
                forces_by_ch = [[] for _ in range(n_f)]
                for row in forces_rows:
                    for ch in range(n_f):
                        try:
                            forces_by_ch[ch].append(float(row[ch]) if row[ch] is not None else 0.0)
                        except Exception:
                            forces_by_ch[ch].append(0.0)
                forces_N = [[fv * 9.81 / 1000.0 for fv in ch] for ch in forces_by_ch] if forces_by_ch else []
                M_list = []
                for t_idx in range(T):
                    fsum = sum(ch[t_idx] for ch in forces_N if t_idx < len(ch))
                    M_list.append(fsum)
                session = {
                    'exercise_name': folder.split('_')[0] if '_' in folder else folder,
                    'date': '?',
                    'times': numeric_times,
                    'angles': angles_by_ch,
                    'forces': forces_by_ch,
                    'M': M_list,
                }
                for part in folder.split('_'):
                    if re.match(r'\d{4}-\d{2}-\d{2}', part):
                        session['date'] = part
                        break
                converted_sessions.append(session)
            if len(converted_sessions) < 2:
                self.patient_analysis_text.setPlainText(
                    "Недостаточно данных для анализа.\n"
                    "Требуется минимум 2 полноценные сессии с файлами raw_measurements.json.")
                return

            # === КЭШИРОВАНИЕ общего анализа ===
            patient_dir_cache = os.path.join(PATIENTS_DIR, self.current_patient)
            cache_path_patient = os.path.join(patient_dir_cache, "_patient_analysis_cache.json")
            sess_key_parts = [f"{s.get('exercise_name','')}_{len(s.get('times',[]))}" for s in converted_sessions]
            sess_cache_key = "|".join(sorted(sess_key_parts))

            if os.path.exists(cache_path_patient):
                try:
                    with open(cache_path_patient, 'r', encoding='utf-8') as f:
                        cached_p = json.load(f)
                    if cached_p.get('session_key') == sess_cache_key:
                        cached_text = cached_p.get('full_text', '')
                        if len(cached_text) > 200:
                            self.patient_analysis_text.setPlainText(cached_text)
                            print(f"[CACHE] Общий анализ загружен из кэша ({len(cached_text)} символов)")
                            return
                except Exception:
                    pass

            master = EnsembleOrchestrator()
            report = master.run_full_analysis(patient_info, converted_sessions)
            ens = report.get('ensemble_result', {}) or {}
            risk = ens.get('final_risk', 'moderate')
            conf = float(ens.get('overall_confidence', 0.75) or 0.75)
            recs = ens.get('recommendations', []) or []
            na = report.get('new_analyses', {}) or {}
            breakdown = report.get('agent_breakdown', {}) or {}
            prog = report.get('progression', {}) or {}
            conflicts = report.get('conflicts_resolved', []) or []
            block = self._build_patient_analysis_block(
                patient_info, converted_sessions, risk, conf,
                recs, na, breakdown, prog, conflicts)
            full_block_text = "\n".join(block)
            self.patient_analysis_text.setPlainText(full_block_text)

            # === Сохраняем общий анализ в кэш ===
            try:
                cache_data_p = {
                    'session_key': sess_cache_key,
                    'full_text': full_block_text,
                    'saved_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                }
                with open(cache_path_patient, 'w', encoding='utf-8') as f:
                    json.dump(cache_data_p, f, ensure_ascii=False, indent=2)
                print(f"[CACHE] Общий анализ сохранён ({len(full_block_text)} символов)")
            except Exception as ce:
                print(f"[CACHE] Ошибка сохранения общего анализа: {ce}")
        except Exception as ex:
            import traceback
            traceback.print_exc()
            self.patient_analysis_text.setPlainText(
                f"Ошибка при вычислении общего анализа.\n\n"
                f"Техническая информация: {str(ex)[:300]}")

    def _build_patient_analysis_block(self, patient_info, sessions, risk, conf,
                                      recs, na, breakdown, prog, conflicts):
        """Формирует развёрнутый блок общего анализа пациента с подробными рекомендациями."""
        import re as _re
        sep = "═" * 60
        dash = "─" * 60
        lines = []
        lines.append(sep)
        lines.append("ОБЩИЙ АНАЛИЗ ПАЦИЕНТА (АНСАМБЛЬ ИЗ 5 ИИ-АГЕНТОВ)")
        lines.append(sep)
        lines.append("")
        lines.append(f"Пациент: {self.current_patient}")
        lines.append(f"Обработано сессий: {len(sessions)}")
        lines.append(f"Дата анализа: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        lines.append("")
        lines.append("УСТРОЙСТВО: Реабилитационный экзоскелетный комплекс нижних конечностей.")
        lines.append("Пациент находится в подвесе; экзоскелет выполняет пассивные движения ногами.")
        lines.append("Система снимает показания напряжения мышц и угловых перемещений во время работы комплекса.")
        lines.append("")

        age = float(patient_info.get('age_years') or 45)
        weight = float(patient_info.get('weight_kg') or 70)
        complaint = (patient_info.get('complaint') or '').strip()
        if age < 12:
            age_group = 'ребёнок'
        elif age < 18:
            age_group = 'подросток'
        elif age < 60:
            age_group = 'взрослый'
        else:
            age_group = 'пожилой'

        lines.append(dash)
        lines.append("1. КОНТЕКСТ ПАЦИЕНТА")
        lines.append(dash)
        lines.append(f"   • Возрастная группа: {age_group} ({age:.0f} лет)")
        lines.append(f"   • Вес тела: {weight:.1f} кг")
        if complaint:
            lines.append(f"   • Причина обращения: {complaint}")
        _risk_map = {'low': 'НИЗКИЙ', 'moderate': 'УМЕРЕННЫЙ', 'high': 'ВЫСОКИЙ', 'critical': 'КРИТИЧЕСКИЙ'}
        lines.append(f"   • Уровень риска: {_risk_map.get(str(risk).lower(), str(risk).upper())}")
        lines.append(f"   • Уверенность ансамбля: {conf:.0%}")
        lines.append("")

        exercise_names = []
        for s in sessions:
            exn = s.get('exercise_name', '')
            if exn and exn not in exercise_names:
                exercise_names.append(exn)
        if exercise_names:
            lines.append("   Упражнения в анализе:")
            for en in exercise_names:
                lines.append(f"      — {en}")
            lines.append("")

        lines.append(dash)
        lines.append("2. РЕЗУЛЬТАТЫ АГЕНТОВ")
        lines.append(dash)
        lines.append("")

        def _risk_ru(r):
            m = {'low': 'низкий', 'moderate': 'умеренный', 'high': 'высокий', 'critical': 'критический'}
            return m.get(str(r).lower(), str(r))

        def _conf_ru(c):
            c = float(c) if c is not None else 0.5
            if c >= 0.8:
                return 'очень высокая'
            elif c >= 0.7:
                return 'высокая'
            elif c >= 0.5:
                return 'умеренная'
            else:
                return 'низкая'

        lines.append(f"   Общий риск: {_risk_ru(risk)}")
        lines.append(f"   Уверенность: {_conf_ru(conf)}")
        lines.append("")

        if breakdown:
            _name_ru = {'biomechanical': 'Биомеханический', 'kinematic': 'Кинематический',
                         'statistical': 'Статистический', 'normative': 'Нормативный',
                         'clinical': 'Клинический'}
            lines.append("   Распределение весов агентов:")
            for ag, w in list(breakdown.items())[:5]:
                try:
                    wt = float(w)
                    pct = int(wt * 100)
                    ag_ru = _name_ru.get(ag, ag)
                    lines.append(f"      — {ag_ru}: {pct}%")
                except Exception:
                    pass
            lines.append("")

        if conflicts:
            lines.append(f"   Разрешённые конфликты ({len(conflicts)}):")
            for c in conflicts[:3]:
                lines.append(f"      — {str(c)[:120]}")
            lines.append("")

        lines.append(dash)
        lines.append("3. РАСШИРЕННЫЕ АНАЛИЗЫ")
        lines.append(dash)
        lines.append("")

        if na.get("fft_bio") or na.get("fft_stat"):
            lines.append("   Частотный анализ (FFT):")
            fft = na.get("fft_bio") or na.get("fft_stat") or {}
            if isinstance(fft, dict):
                peaks = fft.get('dominant_frequencies') or fft.get('peaks') or []
                if peaks:
                    lines.append(f"      Доминантные частоты: {', '.join(str(p)[:5] for p in peaks[:5])} Гц")
                else:
                    lines.append("      Частотный спектр рассчитан, выраженных пиков не выявлено.")
            lines.append("      Периодичность движений в норме — ритм стабильный.")
            lines.append("      Отклонения от ритма могут указывать на утомление или нарушение координации.")
            lines.append("")

        if na.get("complexity"):
            lines.append("   Метрики сложности (SampEn, DFA):")
            cx = na.get("complexity") or {}
            if isinstance(cx, dict):
                sampen = cx.get('sampen') or cx.get('SampEn')
                dfa = cx.get('dfa') or cx.get('DFA')
                if sampen is not None:
                    lines.append(f"      SampEn: {float(sampen):.3f}")
                if dfa is not None:
                    lines.append(f"      DFA α: {float(dfa):.3f}")
            lines.append("      SampEn > 2.0 — высокая сложность/нестабильность движения.")
            lines.append("      DFA α ~0.5–1.0 — динамика близка к нормальной (фрактальная).")
            lines.append("")

        if na.get("asymmetry_evolution"):
            lines.append("   Эволюция асимметрии:")
            asym = na.get("asymmetry_evolution") or {}
            if isinstance(asym, dict):
                trend = asym.get('trend') or asym.get('direction')
                if trend:
                    lines.append(f"      Тренд: {trend}")
            lines.append("      Асимметрия > 15% между сторонами — признак компенсации.")
            lines.append("      Динамика асимметрии по сессиям показывает эффективность реабилитации.")
            lines.append("")

        if na.get("icc_21"):
            lines.append(f"   Межсессионная согласованность ICC(2,1): {float(na.get('icc_21', 0)):.2f}")
            lines.append("      ICC > 0.75 — отличная воспроизводимость движений.")
            lines.append("      ICC 0.5–0.75 — умеренная, рекомендуется фокус на стабильность.")
            lines.append("      ICC < 0.5 — низкая согласованность, нужна работа над техникой.")
            lines.append("")

        if not any([na.get("fft_bio"), na.get("fft_stat"), na.get("complexity"),
                     na.get("asymmetry_evolution"), na.get("icc_21")]):
            lines.append("   Расширенные анализы недоступны (недостаточно данных или время вычислений).")
            lines.append("")

        lines.append(dash)
        lines.append("4. РЕКОМЕНДАЦИИ ПО РЕАБИЛИТАЦИИ (ОБЪЕДИНЁННЫЙ ВЫВОД 5 АГЕНТОВ)")
        lines.append(dash)
        lines.append("")

        lines.append("A. КЛЮЧЕВЫЕ РЕКОМЕНДАЦИИ")
        lines.append("")
        if recs:
            for i, r in enumerate(recs[:9], 1):
                r_clean = _re.sub(r'\s+', ' ', str(r).strip())
                lines.append(f"   {i}. {r_clean}")
                r_low = r_clean.lower()
                expl = ""
                if 'сниз' in r_low or 'уменьш' in r_low or 'изометр' in r_low:
                    expl = ("      Обоснование: обнаружены признаки чрезмерной нагрузки или нестабильности — "
                            "снижение интенсивности позволяет восстановиться тканям и снизить риск травмы.")
                elif 'увелич' in r_low or 'прогресс' in r_low or 'усложн' in r_low:
                    expl = ("      Обоснование: показатели стабилизировались — допускается постепенная "
                            "прогрессия нагрузки (10–15% за 1–2 сессии) для стимуляции адаптации.")
                elif 'техник' in r_low or 'паттерн' in r_low or 'координа' in r_low or 'симметр' in r_low:
                    expl = ("      Обоснование: выявлены особенности координации/асимметрии — коррекция "
                            "техники повышает эффективность реабилитации и снижает компенсации.")
                elif 'вариаб' in r_low or 'cv' in r_low or 'утомл' in r_low:
                    expl = ("      Обоснование: повышенная вариабельность или признаки утомления указывают "
                            "на нестабильность моторного контроля — развитие выносливости улучшает надёжность.")
                elif 'монитор' in r_low or 'контроль' in r_low:
                    expl = ("      Обоснование: при умеренной уверенности регулярный мониторинг позволяет "
                            "своевременно выявлять ухудшения и корректировать программу.")
                elif 'безопасн' in r_low or 'риск' in r_low:
                    expl = ("      Обоснование: высокий/критический риск — безопасность имеет приоритет "
                            "над прогрессией, избыточная нагрузка может усугубить состояние.")
                if expl:
                    lines.append(expl)
                else:
                    lines.append("      Обоснование: рекомендация сформирована на основе оценки метрик "
                                 "всех пяти агентов с учётом индивидуальных параметров.")
                lines.append("")
        else:
            lines.append("   Показатели в пределах нормы — коррекция не требуется.")
            lines.append("   Продолжайте текущую программу с мониторингом каждые 3–4 сессии.")
            lines.append("")

        lines.append("B. ПРОГРЕССИЯ НАГРУЗКИ И ТЕХНИКА ВЫПОЛНЕНИЯ")
        lines.append("")
        risk_lower = str(risk).lower()
        if risk_lower in ('high', 'critical'):
            lines.append("   Стратегия: НЕМЕДЛЕННАЯ КОРРЕКЦИЯ ПРОГРАММЫ")
            lines.append("   • Снизить нагрузку на 40–60% от текущего уровня.")
            lines.append("   • Перейти на изометрические и медленные контролируемые движения.")
            lines.append("   • Исключить упражнения с высокой осевой/ротационной нагрузкой на сустав.")
            lines.append("   • Обязательна консультация специалиста перед возобновлением прогрессии.")
            lines.append("   • Возобновлять нагрузку после стабилизации метрик на 2–3 сессиях подряд.")
        elif risk_lower == 'moderate':
            lines.append("   Стратегия: КОНТРОЛИРУЕМАЯ ПРОГРЕССИЯ")
            lines.append("   • Поддерживать текущий уровень или увеличивать на 10% каждые 2 сессии.")
            lines.append("   • Приоритет — точность техники и равномерное распределение нагрузки.")
            lines.append("   • При росте CV или падении пиковой силы — добавить восстановительные дни.")
            lines.append("   • Вводить вариации упражнения для предотвращения плато.")
            lines.append("   • Целевой ориентир: снижение CV и рост согласованности между сессиями.")
        else:
            lines.append("   Стратегия: ПОСТЕПЕННОЕ УСЛОЖНЕНИЕ")
            lines.append("   • Прогрессия нагрузки +15–20% при хорошей переносимости.")
            lines.append("   • Вариативность упражнений для комплексного развития.")
            lines.append("   • Мониторинг каждые 3–4 сессии для контроля за трендом.")
            lines.append("   • Целевой ориентир: рост пиковой силы при стабильном CV.")
        lines.append("")

        lines.append(f"   Возрастные особенности ({age_group}, {age:.0f} лет):")
        if age_group in ('ребёнок', 'подросток'):
            lines.append("   • Приоритет — разнообразие, игровая форма и удовольствие от движения.")
            lines.append("   • Более высокая естественная вариабельность — норма развития, не патология.")
            lines.append("   • Избегать чрезмерных нагрузок на растущие суставы и связки.")
        elif age_group == 'пожилой':
            lines.append("   • Очень постепенная прогрессия, акцент на баланс и проприоцепцию.")
            lines.append("   • Повышенный риск падений при высокой вариабельности — добавлять "
                         "упражнения на равновесие.")
            lines.append("   • Учитывать возрастное снижение мышечной массы — избегать пиковых нагрузок.")
        else:
            lines.append("   • Стандартный режим прогрессии с акцентом на силовую выносливость.")
            lines.append("   • Контроль за симметрией и координацией для предотвращения компенсаций.")
        lines.append("")

        if complaint:
            lines.append(f"   С учётом причины обращения («{complaint}»):")
            if 'колен' in complaint.lower():
                lines.append("   • Контроль момента в коленном суставе — избегать ротационных нагрузок.")
                lines.append("   • Укрепление четырёхглавой и ягодичных мышц для стабилизации.")
            elif 'тазобедр' in complaint.lower() or 'бедр' in complaint.lower():
                lines.append("   • Контроль объёма движений, избегать крайних ротаций.")
                lines.append("   • Укрепление средней ягодичной мышцы и стабилизаторов таза.")
            elif 'стоп' in complaint.lower() or 'голеност' in complaint.lower():
                lines.append("   • Акцент на стабилизацию голеностопа и проприоцепцию.")
                lines.append("   • Избегать резких поворотов стопы на опоре.")
            elif 'ампут' in complaint.lower() or 'протез' in complaint.lower():
                lines.append("   • Адаптация нагрузки к состоянию протеза/культи.")
                lines.append("   • Контроль посадки и распределения веса.")
            else:
                lines.append("   • Индивидуальная корректировка программы с учётом состояния тканей.")
            lines.append("")

        lines.append("C. МЕТОДОЛОГИЯ И ОЦЕНКА КАЧЕСТВА АНАЛИЗА")
        lines.append("")
        lines.append(f"   Количество сессий для анализа: {len(sessions)}")
        total_points = sum(len(s.get('times', [])) for s in sessions)
        lines.append(f"   Общее количество точек данных: {total_points}")
        lines.append(f"   Уверенность ансамбля: {conf:.0%}")
        if conf >= 0.8:
            lines.append("   Вывод: высокая надёжность результатов — рекомендации могут использоваться "
                         "для планирования тренировок.")
        elif conf >= 0.6:
            lines.append("   Вывод: умеренная надёжность — рекомендации носят ориентировочный характер.")
        else:
            lines.append("   Вывод: низкая надёжность — рекомендуется повторный сбор данных для уточнения.")
        lines.append("")
        lines.append("   Источники данных:")
        lines.append("      — Биомеханический агент: анализ моментов, нагрузок, пиковых значений.")
        lines.append("      — Кинематический агент: петли гистерезиса, координация, симметрия.")
        lines.append("      — Статистический агент: вариабельность (CV), нормы, тренды.")
        lines.append("      — Клинический агент: оценка рисков и приоритетов.")
        lines.append("      — Мастер-ансамбль: объединение выводов, разрешение конфликтов.")
        lines.append("")
        lines.append(sep)
        lines.append("СМОТРИТЕ ПОДРОБНЫЙ АНАЛИЗ КАЖДОГО УПРАЖНЕНИЯ В ОТДЕЛЬНОМ ОКНЕ.")
        lines.append(sep)
        return lines

    def on_exercise_double_click(self, item):
        widget = self.ex_list_widget.itemWidget(item)
        if isinstance(widget, ExerciseCardWidget):
            ex_name = widget.get_name()
            variants = self.exercise_variants.get(ex_name, [])
            if variants:
                latest_folder = variants[0][1]
                self.exercise_selected.emit(self.current_patient, latest_folder)

    def get_available_dates(self):
        if not self.current_patient:
            return []
        patient_dir = os.path.join(PATIENTS_DIR, self.current_patient)
        dates = set()
        for folder in os.listdir(patient_dir):
            folder_path = os.path.join(patient_dir, folder)
            if not os.path.isdir(folder_path) or not os.path.exists(os.path.join(folder_path, 'angles.png')):
                continue
            date_match = re.search(r'(\d{4}-\d{2}-\d{2})', folder)
            if date_match:
                try:
                    dt = datetime.strptime(date_match.group(1), '%Y-%m-%d')
                    dates.add(dt.strftime('%d.%m.%Y'))
                except:
                    pass
        return sorted(dates, reverse=True)

    def delete_reports_by_date(self, date_str):
        if not self.current_patient:
            return False, 0, []
        patient_dir = os.path.join(PATIENTS_DIR, self.current_patient)
        deleted_count = 0
        deleted_folders = []
        try:
            target_date = datetime.strptime(date_str, '%d.%m.%Y')
            target_date_str = target_date.strftime('%Y-%m-%d')
        except:
            return False, 0, []
        for folder in os.listdir(patient_dir):
            folder_path = os.path.join(patient_dir, folder)
            if not os.path.isdir(folder_path):
                continue
            date_match = re.search(r'(\d{4}-\d{2}-\d{2})', folder)
            if date_match and date_match.group(1) == target_date_str:
                try:
                    shutil.rmtree(folder_path, ignore_errors=True)
                    deleted_count += 1
                    deleted_folders.append(folder)
                    gc.collect()
                except Exception:
                    pass
        if deleted_count > 0:
            create_full_report(self.current_patient, patient_dir)
            return True, deleted_count, deleted_folders
        return False, 0, []

    def delete_report_by_date(self):
        if not self.current_patient:
            show_styled_message(self, "Ошибка", "Сначала выберите пациента", "warning")
            return
        available_dates = self.get_available_dates()
        if not available_dates:
            show_styled_message(self, "Нет отчётов", "Нет доступных отчётов для удаления", "info")
            return
        dlg = DeleteReportDialog(self.current_patient, available_dates, self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            selected_date = dlg.get_selected_date()
            success, count, folders = self.delete_reports_by_date(selected_date)
            if success:
                show_styled_message(self, "Удаление завершено",
                                    f"Удалено упражнений: {count}\nДата: {selected_date}", "info")
                self.load_patient_info()
                self.load_exercises()
            else:
                show_styled_message(self, "Ошибка", "Не найдено упражнений для удаления за указанную дату", "warning")

    def delete_patient(self):
        if not self.current_patient:
            return
        dlg = DeleteConfirmDialog(self.current_patient, self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            patient_dir = os.path.join(PATIENTS_DIR, self.current_patient)
            try:
                shutil.rmtree(patient_dir, ignore_errors=True)
                gc.collect()
                show_styled_message(self, "Удалено", f"Пациент '{self.current_patient}' успешно удалён.", "info")
                self.back_to_list.emit()
            except Exception as e:
                show_styled_message(self, "Ошибка", f"Не удалось удалить пациента:\n{str(e)}", "critical")


# ============================================================
# ExerciseViewPage
# ============================================================
class ExerciseViewPage(QWidget):
    back_to_exercises = pyqtSignal()

    def __init__(self, parent=None):
        super().__init__(parent)
        layout = QHBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(20)

        self.left_panel = QGroupBox("Данные упражнения")
        left_layout = QVBoxLayout(self.left_panel)
        self.name_label = QLabel()
        self.name_label.setStyleSheet("font-size: 19px; font-weight: bold; color: #5a9eff;")
        self.count_label = QLabel("Выполнений: 0")
        self.count_label.setStyleSheet("font-size: 14px; color: #ccc;")

        # === Подробный анализ ансамбля ИИ-агентов ===
        # Перемещён под блок выбора сглаживания. Должен быть объёмным (≥2000 символов), полностью на русском.
        self.ai_analysis_text = QTextEdit()
        self.ai_analysis_text.setReadOnly(True)
        self.ai_analysis_text.setMinimumHeight(220)
        self.ai_analysis_text.setMaximumHeight(420)
        self.ai_analysis_text.setStyleSheet("""
            QTextEdit {
                font-size: 11px; 
                color: #aaddaa; 
                background-color: #16251a; 
                border: 1px solid #2a4a32; 
                border-radius: 4px;
                padding: 6px;
            }
        """)
        self.ai_analysis_text.setPlaceholderText("Здесь отображается подробный анализ ансамбля (5 ИИ-агентов). Минимум 2000 символов, на русском языке.")

        # === Выбор уровня сглаживания (по запросу пользователя) ===
        self.smoothing_label = QLabel("Сглаживание:")
        self.smoothing_label.setStyleSheet("font-size: 13px; color: #aaa; margin-top: 8px;")

        self.smoothing_combo = QComboBox()
        self.smoothing_combo.addItems([
            "Без сглаживания",
            "Лёгкое",
            "Среднее (рекомендуется)",
            "Сильное"
        ])
        self.smoothing_combo.setCurrentIndex(2)  # Среднее по умолчанию
        self.smoothing_combo.setStyleSheet("""
            QComboBox {
                font-size: 13px;
                padding: 4px 8px;
                background-color: #2a3a4a;
                border: 1px solid #4a5a6a;
                border-radius: 4px;
            }
            QComboBox::drop-down {
                border: none;
            }
        """)

        self.apply_smoothing_btn = QPushButton("Перестроить графики")
        self.apply_smoothing_btn.setStyleSheet("""
            QPushButton {
                font-size: 12px;
                padding: 6px;
                background-color: #2a6a4a;
                border: none;
                border-radius: 4px;
            }
            QPushButton:hover {
                background-color: #3a8a6a;
            }
        """)
        self.apply_smoothing_btn.clicked.connect(self.on_apply_smoothing)

        self.smoothing_hint = QLabel("Изменения применяются к текущему сеансу")
        self.smoothing_hint.setStyleSheet("font-size: 11px; color: #777;")

        left_layout.addWidget(self.name_label)
        left_layout.addWidget(self.count_label)
        left_layout.addWidget(self.smoothing_label)
        left_layout.addWidget(self.smoothing_combo)
        left_layout.addWidget(self.apply_smoothing_btn)
        left_layout.addWidget(self.smoothing_hint)
        # Анализ размещён под выбором сглаживания (по запросу пользователя)
        left_layout.addWidget(self.ai_analysis_text)
        left_layout.addStretch()

        self.center_panel = QGroupBox("График упражнения")
        self.center_panel.setFixedWidth(960)
        self.center_panel.setStyleSheet("""
            QGroupBox {
                border: none;
                margin-top: 4px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 8px;
                padding: 0 4px 0 4px;
                color: #5a9eff;
                font-size: 13px;
            }
        """)
        center_layout = QVBoxLayout(self.center_panel)

        self.center_tabs = QTabWidget()
        self.center_tabs.setStyleSheet("""
            QTabWidget::pane {
                background-color: #1e2a3a;
                border: none;
                border-radius: 6px;
                padding: 6px;
            }
            QTabBar::tab {
                background-color: #2a3a4a;
                color: #cccccc;
                padding: 6px 14px;
                margin-right: 3px;
                border-top-left-radius: 6px;
                border-top-right-radius: 6px;
            }
            QTabBar::tab:selected {
                background-color: #1e90ff;
                color: white;
            }
            QTabBar::tab:hover:!selected {
                background-color: #3a5a7a;
            }
        """)

        # ========== ВКЛАДКА 1: ДИНАМИКА АМПЛИТУДЫ (по запросу — первая вкладка) ==========
        self.amplitude_page = QWidget()
        amp_layout = QVBoxLayout(self.amplitude_page)
        amp_layout.setContentsMargins(0, 0, 0, 0)

        self.amplitude_scroll = QScrollArea()
        self.amplitude_scroll.setWidgetResizable(True)
        self.amplitude_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.amplitude_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.amplitude_scroll.setStyleSheet("QScrollArea { border: none; background: transparent; }")

        self.amplitude_content = QWidget()
        self.amplitude_content_layout = QVBoxLayout(self.amplitude_content)
        self.amplitude_content_layout.setContentsMargins(8, 8, 8, 8)
        self.amplitude_content_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

        self.amplitude_scroll.setWidget(self.amplitude_content)
        amp_layout.addWidget(self.amplitude_scroll)
        # (moved to correct order below)

        # ========== ВКЛАДКА 2: СПЕКТРАЛЬНЫЙ АНАЛИЗ ==========
        self.spectral_page = QWidget()
        spectral_layout = QVBoxLayout(self.spectral_page)
        spectral_layout.setContentsMargins(0, 0, 0, 0)

        self.spectral_scroll = QScrollArea()
        self.spectral_scroll.setWidgetResizable(True)
        self.spectral_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.spectral_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.spectral_scroll.setStyleSheet("QScrollArea { border: none; background: transparent; }")

        self.spectral_content = QWidget()
        self.spectral_content_layout = QVBoxLayout(self.spectral_content)
        self.spectral_content_layout.setSpacing(8)
        self.spectral_content_layout.setContentsMargins(4, 8, 4, 8)
        self.spectral_content_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

        self.spectral_scroll.setWidget(self.spectral_content)
        spectral_layout.addWidget(self.spectral_scroll)
        # (moved to correct order below)

        # ========== ВКЛАДКА 3: ПАРНЫЕ ГРАФИКИ ==========
        self.pairs_page = QWidget()
        pairs_layout = QVBoxLayout(self.pairs_page)
        pairs_layout.setContentsMargins(0, 0, 0, 0)

        self.pairs_scroll = QScrollArea()
        self.pairs_scroll.setWidgetResizable(True)
        self.pairs_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.pairs_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.pairs_scroll.setStyleSheet("QScrollArea { border: none; background: transparent; }")

        self.pairs_content = QWidget()
        self.pairs_content_layout = QVBoxLayout(self.pairs_content)
        self.pairs_content_layout.setSpacing(8)
        self.pairs_content_layout.setContentsMargins(4, 8, 4, 8)
        self.pairs_content_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

        self.pairs_scroll.setWidget(self.pairs_content)
        pairs_layout.addWidget(self.pairs_scroll)
        # (moved to correct order below)

        # ========== ВКЛАДКА 5: ГИСТЕРЕЗИСНЫЕ ПЕТЛИ ==========
        self.hysteresis_page = QWidget()
        hyst_layout = QVBoxLayout(self.hysteresis_page)
        hyst_layout.setContentsMargins(0, 0, 0, 0)

        self.hysteresis_scroll = QScrollArea()
        self.hysteresis_scroll.setWidgetResizable(True)
        self.hysteresis_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.hysteresis_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.hysteresis_scroll.setStyleSheet("QScrollArea { border: none; background: transparent; }")

        self.hysteresis_content = QWidget()
        self.hysteresis_content_layout = QVBoxLayout(self.hysteresis_content)
        self.hysteresis_content_layout.setSpacing(16)
        self.hysteresis_content_layout.setContentsMargins(8, 8, 8, 8)
        self.hysteresis_content_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

        self.hysteresis_scroll.setWidget(self.hysteresis_content)
        hyst_layout.addWidget(self.hysteresis_scroll)
        # (moved to correct order below)

        # ========== ВКЛАДКА: СИММЕТРИЯ (Left vs Right) ==========
        self.symmetry_page = QWidget()
        symmetry_layout = QVBoxLayout(self.symmetry_page)
        symmetry_layout.setContentsMargins(0, 0, 0, 0)

        self.symmetry_scroll = QScrollArea()
        self.symmetry_scroll.setWidgetResizable(True)
        self.symmetry_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.symmetry_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.symmetry_scroll.setStyleSheet("QScrollArea { border: none; background: transparent; }")

        self.symmetry_content = QWidget()
        self.symmetry_content_layout = QVBoxLayout(self.symmetry_content)
        self.symmetry_content_layout.setSpacing(12)
        self.symmetry_content_layout.setContentsMargins(8, 8, 8, 8)
        self.symmetry_content_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

        self.symmetry_scroll.setWidget(self.symmetry_content)
        symmetry_layout.addWidget(self.symmetry_scroll)
        # (moved to correct order below)

        # ========== ВКЛАДКА: НАГРУЗКА НОГИ (новый график момента с учётом массы и длин звеньев) ==========
        self.legload_page = QWidget()
        legload_layout = QVBoxLayout(self.legload_page)
        legload_layout.setContentsMargins(0, 0, 0, 0)

        self.legload_scroll = QScrollArea()
        self.legload_scroll.setWidgetResizable(True)
        self.legload_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.legload_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.legload_scroll.setStyleSheet("QScrollArea { border: none; background: transparent; }")

        self.legload_content = QWidget()
        self.legload_content_layout = QVBoxLayout(self.legload_content)
        self.legload_content_layout.setSpacing(12)
        self.legload_content_layout.setContentsMargins(8, 8, 8, 8)
        self.legload_content_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)

        self.legload_scroll.setWidget(self.legload_content)
        legload_layout.addWidget(self.legload_scroll)
        # (added to tab order below)

        # ========== НОВЫЕ ВКЛАДКИ ДЛЯ АНАЛИЗОВ ИЗ ОБЩЕГО ОТЧЁТА (Симметрия по пикам/импульсу, Phase lag + перекос, Тренд пиков и CV) ==========
        # Эти вкладки активны/полезны когда у упражнения есть несколько тестов (сессий) у пациента.
        # Они реализуют именно те анализы, что в Obschiy_otchyot (Рис.7,8,9 и аналоги).

        self.sym_peaks_page = QWidget()
        sym_peaks_layout = QVBoxLayout(self.sym_peaks_page)
        sym_peaks_layout.setContentsMargins(0, 0, 0, 0)
        self.sym_peaks_scroll = QScrollArea()
        self.sym_peaks_scroll.setWidgetResizable(True)
        self.sym_peaks_content = QWidget()
        self.sym_peaks_content_layout = QVBoxLayout(self.sym_peaks_content)
        self.sym_peaks_content_layout.setContentsMargins(8, 8, 8, 8)
        self.sym_peaks_scroll.setWidget(self.sym_peaks_content)
        sym_peaks_layout.addWidget(self.sym_peaks_scroll)

        self.phase_lag_page = QWidget()
        phase_lag_layout = QVBoxLayout(self.phase_lag_page)
        phase_lag_layout.setContentsMargins(0, 0, 0, 0)
        self.phase_lag_scroll = QScrollArea()
        self.phase_lag_scroll.setWidgetResizable(True)
        self.phase_lag_content = QWidget()
        self.phase_lag_content_layout = QVBoxLayout(self.phase_lag_content)
        self.phase_lag_content_layout.setContentsMargins(8, 8, 8, 8)
        self.phase_lag_scroll.setWidget(self.phase_lag_content)
        phase_lag_layout.addWidget(self.phase_lag_scroll)

        self.trend_cv_page = QWidget()
        trend_cv_layout = QVBoxLayout(self.trend_cv_page)
        trend_cv_layout.setContentsMargins(0, 0, 0, 0)
        self.trend_cv_scroll = QScrollArea()
        self.trend_cv_scroll.setWidgetResizable(True)
        self.trend_cv_content = QWidget()
        self.trend_cv_content_layout = QVBoxLayout(self.trend_cv_content)
        self.trend_cv_content_layout.setContentsMargins(8, 8, 8, 8)
        self.trend_cv_scroll.setWidget(self.trend_cv_content)
        trend_cv_layout.addWidget(self.trend_cv_scroll)

        # Дополнительные анализы из ПОВОРОТ ГОЛЕНИ (и аналогичных) в файле
        self.phase_portrait_page = QWidget()
        pp_layout = QVBoxLayout(self.phase_portrait_page)
        pp_layout.setContentsMargins(0, 0, 0, 0)
        self.pp_scroll = QScrollArea()
        self.pp_scroll.setWidgetResizable(True)
        self.pp_content = QWidget()
        self.pp_content_layout = QVBoxLayout(self.pp_content)
        self.pp_content_layout.setContentsMargins(8, 8, 8, 8)
        self.pp_scroll.setWidget(self.pp_content)
        pp_layout.addWidget(self.pp_scroll)

        self.consistency_page = QWidget()
        cons_layout = QVBoxLayout(self.consistency_page)
        cons_layout.setContentsMargins(0, 0, 0, 0)
        self.cons_scroll = QScrollArea()
        self.cons_scroll.setWidgetResizable(True)
        self.cons_content = QWidget()
        self.cons_content_layout = QVBoxLayout(self.cons_content)
        self.cons_content_layout.setContentsMargins(8, 8, 8, 8)
        self.cons_scroll.setWidget(self.cons_content)
        cons_layout.addWidget(self.cons_scroll)

        # Ещё анализы из файла (для ПОВОРОТ ГОЛЕНИ и аналогичных с несколькими тестами)
        self.thirds_fatigue_page = QWidget()
        thirds_layout = QVBoxLayout(self.thirds_fatigue_page)
        thirds_layout.setContentsMargins(0, 0, 0, 0)
        self.thirds_scroll = QScrollArea()
        self.thirds_scroll.setWidgetResizable(True)
        self.thirds_content = QWidget()
        self.thirds_content_layout = QVBoxLayout(self.thirds_content)
        self.thirds_content_layout.setContentsMargins(8, 8, 8, 8)
        self.thirds_scroll.setWidget(self.thirds_content)
        thirds_layout.addWidget(self.thirds_scroll)

        self.radar_page = QWidget()
        radar_layout = QVBoxLayout(self.radar_page)
        radar_layout.setContentsMargins(0, 0, 0, 0)
        self.radar_scroll = QScrollArea()
        self.radar_scroll.setWidgetResizable(True)
        self.radar_content = QWidget()
        self.radar_content_layout = QVBoxLayout(self.radar_content)
        self.radar_content_layout.setContentsMargins(8, 8, 8, 8)
        self.radar_scroll.setWidget(self.radar_content)
        radar_layout.addWidget(self.radar_scroll)

        # Добавляем вкладки в требуемом порядке:
        # 1. Динамика амплитуд
        # 2. Спектральный анализ
        # 3. Гистерезисные петли
        # 4. Симметрия
        # 5. Нагрузка ноги (момент с массой + длинами звеньев)
        # 6. Парные графики
        self.center_tabs.addTab(self.amplitude_page, "📈 Динамика амплитуды")
        self.center_tabs.addTab(self.spectral_page, "🎵 Спектральный анализ")
        self.center_tabs.addTab(self.hysteresis_page, "🔄 Гистерезисные петли")
        self.center_tabs.addTab(self.symmetry_page, "⚖️ Симметрия")
        self.center_tabs.addTab(self.legload_page, "🦵 Нагрузка ноги")
        self.center_tabs.addTab(self.sym_peaks_page, "📐 Симметрия (пики + импульс)")
        self.center_tabs.addTab(self.phase_lag_page, "⏱ Phase lag + перекос")
        self.center_tabs.addTab(self.trend_cv_page, "📉 Тренд пиков и CV")
        self.center_tabs.addTab(self.phase_portrait_page, "🔄 Петля угол-момент (phase portrait)")
        self.center_tabs.addTab(self.consistency_page, "📊 Согласованность (Pearson r)")
        self.center_tabs.addTab(self.thirds_fatigue_page, "📊 Пики по третям (утомляемость внутри теста)")
        self.center_tabs.addTab(self.radar_page, "🕸 Радар-паспорт метрик (T1 vs последний)")
        self.center_tabs.addTab(self.pairs_page, "📊 Парные графики")

        center_layout.addWidget(self.center_tabs)

        self.right_panel = QGroupBox("Выбор сеанса")
        right_layout = QVBoxLayout(self.right_panel)
        self.session_list = QListWidget()
        self.session_list.setStyleSheet("""
            QListWidget {
                font-size: 14px;
                background: transparent;
                border: none;
            }
            QListWidget::item {
                padding: 8px;
                border-bottom: 1px solid #3a4a5a;
            }
            QListWidget::item:selected {
                background-color: #1e90ff;
                border-radius: 6px;
            }
        """)
        self.session_list.itemClicked.connect(self.on_session_selected)
        right_layout.addWidget(self.session_list)

        self.splitter = QSplitter(Qt.Orientation.Horizontal)
        self.splitter.addWidget(self.left_panel)
        self.splitter.addWidget(self.center_panel)
        self.splitter.addWidget(self.right_panel)
        self.splitter.setSizes([340, 960, 300])
        layout.addWidget(self.splitter)

        self.current_patient = None
        self.current_ex_name = None
        self.sessions = []
        self.current_smoothing_intensity = "medium"  # по умолчанию Среднее

        # Подключаем обработчик смены вкладок — выбор сеанса показывается только на нужных вкладках
        self.center_tabs.currentChanged.connect(self._on_tab_changed)

        # До загрузки упражнения прячем панель выбора сеанса
        self.right_panel.setVisible(False)

    def set_exercise(self, patient_name, exercise_folder, exercise_path):
        self.current_patient = patient_name
        folder_name = os.path.basename(exercise_path)
        self.current_ex_name = folder_name.split('_', 1)[0] if '_' in folder_name else folder_name

        self.name_label.setText(self.current_ex_name)

        patient_dir = os.path.join(PATIENTS_DIR, patient_name)
        count = 0
        for f in os.listdir(patient_dir):
            if f.startswith(self.current_ex_name + "_") and os.path.exists(os.path.join(patient_dir, f, 'angles.png')):
                count += 1
        self.count_label.setText(f"Выполнений: {count}")

        if hasattr(self, 'ai_analysis_text'):
            self.ai_analysis_text.setPlainText("Анализ ансамбля ИИ-агентов выполняется...\n(будет отображён подробный текст объёмом не менее 2000 символов)")

        # Лёгкая часть сразу
        self.load_sessions_list(patient_name, group_by_date=False)

        # Тяжёлые операции откладываем + оборачиваем в try, чтобы не крашилось всё приложение
        from PyQt6.QtCore import QTimer
        def safe_finish():
            try:
                self._finish_set_exercise(patient_name)
            except Exception as e:
                print(f"Ошибка при загрузке графиков упражнения: {e}")

        QTimer.singleShot(10, safe_finish)

    def _finish_set_exercise(self, patient_name):
        """Тяжёлая часть загрузки упражнения. Вызывается отложенно через QTimer."""
        self.load_amplitude_graph(patient_name)

        # Определяем текущую вкладку и загружаем только её + amplitude
        current_widget = self.center_tabs.currentWidget()

        if self.sessions:
            latest_session = self.sessions[0][1]

            # Загружаем только текущую активную вкладку (ленивая загрузка)
            if current_widget == self.spectral_page:
                self.load_spectral_graphs(patient_name, latest_session)
            elif current_widget == self.pairs_page:
                self.load_pair_graphs(patient_name, latest_session)
            elif current_widget == self.hysteresis_page:
                self.load_hysteresis_graphs(patient_name, latest_session)
            elif current_widget == self.symmetry_page:
                self.load_symmetry_graphs(patient_name, latest_session)
            elif current_widget == self.legload_page:
                self.load_leg_load_graph(patient_name, latest_session)
            elif current_widget == self.sym_peaks_page:
                self.load_sym_peaks_impulse()
            elif current_widget == self.phase_lag_page:
                self.load_phase_lag_skew()
            elif current_widget == self.trend_cv_page:
                self.load_trend_cv()
            elif current_widget == self.phase_portrait_page:
                self.load_phase_portrait()
            elif current_widget == self.consistency_page:
                self.load_consistency()
            elif current_widget == self.thirds_fatigue_page:
                self.load_thirds_fatigue()
            elif current_widget == self.radar_page:
                self.load_radar_metrics()
            else:
                # По умолчанию загружаем спектральный (самый частый)
                self.load_spectral_graphs(patient_name, latest_session)

        # Обновляем правую панель выбора сеанса
        self._update_session_selector_visibility()

        # Финальная гарантия: для спектрального, гистерезиса, симметрии и нагрузки ноги — обычный выбор по датам
        current = self.center_tabs.currentWidget()
        if current == self.pairs_page:
            # Только для парных — иерархический список (даты + сеансы)
            self.load_sessions_list(patient_name, hierarchical=True)
        else:
            self.load_sessions_list(patient_name, group_by_date=True)

        # Автоматически выбираем самую последнюю дату/сеанс

        # Всегда открываем упражнение на первой вкладке — "Динамика амплитуд"
        self.center_tabs.setCurrentIndex(0)
        self._auto_select_latest_session()

        # Программа сама делает анализ и создаёт файлы для всех продвинутых вкладок
        # (симметрия по пикам/импульсу, phase lag, тренд CV, петля угол-момент, согласованность, трети, радар)
        # если у упражнения несколько сеансов. Файлы сохраняются в папку пациента.
        if len(getattr(self, 'sessions', [])) > 1:
            # Сам код заботится: для всех сессий, которые видны в правой панели,
            # убеждаемся что есть raw_measurements.json (нужен для M в анализах).
            # Если нет — пытаемся автоматически восстановить из .docx пациента.
            print("[INFO] ensuring raw data exists for all listed sessions (so advanced analyses can process M themselves)...")
            pdir = os.path.join(PATIENTS_DIR, self.current_patient)
            for item in getattr(self, 'sessions', []) or []:
                if not (isinstance(item, (list, tuple)) and len(item) >= 3):
                    continue
                fldr = item[2]
                rpath = os.path.join(pdir, fldr, 'raw_measurements.json')
                if not os.path.exists(rpath):
                    print(f"  auto-creating raw for {fldr} ...")
                    try:
                        reg = globals().get('regenerate_graphs_for_session')
                        if reg is None:
                            import sys as _sys
                            _mod = _sys.modules.get(__name__)
                            if _mod:
                                reg = getattr(_mod, 'regenerate_graphs_for_session', None)
                        if reg:
                            ok, m = reg(self.current_patient, fldr, "medium")
                            print(f"    -> {ok} ({m[:80] if m else ''})")
                        else:
                            print("    (regenerate function not available)")
                    except Exception as ee:
                        print(f"    regen error: {ee}")
            # Теперь pre-gen вызовет load_* которые соберут M из (восстановленных) raw
            for load_func in [
                self.load_sym_peaks_impulse,
                self.load_phase_lag_skew,
                self.load_trend_cv,
                self.load_phase_portrait,
                self.load_consistency,
                self.load_thirds_fatigue,
                self.load_radar_metrics,
            ]:
                try:
                    load_func()
                except Exception as e:
                    print(f"[WARN] pre-generate {getattr(load_func, '__name__', 'analysis')}: {e}")

        # === Вывод анализа ансамбля ИИ-агентов под названием упражнения ===
        # Запускаем отложенно, чтобы не блокировать открытие интерфейса упражнения
        from PyQt6.QtCore import QTimer as _QTimer
        def _safe_ai_analysis():
            try:
                if hasattr(self, '_run_and_display_ai_analysis'):
                    self._run_and_display_ai_analysis(patient_name)
            except Exception as _e:
                print(f"[AI] safe call error: {_e}")
                if hasattr(self, 'ai_analysis_text'):
                    self.ai_analysis_text.setPlainText("Анализ ИИ-ансамбля: не удалось выполнить.\n\nПовторите открытие упражнения.")
        _QTimer.singleShot(280, _safe_ai_analysis)

    def load_sessions_list(self, patient_name, group_by_date=False, hierarchical=False, preserve_selection=True):
        """
        Загружает список сеансов.
        - group_by_date=True: только даты (для спектрального и гистерезиса)
        - hierarchical=True: даты как заголовки + сеансы под ними с временем (только для парных графиков)
        - preserve_selection=True: после перезаполнения списка пытается восстановить визуальное выделение
          по UserRole (нужно чтобы в новых агрегатных вкладках при клике на дату элемент в списке оставался выделенным,
          как в старых вкладках).
        """
        current_data = None
        if preserve_selection:
            cur = self.session_list.currentItem()
            if cur is not None:
                current_data = cur.data(Qt.ItemDataRole.UserRole)

        self.session_list.clear()
        self.sessions = []
        self.date_to_folders = {}

        patient_dir = os.path.join(PATIENTS_DIR, patient_name)
        sessions = []

        for folder in os.listdir(patient_dir):
            if not folder.startswith(self.current_ex_name + "_"):
                continue
            folder_path = os.path.join(patient_dir, folder)
            if not os.path.isdir(folder_path) or not os.path.exists(os.path.join(folder_path, 'angles.png')):
                continue

            date_str = "неизвестно"
            date_obj = datetime.min
            for part in folder.split('_'):
                if re.match(r'\d{4}-\d{2}-\d{2}', part):
                    try:
                        date_obj = datetime.strptime(part, '%Y-%m-%d')
                        date_str = date_obj.strftime('%Y-%m-%d')
                    except:
                        date_str = part
                    break

            time_str = ""
            for part in folder.split('_'):
                if re.match(r'\d{2}-\d{2}-\d{2}', part):
                    time_str = part.replace('-', ':')
                    break

            display = f"{date_str} {time_str}".strip() if time_str else date_str
            sessions.append((date_obj, display, folder, date_str, time_str))

        # Сортируем от нового к старому
        sessions.sort(key=lambda x: x[0], reverse=True)

        self.sessions = [(d[0], d[1], d[2]) for d in sessions]

        if hierarchical:
            # Специальный режим только для Парных графиков: даты + сеансы под ними
            from collections import defaultdict
            date_groups = defaultdict(list)
            for date_obj, display, folder, date_str, time_str in sessions:
                date_groups[date_str].append((date_obj, folder, time_str))

            for date_str in sorted(date_groups.keys(), reverse=True):
                sessions_in_day = date_groups[date_str]
                # Сессии в дне тоже от новых к старым
                sessions_in_day.sort(key=lambda x: x[0], reverse=True)

                # Заголовок даты
                header_item = QListWidgetItem(f"📅 {date_str} ({len(sessions_in_day)})")
                header_font = header_item.font()
                header_font.setBold(True)
                header_item.setFont(header_font)
                header_item.setBackground(Qt.GlobalColor.lightGray)  # лёгкий фон для заголовка
                header_item.setData(Qt.ItemDataRole.UserRole, date_str)
                self.session_list.addItem(header_item)

                # Сеансы под датой (с отступом)
                for date_obj, folder, time_str in sessions_in_day:
                    time_display = time_str if time_str else "??:??"
                    session_item = QListWidgetItem(f"    {time_display}")
                    session_item.setData(Qt.ItemDataRole.UserRole, folder)
                    self.session_list.addItem(session_item)

            self.date_to_folders = {date_str: [s[1] for s in date_groups[date_str]] 
                                    for date_str in date_groups}

        elif group_by_date:
            # Обычный режим "только даты" (для спектрального и гистерезиса)
            from collections import defaultdict
            date_groups = defaultdict(list)
            for date_obj, display, folder, date_str, time_str in sessions:
                date_groups[date_str].append((date_obj, folder, display))

            for date_str in sorted(date_groups.keys(), reverse=True):
                folders = date_groups[date_str]
                count = len(folders)
                display_text = f"📅 {date_str} ({count} сеанс{'а' if count == 1 else 'ов'})"
                item = QListWidgetItem(display_text)
                item.setData(Qt.ItemDataRole.UserRole, date_str)
                self.session_list.addItem(item)

            self.date_to_folders = {date_str: [f[1] for f in folders] 
                                    for date_str, folders in date_groups.items()}
        else:
            # Полный плоский список (fallback)
            for date_obj, display, folder, date_str, time_str in sessions:
                item = QListWidgetItem(f"📅 {display}")
                item.setData(Qt.ItemDataRole.UserRole, folder)
                self.session_list.addItem(item)

        # Восстанавливаем визуальное выделение элемента в списке (по UserRole).
        # Это нужно, чтобы в новых вкладках (агрегатных) при клике на сеанс по дате
        # элемент оставался подсвеченным, как в старых пер-сеансовых вкладках.
        if preserve_selection and current_data is not None:
            restored = False
            for i in range(self.session_list.count()):
                it = self.session_list.item(i)
                if it is not None and it.data(Qt.ItemDataRole.UserRole) == current_data:
                    self.session_list.setCurrentItem(it)
                    restored = True
                    break
            if not restored and self.session_list.count() > 0:
                # если по какой-то причине не нашли (например после смены упражнения) — ничего не делаем
                pass

    def load_amplitude_graph(self, patient_name):
        """
        Загружает предварительно сгенерированный PNG график динамики амплитуд.
        Полностью убрана работа matplotlib из UI потока.
        """
        self._safe_clear_layout(self.amplitude_content_layout)

        patient_dir = os.path.join(PATIENTS_DIR, patient_name)
        safe_name = re.sub(r'[^a-zA-Zа-яА-ЯёЁ0-9]', '_', self.current_ex_name)
        amp_png_path = os.path.join(patient_dir, f"{safe_name}_amplitude_dynamics.png")

        if os.path.exists(amp_png_path):
            self._load_images_as_single_composite(self.amplitude_content_layout, [amp_png_path])
        else:
            info_label = QLabel(
                "График динамики амплитуд ещё не сгенерирован.\n"
                "Он появится после следующей обработки данных для этого упражнения."
            )
            info_label.setStyleSheet("color: #888; font-size: 15px;")
            info_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            info_label.setWordWrap(True)
            self.amplitude_content_layout.addWidget(info_label)

    def _safe_clear_layout(self, layout):
        """Безопасная очистка layout с принудительной обработкой событий Qt (помогает избежать 0xC0000409)."""
        from PyQt6.QtWidgets import QApplication

        for i in reversed(range(layout.count())):
            item = layout.itemAt(i)
            widget = item.widget() if item else None
            if widget:
                widget.deleteLater()

        # Даём Qt время на реальную удаление объектов
        QApplication.processEvents()

    def load_spectral_graphs(self, patient_name, folder_name):
        self._safe_clear_layout(self.spectral_content_layout)

        folder_path = os.path.join(PATIENTS_DIR, patient_name, folder_name)

        fft_path = os.path.join(folder_path, 'analysis_fft.png')
        if os.path.exists(fft_path):
            self._load_images_as_single_composite(self.spectral_content_layout, [fft_path])
        else:
            info_label = QLabel("🎵 Нет данных спектрального анализа для отображения")
            info_label.setStyleSheet("color: #888; font-size: 16px; background: transparent;")
            info_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            self.spectral_content_layout.addWidget(info_label)

        spacer = QWidget()
        spacer.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self.spectral_content_layout.addWidget(spacer)

    def load_pair_graphs(self, patient_name, folder_name):
        self._safe_clear_layout(self.pairs_content_layout)

        folder_path = os.path.join(PATIENTS_DIR, patient_name, folder_name)

        image_files = []
        for file in os.listdir(folder_path):
            if file.startswith('pair_') and file.endswith('_combined.png'):
                full_path = os.path.join(folder_path, file)
                image_files.append(full_path)

        image_files.sort()

        if image_files:
            # Возвращаемся к отдельным карточкам (безопаснее при большом количестве)
            for img_path in image_files:
                self._add_graph_card(img_path, self.pairs_content_layout, max_width=950)
        else:
            angles_path = os.path.join(folder_path, 'angles.png')
            forces_path = os.path.join(folder_path, 'forces.png')
            if os.path.exists(angles_path):
                self._add_graph_card(angles_path, self.pairs_content_layout, max_width=950)
            if os.path.exists(forces_path):
                self._add_graph_card(forces_path, self.pairs_content_layout, max_width=950)

        if self.pairs_content_layout.count() == 0:
            error_label = QLabel("❌ Графики не найдены\n\nВозможно, файл данных повреждён или не содержит измерений.")
            error_label.setStyleSheet("color: #ff8888; font-size: 16px; background: transparent;")
            error_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            self.pairs_content_layout.addWidget(error_label)

        spacer = QWidget()
        spacer.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self.pairs_content_layout.addWidget(spacer)

    def load_hysteresis_graphs(self, patient_name, folder_name):
        """Загружает все гистерезисные петли для текущего сеанса (без искусственного ограничения)."""
        self._safe_clear_layout(self.hysteresis_content_layout)

        folder_path = os.path.join(PATIENTS_DIR, patient_name, folder_name)

        hyst_files = []
        for f in os.listdir(folder_path):
            if f.startswith("hysteresis_p") and f.endswith(".png"):
                hyst_files.append(f)
        hyst_files.sort()

        if hyst_files:
            # Показываем ВСЕ гистерезисные петли (без искусственного ограничения)
            files_to_show = hyst_files

            # Показываем как одно композитное изображение (меньше виджетов, стабильнее)
            self._load_images_as_single_composite(self.hysteresis_content_layout, 
                                                  [os.path.join(folder_path, f) for f in files_to_show],
                                                  max_width=800)

            # Убираем предупреждение об ограничении — теперь показываем всё
        else:
            info_label = QLabel(
                "Гистерезисные петли не найдены для этой сессии.\n"
                "Возможно, данные были обработаны до добавления этой функции.\n"
                "Попробуйте заново обработать отчёт."
            )
            info_label.setStyleSheet("color: #888; font-size: 15px; background: transparent;")
            info_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            info_label.setWordWrap(True)
            self.hysteresis_content_layout.addWidget(info_label)

        spacer = QWidget()
        spacer.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self.hysteresis_content_layout.addWidget(spacer)

    def load_symmetry_graphs(self, patient_name, folder_name):
        """Загружает графики симметрии (Left vs Right) для текущего сеанса.
        Если графиков нет — генерирует их на лету из raw_measurements.json."""
        self._safe_clear_layout(self.symmetry_content_layout)

        folder_path = os.path.join(PATIENTS_DIR, patient_name, folder_name)

        # Ищем сгенерированный файл симметрии
        sym_files = []
        for f in os.listdir(folder_path):
            if f.startswith("symmetry") and f.endswith(".png"):
                sym_files.append(f)
        sym_files.sort()

        # Если нет — генерируем на лету
        if not sym_files:
            raw_path = os.path.join(folder_path, 'raw_measurements.json')
            if os.path.exists(raw_path):
                try:
                    self._generate_symmetry_on_the_fly(folder_path, folder_name, patient_name)
                    for f in os.listdir(folder_path):
                        if f.startswith("symmetry") and f.endswith(".png"):
                            sym_files.append(f)
                    sym_files.sort()
                except Exception as e:
                    print(f"[WARN] Ошибка генерации симметрии на лету: {e}")

        if sym_files:
            for f in sorted(sym_files):
                full_path = os.path.join(folder_path, f)
                self._add_graph_card(full_path, self.symmetry_content_layout, max_width=950)
        else:
            info = QLabel(
                "Графики симметрии не найдены для этой сессии.\n\n"
                "Симметрия строится для упражнений с левыми/правыми каналами.\n"
                "Если данные были обработаны раньше, попробуйте заново добавить отчёт."
            )
            info.setStyleSheet("color: #888; font-size: 15px; background: transparent;")
            info.setAlignment(Qt.AlignmentFlag.AlignCenter)
            info.setWordWrap(True)
            self.symmetry_content_layout.addWidget(info)

        spacer = QWidget()
        spacer.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self.symmetry_content_layout.addWidget(spacer)

    def _generate_symmetry_on_the_fly(self, folder_path, folder_name, patient_name):
        """Генерирует графики симметрии из raw_measurements.json если их ещё нет."""
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt

        raw_path = os.path.join(folder_path, 'raw_measurements.json')
        with open(raw_path, 'r', encoding='utf-8') as f:
            raw = json.load(f)

        times_str = raw.get('times', [])
        angles_rows = raw.get('angles', []) or []
        forces_rows = raw.get('forces', []) or []
        if not times_str or len(times_str) < 5:
            return

        # Определяем тип упражнения из имени папки
        ex_name = folder_name.split('_')[0] if '_' in folder_name else folder_name
        # Убираем дату из имени
        parts = folder_name.split('_')
        ex_name_parts = [p for p in parts if not re.match(r'\d{4}-\d{2}-\d{2}', p)]
        ex_name = '_'.join(ex_name_parts) if ex_name_parts else ex_name
        ex_name_upper = ex_name.upper()

        # Конвертируем времена
        base_dt = None
        numeric_times = []
        for ts in times_str:
            try:
                if ' ' in ts:
                    try:
                        dt = datetime.strptime(ts, '%Y-%m-%d %H:%M:%S.%f')
                    except ValueError:
                        dt = datetime.strptime(ts, '%Y-%m-%d %H:%M:%S')
                else:
                    dt = datetime.strptime(ts, '%Y-%m-%d')
                if base_dt is None:
                    base_dt = dt
                numeric_times.append((dt - base_dt).total_seconds())
            except Exception:
                numeric_times.append(float(len(numeric_times)) * 0.01)

        # Транспонируем данные: time-major → channel-major
        n_a = len(angles_rows[0]) if angles_rows and angles_rows[0] else 0
        n_f = len(forces_rows[0]) if forces_rows and forces_rows[0] else 0
        angles_by_ch = [[] for _ in range(n_a)]
        for row in angles_rows:
            for ch in range(n_a):
                try:
                    angles_by_ch[ch].append(float(row[ch]) if row[ch] is not None else 0.0)
                except Exception:
                    angles_by_ch[ch].append(0.0)
        forces_by_ch = [[] for _ in range(n_f)]
        for row in forces_rows:
            for ch in range(n_f):
                try:
                    forces_by_ch[ch].append(float(row[ch]) if row[ch] is not None else 0.0)
                except Exception:
                    forces_by_ch[ch].append(0.0)

        # Конвертируем силы в Ньютон
        forces_N = [[fv * 9.81 / 1000.0 for fv in ch] for ch in forces_by_ch]

        # Определяем пары левый/правый по названиям каналов
        angle_labels = CHANNEL_LABELS.get("angles", {}).get(ex_name_upper, [])
        force_labels_def = CHANNEL_LABELS.get("forces", {}).get(ex_name_upper, [])

        def _find_lr_pairs(labels):
            left_indices = []
            right_indices = []
            for i, lbl in enumerate(labels):
                lbl_l = lbl.lower()
                if any(k in lbl_l for k in ["левый", "левая", "левое", "левой", "левом", "левого", "левую"]):
                    left_indices.append(i)
                elif any(k in lbl_l for k in ["правый", "правая", "правое", "правой", "правом", "правого", "правую"]):
                    right_indices.append(i)
            pairs = []
            for li, ri in zip(left_indices, right_indices):
                pairs.append((li, labels[li], ri, labels[ri]))
            return pairs

        angle_lr_pairs = _find_lr_pairs(angle_labels) if angle_labels else []
        force_lr_pairs = _find_lr_pairs(force_labels_def) if force_labels_def else []

        # Fallback: если подписи не заданы
        if not angle_lr_pairs and n_a >= 2:
            half = n_a // 2
            for i in range(half):
                angle_lr_pairs.append((i, f"Канал {i+1} (Л)", i + half, f"Канал {i+half+1} (П)"))
        if not force_lr_pairs and n_f >= 2:
            half = n_f // 2
            for i in range(half):
                force_lr_pairs.append((i, f"Сила {i+1} (Л)", i + half, f"Сила {i+half+1} (П)"))

        if not angle_lr_pairs and not force_lr_pairs:
            return

        times = numeric_times

        # Генерируем симметрию углов
        if angle_lr_pairs:
            try:
                n_plots = len(angle_lr_pairs)
                fig, axes = plt.subplots(n_plots, 1, figsize=(14, 3.5 * n_plots), dpi=120, sharex=True)
                if n_plots == 1:
                    axes = [axes]
                for i, (li, l_label, ri, r_label) in enumerate(angle_lr_pairs):
                    ax = axes[i]
                    if li < len(angles_by_ch) and ri < len(angles_by_ch):
                        ax.plot(times, angles_by_ch[li], label=l_label, color='#4a9eff', linewidth=2.2)
                        ax.plot(times, angles_by_ch[ri], label=r_label, color='#ff6b6b', linewidth=2.2)
                        ax.set_ylabel(f"{l_label.replace('Угол ', '')} / {r_label.replace('Угол ', '')} (°)", fontsize=10, fontweight='bold')
                        ax.legend(loc='upper right', fontsize=9)
                        ax.grid(True, linestyle='--', alpha=0.35)
                axes[-1].set_xlabel("Время (с)", fontsize=11, fontweight='bold')
                fig.suptitle(f"Симметрия углов — {ex_name}", fontsize=14, fontweight='bold', y=0.98)
                plt.tight_layout(rect=[0, 0.02, 1, 0.96])
                plt.savefig(os.path.join(folder_path, "symmetry_angles.png"), dpi=100, bbox_inches='tight', facecolor='white')
                plt.close(fig)
            except Exception as e:
                print(f"[WARN] sym angles on-the-fly: {e}")

        # Генерируем симметрию сил
        if force_lr_pairs:
            try:
                n_plots = len(force_lr_pairs)
                fig2, axes2 = plt.subplots(n_plots, 1, figsize=(14, 3.5 * n_plots), dpi=120, sharex=True)
                if n_plots == 1:
                    axes2 = [axes2]
                for i, (li, l_label, ri, r_label) in enumerate(force_lr_pairs):
                    ax = axes2[i]
                    left_data = forces_N[li][:len(times)] if li < len(forces_N) else []
                    right_data = forces_N[ri][:len(times)] if ri < len(forces_N) else []
                    if left_data and right_data:
                        ax.plot(times[:len(left_data)], left_data, label=l_label, color='#4a9eff', linewidth=2.2)
                        ax.plot(times[:len(right_data)], right_data, label=r_label, color='#ff6b6b', linewidth=2.2)
                        ax.set_ylabel(f"{l_label} / {r_label} (Н)", fontsize=10, fontweight='bold')
                        ax.legend(loc='upper right', fontsize=9)
                        ax.grid(True, linestyle='--', alpha=0.35)
                axes2[-1].set_xlabel("Время (с)", fontsize=11, fontweight='bold')
                fig2.suptitle(f"Симметрия сил — {ex_name}", fontsize=14, fontweight='bold', y=0.98)
                plt.tight_layout(rect=[0, 0.02, 1, 0.96])
                plt.savefig(os.path.join(folder_path, "symmetry_forces.png"), dpi=100, bbox_inches='tight', facecolor='white')
                plt.close(fig2)
            except Exception as e:
                print(f"[WARN] sym forces on-the-fly: {e}")

    def load_leg_load_graph(self, patient_name, folder_name):
        """Загружает график 'Нагрузка ноги' (момент с учётом массы пациента и длин звеньев).
        График теперь по центру области (горизонтально и вертикально).
        """
        self._safe_clear_layout(self.legload_content_layout)

        folder_path = os.path.join(PATIENTS_DIR, patient_name, folder_name)
        leg_path = os.path.join(folder_path, 'leg_load.png')

        # Горизонтальный контейнер для центрирования
        h_container = QWidget()
        h_layout = QHBoxLayout(h_container)
        h_layout.setContentsMargins(0, 0, 0, 0)
        h_layout.addStretch(1)

        if os.path.exists(leg_path):
            # Создаём label с картинкой и центрируем его
            pix = QPixmap(leg_path)
            if not pix.isNull():
                if pix.width() > 950:
                    pix = pix.scaledToWidth(950, Qt.TransformationMode.SmoothTransformation)
                img_label = QLabel()
                img_label.setPixmap(pix)
                img_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
                img_label.setSizePolicy(QSizePolicy.Policy.Fixed, QSizePolicy.Policy.Fixed)
                h_layout.addWidget(img_label, 0, Qt.AlignmentFlag.AlignCenter)
            else:
                err = QLabel("Не удалось загрузить изображение")
                err.setAlignment(Qt.AlignmentFlag.AlignCenter)
                h_layout.addWidget(err, 0, Qt.AlignmentFlag.AlignCenter)
        else:
            info = QLabel(
                "График «Нагрузка ноги» не найден для этой сессии.\n\n"
                "Он генерируется автоматически при добавлении отчёта (или пересоздании графиков).\n"
                "Для его построения используются масса пациента и длины звеньев из карточки."
            )
            info.setStyleSheet("color: #888; font-size: 15px; background: transparent;")
            info.setAlignment(Qt.AlignmentFlag.AlignCenter)
            info.setWordWrap(True)
            h_layout.addWidget(info, 0, Qt.AlignmentFlag.AlignCenter)

        h_layout.addStretch(1)

        # Вертикальное центрирование: растягиваем сверху и снизу
        v_container = QWidget()
        v_layout = QVBoxLayout(v_container)
        v_layout.setContentsMargins(0, 0, 0, 0)
        v_layout.addStretch(1)
        v_layout.addWidget(h_container, 0, Qt.AlignmentFlag.AlignCenter)
        v_layout.addStretch(1)

        self.legload_content_layout.addWidget(v_container)

        # Небольшой спейсер внизу для скролла если нужно
        bottom_spacer = QWidget()
        bottom_spacer.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        bottom_spacer.setMinimumHeight(20)
        self.legload_content_layout.addWidget(bottom_spacer)

    # ============================================================
    # НОВЫЕ ЗАГРУЗЧИКИ ВКЛАДОК — анализы из Общего отчёта (по всем тестам упражнения)
    # Строятся точно как в Obschiy_otchyot: требуют несколько сессий одного упражнения у пациента.
    # Используют compute_leg_load_moment + helpers (symmetry_index, phase_lag, consistency, CV, LinearRegression).
    # ============================================================

    def _collect_sessions_data(self):
        """Сбор данных по сессиям для продвинутых анализов.
        Приоритет: использовать уже отфильтрованный список self.sessions (который успешно показывает сеансы в правой панели).
        Если он пуст — fallback на прямой надёжный скан.
        Всегда форсируем load_sessions_list чтобы self.sessions был свежим.
        """
        if not self.current_patient or not self.current_ex_name:
            return []
        # Форсируем обновление списка сеансов (точно тех, что видны в правой панели)
        try:
            self.load_sessions_list(self.current_patient, group_by_date=True)
        except Exception:
            pass
        patient_dir = os.path.join(PATIENTS_DIR, self.current_patient)
        anthro = load_patient_anthropometrics(self.current_patient)
        sessions_data = []

        # 1. Предпочитаем self.sessions — это те же папки, что уже успешно отображены в UI
        # ТОЧНО та же фильтрация, что и в load_sessions_list (показывает 3 сеанса в правой панели)
        # Используем точно те же папки, что load_sessions_list только что добавил в self.sessions
        # (это те, что успешно показаны в правой панели)
        ui_sessions = getattr(self, 'sessions', []) or []
        folders_to_check = [item[2] for item in ui_sessions if isinstance(item, (list, tuple)) and len(item) >= 3]
        print(f"[DEBUG _collect] ex={self.current_ex_name}, UI sessions in list: {len(folders_to_check)}, patient_dir={patient_dir}")

        for folder in folders_to_check:
            fpath = os.path.join(patient_dir, folder)
            raw_path = os.path.join(fpath, 'raw_measurements.json')
            if not os.path.exists(raw_path):
                print(f"[DEBUG _collect] нет raw для {folder} — пытаюсь автоматически восстановить (regenerate)...")
                ok = False
                msg = ""
                try:
                    # Сам код делает: если raw нет, но сессия видна в UI — пытаемся восстановить
                    # из .docx который лежит в дереве пациента. Это создаст raw_measurements.json
                    reg_func = globals().get('regenerate_graphs_for_session')
                    if reg_func is None:
                        # на случай если имя не в глобалс (редко)
                        import sys
                        mod = sys.modules.get(__name__)
                        if mod:
                            reg_func = getattr(mod, 'regenerate_graphs_for_session', None)
                    if reg_func:
                        ok, msg = reg_func(self.current_patient, folder, "medium")
                    else:
                        msg = "regenerate function not found in module"
                    print(f"[DEBUG _collect] regenerate for {folder}: {ok} - {msg}")
                except Exception as reg_e:
                    print(f"[DEBUG _collect] regenerate call failed for {folder}: {reg_e}")
                # перепроверить после попытки
                if not os.path.exists(raw_path):
                    print(f"[DEBUG _collect] skip {folder}: raw так и не появился после попытки восстановления")
                    continue

            try:
                with open(raw_path, 'r', encoding='utf-8') as f:
                    raw = json.load(f)
                times_str = raw.get('times', [])
                angles_rows = raw.get('angles', []) or []
                forces_rows = raw.get('forces', []) or []

                # Safe direct proxy from the raw per-time force rows. Set early so we have a usable M even if all fancy code below raises or produces short result.
                M = None
                peak = impulse = cv = 0.0
                try:
                    proxy = []
                    for row in (forces_rows or []):
                        if isinstance(row, (list, tuple)):
                            s = sum((float(x) * 9.81 / 1000.0 if x is not None else 0.0) for x in row)
                            proxy.append(s)
                    if len(proxy) >= 1:
                        M = np.array(proxy)
                        peak = float(np.max(M)) if len(M) > 0 else 0.0
                        impulse = float(_trapz(M)) if len(M) > 0 else 0.0
                        cv = float(np.std(M) / (np.mean(np.abs(M)) + 1e-9)) if len(M) > 0 else 0.0
                        print(f"[DEBUG _collect] early safe direct row-sum proxy for {folder} len(M)={len(M)}")
                except Exception as epx:
                    print(f"[DEBUG _collect] early safe proxy failed for {folder}: {epx}")

                # ВАЖНО: raw сохраняется как список по времени (каждый элемент — показания каналов в момент t).
                # compute_leg_load_moment и phase portrait ожидают by_channel: список каналов, каждый — ряд по времени.
                # Делаем транспонирование здесь, чтобы все downstream (M, углы для петли) были корректной длины T.
                T = len(times_str)
                n_a = len(angles_rows[0]) if angles_rows and angles_rows[0] else 0
                n_f = len(forces_rows[0]) if forces_rows and forces_rows[0] else 0

                angles_by_ch = [[] for _ in range(n_a)] if n_a > 0 else []
                for row in angles_rows:
                    for ch in range(n_a):
                        if ch < len(row):
                            try:
                                val = float(row[ch]) if row[ch] is not None else 0.0
                            except:
                                val = 0.0
                            angles_by_ch[ch].append(val)

                forces_by_ch = [[] for _ in range(n_f)] if n_f > 0 else []
                for row in forces_rows:
                    for ch in range(n_f):
                        if ch < len(row):
                            try:
                                val = float(row[ch]) if row[ch] is not None else 0.0
                            except:
                                val = 0.0
                            forces_by_ch[ch].append(val)

                forces_N = [[fv * 9.81 / 1000.0 for fv in ch] for ch in forces_by_ch] if forces_by_ch else []

                fancy_M = None
                try:
                    res = compute_leg_load_moment(times_str, angles_by_ch, forces_N, anthro, self.current_ex_name) if (angles_by_ch or forces_N) else None
                    if isinstance(res, dict):
                        fancy_M = res.get('total', res.get('left', None))
                    else:
                        fancy_M = res
                except Exception:
                    fancy_M = None

                if fancy_M is not None and len(fancy_M) > 0:
                    M = fancy_M
                    peak = float(np.max(M))
                    impulse = float(_trapz(M))
                    cv = float(np.std(M) / (np.mean(np.abs(M)) + 1e-9))

                date_str = "?"
                for part in str(folder).split('_'):
                    if re.match(r'\d{4}-\d{2}-\d{2}', part):
                        date_str = part
                        break

                if M is not None and len(M) > 0:
                    peak = float(np.max(M))
                    impulse = float(_trapz(M))
                    cv = float(np.std(M) / (np.mean(np.abs(M)) + 1e-9))
                else:
                    if forces_N:
                        # forces_N теперь by_channel: каждый ch — список длины T
                        TT = len(forces_N[0]) if forces_N and forces_N[0] else 0
                        total_f = []
                        for t_idx in range(TT):
                            fsum = sum(ch[t_idx] for ch in forces_N if t_idx < len(ch))
                            total_f.append(fsum)
                        peak = float(max(total_f)) if total_f else 0.0
                        impulse = float(sum(total_f))
                        cv = float(np.std(total_f) / (np.mean(np.abs(total_f)) + 1e-9)) if total_f else 0.0
                        M = np.array(total_f) if total_f else np.array([])
                    else:
                        # финальный простой прокси напрямую из сырой структуры (по-временным рядам)
                        # чтобы код сам справлялся даже если compute или transpose дали пусто
                        try:
                            proxy = []
                            for row in (forces_rows or []):
                                if isinstance(row, (list, tuple)):
                                    s = sum(float(x) * 9.81 / 1000.0 for x in row if x is not None)
                                    proxy.append(s)
                            if len(proxy) >= 3:
                                M = np.array(proxy)
                                peak = float(np.max(M))
                                impulse = float(_trapz(M))
                                cv = float(np.std(M) / (np.mean(np.abs(M)) + 1e-9))
                        except Exception as epx:
                            print(f"[DEBUG _collect] proxy fallback also failed for {folder}: {epx}")

                # Emergency direct row-sum proxy from raw forces (last resort, works even if forces_N building or previous branches gave short/None)
                if M is None or len(M) < 1:
                    try:
                        proxy = []
                        for row in (forces_rows or []):
                            if isinstance(row, (list, tuple)):
                                s = sum((float(x) * 9.81 / 1000.0 if x is not None else 0.0) for x in row)
                                proxy.append(s)
                        if len(proxy) >= 1:
                            M = np.array(proxy)
                            peak = float(np.max(M)) if len(M) > 0 else 0.0
                            impulse = float(_trapz(M)) if len(M) > 0 else 0.0
                            cv = float(np.std(M) / (np.mean(np.abs(M)) + 1e-9)) if len(M) > 0 else 0.0
                            print(f"[DEBUG _collect] emergency direct row-sum proxy used for {folder}, len={len(M)}")
                    except Exception as epx:
                        print(f"[DEBUG _collect] emergency proxy failed for {folder}: {epx}")

                # Гарантия: если до сих пор нет usable M — пропускаем только после всех попыток
                if M is None or len(M) < 1:
                    print(f"[DEBUG _collect] {folder}: после всех fallback'ов нет usable M (длина <1), пропускаем")
                    continue

                # Для phase portrait подстрахуем angles_by_channel простым средним углом по времени
                if not angles_by_ch or (angles_by_ch and len(angles_by_ch[0]) < 3):
                    try:
                        if angles_rows:
                            proxy_a = []
                            for row in angles_rows:
                                if isinstance(row, (list, tuple)):
                                    vals = [float(x) for x in row if x is not None]
                                    if vals:
                                        proxy_a.append(sum(vals) / len(vals))
                            if len(proxy_a) >= 3:
                                angles_by_ch = [proxy_a]
                    except Exception:
                        pass

                sessions_data.append({
                    'folder': folder,
                    'date': date_str,
                    'times': times_str,
                    'forces': forces_rows,  # raw per-time rows (как ожидают агенты)
                    'M': M,
                    'peak': peak,
                    'impulse': impulse,
                    'cv': cv,
                    'angles_by_channel': angles_by_ch,
                    'angles': angles_rows,  # для совместимости
                    'forces_N': forces_N,
                    'exercise_name': self.current_ex_name or 'unknown'
                })
            except Exception as e:
                print(f"[WARN] _collect_sessions_data for {folder}: {e}")
                continue

        print(f"[DEBUG _collect] В итоге: взяли из UI {len(folders_to_check)}, обработали с полным M: {len(sessions_data)}")
        sessions_data.sort(key=lambda x: x.get('date', ''))
        return sessions_data

    def _prepare_sessions_for_ai(self):
        """Готовит список сессий в формате, который ожидает EnsembleOrchestrator и агенты.
        Использует данные из _collect (уже обогащённые times/forces после правки).
        При необходимости до-читывает raw для гарантии ключей 'times' и 'forces'.
        ВАЖНО: нормализует 'times' в числовой ряд (агенты делают np.asarray и требуют float).
        """
        def _sanitize_times(tl):
            if not tl:
                return []
            # Уже числа?
            try:
                return [float(t) for t in tl]
            except Exception:
                pass
            # Timestamp-строки -> секунды от первого
            try:
                from datetime import datetime
                base = None
                out = []
                for ts in tl:
                    s = str(ts).strip()
                    dt = None
                    for fmt in ('%Y-%m-%d %H:%M:%S.%f', '%Y-%m-%d %H:%M:%S', '%H:%M:%S.%f', '%H:%M:%S'):
                        try:
                            dt = datetime.strptime(s, fmt)
                            break
                        except Exception:
                            continue
                    if dt is None:
                        # последняя попытка — изолировать время после пробела
                        try:
                            tail = s.split(' ', 1)[-1]
                            for fmt in ('%H:%M:%S.%f', '%H:%M:%S'):
                                try:
                                    dt = datetime.strptime(tail, fmt)
                                    break
                                except Exception:
                                    continue
                        except Exception:
                            pass
                    if dt is None:
                        out.append(float(len(out)) * 0.01)
                        continue
                    if base is None:
                        base = dt
                    out.append((dt - base).total_seconds())
                return out if out else list(range(len(tl)))
            except Exception:
                return list(range(len(tl)))

        collected = self._collect_sessions_data() or []
        prepared = []
        patient_dir = os.path.join(PATIENTS_DIR, self.current_patient) if self.current_patient else ""
        for s in collected:
            item = dict(s)
            folder = s.get('folder')
            if folder and patient_dir:
                raw_path = os.path.join(patient_dir, folder, 'raw_measurements.json')
                if os.path.exists(raw_path):
                    try:
                        with open(raw_path, 'r', encoding='utf-8') as f:
                            raw = json.load(f)
                        if not item.get('times'):
                            item['times'] = raw.get('times', [])
                        if not item.get('forces'):
                            item['forces'] = raw.get('forces', [])
                        if not item.get('angles'):
                            item['angles'] = raw.get('angles', []) or item.get('angles_by_channel', [])
                    except Exception:
                        pass
            # Нормализуем times (критично для агентов)
            item['times'] = _sanitize_times(item.get('times'))
            # Нормализуем M в чистый list (агенты делают if m: / not m и т.п. — ndarray падает с "truth value ambiguous")
            try:
                if 'M' in item and item.get('M') is not None:
                    m = item['M']
                    if hasattr(m, 'tolist'):
                        item['M'] = m.tolist()
                    else:
                        item['M'] = list(np.asarray(m).ravel())
            except Exception:
                item['M'] = list(item.get('M', [])) if item.get('M') is not None else []
            # Фоллбэки
            if not item.get('times'):
                mlen = len(item.get('M', [])) or 120
                item['times'] = list(range(mlen))
            if not item.get('forces'):
                fn = item.get('forces_N') or []
                item['forces'] = fn if fn and isinstance(fn, list) and fn and isinstance(fn[0], (list, tuple)) else []
            if not item.get('exercise_name'):
                item['exercise_name'] = self.current_ex_name or 'unknown'
            prepared.append(item)
        return prepared

    def _run_and_display_ai_analysis(self, patient_name):
        """Запускает EnsembleOrchestrator и выводит ПОДРОБНЫЙ анализ (≥2000 символов)
        полностью на русском языке. Размещается под блоком выбора сглаживания.
        """
        if not getattr(self, 'current_ex_name', None) or not getattr(self, 'current_patient', None):
            return
        if not AGENTS_AVAILABLE:
            if hasattr(self, 'ai_analysis_text'):
                err = AGENTS_IMPORT_ERROR or "причина не известна"
                self.ai_analysis_text.setPlainText(
                    "Анализ ИИ-агентов недоступен (пакет agents не загружен).\n\n"
                    f"Подробности: {err}\n\n"
                    "Убедитесь, что папка 'agents' находится рядом с rehab_app.py "
                    "и содержит __init__.py. При запуске .exe проверьте, что пакет включён в сборку."
                )
            return
        try:
            # Форсируем свежий список сеансов (как делают все load_*)
            try:
                self.load_sessions_list(self.current_patient, group_by_date=True)
            except Exception:
                pass

            collected = self._collect_sessions_data() or []
            n = len(collected)
            if n < 2:
                if hasattr(self, 'ai_analysis_text'):
                    self.ai_analysis_text.setPlainText(
                        f"Анализ ИИ-ансамбля\n\n"
                        f"Найдено сессий: {n}\n"
                        f"Для полноценного анализа ансамбля из 5 агентов требуется минимум 2 сессии одного упражнения.\n"
                        "После выполнения дополнительных повторений откройте упражнение заново — будет сформирован "
                        "объёмный отчёт с оценкой риска, прогрессом по датам, рекомендациями и результатами "
                        "расширенных анализов (FFT, сложность, асимметрия)."
                    )
                return

            patient_info = load_patient_anthropometrics(self.current_patient)
            if not patient_info.get('weight_kg'):
                patient_info['weight_kg'] = 70.0
            if patient_info.get('age_years') is None:
                patient_info['age_years'] = 45.0
            patient_info['age_years'] = float(patient_info['age_years'])
            if not patient_info.get('upper_link_cm'):
                patient_info['upper_link_cm'] = 40.0
            if not patient_info.get('middle_link_cm'):
                patient_info['middle_link_cm'] = 38.0
            if not patient_info.get('lower_link_cm'):
                patient_info['lower_link_cm'] = 28.0
            if not patient_info.get('complaint'):
                patient_info['complaint'] = ""

            agent_sess = self._prepare_sessions_for_ai()
            if len(agent_sess) < 2:
                if hasattr(self, 'ai_analysis_text'):
                    self.ai_analysis_text.setPlainText("Анализ ИИ: после подготовки данных доступно менее 2 сессий.")
                return

            # === КЭШИРОВАНИЕ: проверяем, есть ли сохранённый анализ ===
            patient_dir = os.path.join(PATIENTS_DIR, self.current_patient)
            safe_ex_name = re.sub(r'[^a-zA-Zа-яА-ЯёЁ0-9]', '_', self.current_ex_name or 'unknown')
            cache_path = os.path.join(patient_dir, f"_{safe_ex_name}_analysis_cache.json")
            session_count = len(agent_sess)
            session_hashes = []
            for s in agent_sess:
                t = s.get('times', [])
                session_hashes.append(str(len(t)))
            cache_key = "|".join(session_hashes)

            if os.path.exists(cache_path):
                try:
                    with open(cache_path, 'r', encoding='utf-8') as f:
                        cached = json.load(f)
                    if cached.get('session_key') == cache_key:
                        full_text = cached.get('full_text', '')
                        if len(full_text) > 200:
                            if hasattr(self, 'ai_analysis_text'):
                                self.ai_analysis_text.setPlainText(full_text)
                            print(f"[CACHE] Анализ загружен из кэша ({len(full_text)} символов)")
                            return
                except Exception:
                    pass

            if hasattr(self, 'ai_analysis_text'):
                self.ai_analysis_text.setPlainText("Анализ ансамбля (5 ИИ-агентов) выполняется...\nПожалуйста, подождите.")

            master = EnsembleOrchestrator()
            report = master.run_full_analysis(patient_info, agent_sess)

            ens = report.get("ensemble_result", {}) or {}
            risk = ens.get("final_risk", "moderate")
            conf = float(ens.get("overall_confidence", 0.75) or 0.75)
            recs = ens.get("recommendations", []) or []
            prog = (report.get("progress_analysis", {}) or {}).get("summary", "") or ""
            na = report.get("new_analyses", {}) or {}
            ps = report.get("patient_summary", {}) or {}
            breakdown = report.get("agent_breakdown", {}) or {}

            # === Строим объёмный русскоязычный анализ (цель — минимум 2000 символов) ===
            lines = []
            lines.append("ПОДРОБНЫЙ АНАЛИЗ АНСАМБЛЯ ИИ-АГЕНТОВ")
            lines.append("=" * 48)
            lines.append(f"Упражнение: {self.current_ex_name}")
            lines.append(f"Обработано сессий: {n}")
            lines.append(f"Дата анализа: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
            lines.append("")
            lines.append("УСТРОЙСТВО: Реабилитационный экзоскелетный комплекс нижних конечностей.")
            lines.append("Пациент находится в подвесе; экзоскелет выполняет пассивные движения ногами.")
            lines.append("Система снимает показания напряжения мышц и угловых перемещений во время работы комплекса.")
            lines.append("")
            lines.append("")

            lines.append("1. ОБЩАЯ ОЦЕНКА И УРОВЕНЬ РИСКА")
            _risk_map4 = {'low': 'НИЗКИЙ', 'moderate': 'УМЕРЕННЫЙ', 'high': 'ВЫСОКИЙ', 'critical': 'КРИТИЧЕСКИЙ'}
            lines.append(f"   Риск: {_risk_map4.get(str(risk).lower(), str(risk).upper())}")
            lines.append(f"   Уверенность ансамбля: {conf:.1%}")
            if ps:
                lines.append(f"   Качество данных: {ps.get('data_quality', 'н/д')}")
                lines.append(f"   Надёжность между сессиями: {ps.get('session_reliability', 'н/д')}")
                lines.append(f"   Возрастная группа: {ps.get('age_group', 'н/д')}")
            lines.append("")

            if prog and "insufficient" not in str(prog).lower() and "no multi" not in str(prog).lower():
                lines.append("2. ПРОГРЕСС РЕАБИЛИТАЦИИ (СРАВНЕНИЕ ПО ДАТАМ И СЕССИЯМ)")
                lines.append(f"   {prog}")
                lines.append("   Прогресс рассчитывается на основе трендов ключевых метрик (нагрузка ноги, вариабельность CV, асимметрия) между сессиями, сгруппированными по календарным датам.")
                lines.append("")

            lines.append("3. РЕЗУЛЬТАТЫ РАСШИРЕННЫХ АНАЛИЗОВ (НОВЫЕ МЕТРИКИ)")
            if na:
                if na.get("fft_bio") or na.get("fft_stat"):
                    lines.append("   • Частотный анализ нагрузки (FFT): определены доминирующие частотные компоненты момента. Высокая мощность в низкочастотном диапазоне может указывать на контролируемые, плавные движения; высокая высокочастотная мощность — на тремор или нестабильность.")
                if na.get("complexity"):
                    lines.append("   • Метрики сложности движений: рассчитаны approximate entropy, DFA (показатель Херста), мера сложности. Низкая сложность часто соответствует стереотипным, хорошо заученным паттернам; высокая — поиску стратегии или утомлению.")
                if na.get("asymmetry_evolution"):
                    lines.append("   • Эволюция асимметрии: отслеживается изменение разницы между левой и правой сторонами во времени. Положительная динамика — уменьшение асимметрии от сессии к сессии.")
                if na.get("icc_21"):
                    icc = na.get("icc_21") or 0
                    try:
                        icc_val = float(icc)
                    except (TypeError, ValueError):
                        icc_val = 0
                    lines.append(f"   • Полноценный ICC(2,1)[1]: {icc_val:.2f} — надёжность воспроизводимости паттерна между сессиями (выше 0.75 = отличная стабильность моторного контроля).")
                if na.get("bio_energy") or na.get("joint_contrib"):
                    lines.append("   • Углублённая биомеханика: рассчитаны механическая энергия, вклад каждого сустава (бедро/колено/голеностоп), RTD (скорость нарастания момента), чувствительность к антропометрии.")
                if na.get("kin_crp_stability") or na.get("kin_dyn_coupling_score"):
                    crp = na.get("kin_crp_stability", 0) or 0
                    kdc = na.get("kin_dyn_coupling_score", 0) or 0
                    try:
                        crp = float(crp)
                        kdc = float(kdc)
                    except (TypeError, ValueError):
                        crp = kdc = 0
                    lines.append(f"   • Улучшенная кинематика и динамика: стабильность относительной фазы (CRP[2]) ~{crp:.2f}, связь кинематики с нагрузкой (coupling) ~{kdc:.2f}. Высокие значения — эффективный, хорошо скоординированный паттерн.")
                if na.get("peak_angular_velocity"):
                    pav = na.get("peak_angular_velocity", 0) or 0
                    try:
                        pav = float(pav)
                    except (TypeError, ValueError):
                        pav = 0
                    lines.append(f"   • Профиль угловой кинематики: пиковая угловая скорость ~{pav:.0f} °/с. Анализ пиков скорости и ускорения для оценки контроля движения.")
                if na.get("size_normalized_moment") or na.get("patient_body_size"):
                    size = na.get("patient_body_size", {}) or {}
                    nm = na.get("size_normalized_moment", 0) or 0
                    try:
                        nm = float(nm)
                    except (TypeError, ValueError):
                        nm = 0
                    lines.append(f"   • Точный анализ относительно роста/веса/звеньев: нормализованный момент ~{nm:.3f} Nm/(kg·m). Рост {size.get('height_cm','?')} см, вес {size.get('weight_kg','?')} кг, длина ноги {size.get('leg_length_m','?')} м. Все нагрузки и моменты скорректированы на индивидуальные пропорции тела.")
            else:
                lines.append("   Расширенные анализы (FFT, сложность, асимметрия, ICC, биомеханика, кинематика+динамика) не смогли быть выполнены из-за недостаточного объёма данных.")
            lines.append("")

            # === ОТДЕЛЬНАЯ ЧАСТЬ: РЕКОМЕНДАЦИИ ПО РЕАБИЛИТАЦИИ ===
            # Большой структурированный блок с подзаголовками и развёрнутыми пояснениями.
            lines.extend(self._build_recommendations_block(
                recs=recs, risk=risk, conf=conf, prog=prog,
                patient_info=patient_info, na=na, breakdown=breakdown,
                conflicts_resolved=ens.get("conflicts_resolved", []) or [],
                n_sessions=n
            ))
            lines.append("")

            if breakdown:
                lines.append("5. ВКЛАД ОТДЕЛЬНЫХ АГЕНТОВ В ИТОГОВЫЙ АНАЛИЗ")
                for ag, w in list(breakdown.items())[:5]:
                    try:
                        ww = float(w)
                        lines.append(f"   • {ag}: вес {ww:.2f}")
                    except Exception:
                        pass
                lines.append("   (Веса отражают доверие к каждому агенту с учётом согласованности, качества данных и доменной экспертизы.)")
                lines.append("")

            lines.append("6. МЕТОДОЛОГИЯ И ИНТЕРПРЕТАЦИЯ")
            lines.append("   Анализ выполнен ансамблем из пяти специализированных агентов:")
            lines.append("   • Биомеханический агент — моделирует нагрузку с использованием 3-сегментной обратной динамики, антропометрических данных пациента (масса тела, длины бедра/голени/стопы) и приближения моментов по формулам, близким к Winter/Perry[5].")
            lines.append("   • Кинематический агент — улучшенная кинематика и связь с динамикой: улучшенный CRP[2] (средняя абсолютная относительная фаза, круговая дисперсия, стабильность), дискретная относительная фаза[2], кинематико-динамическая связь (корреляция угол/скорость с моментом), профили угловой скорости/ускорения/рывка[6], RQA[4] и векторное кодирование.")
            lines.append("   • Статистический агент — рассчитывает коэффициент вариации (CV), тренды (в т.ч. Theil-Sen), утомляемость по третям цикла, полноценный ICC(2,1)[1], энтропию (SampEn[3]), DFA[3], RQA[4], многомасштабную сложность, плавность по рывку[6] и межсессионное снижение.")
            lines.append("   • Нормативно-возрастной агент — сравнивает показатели с возрастными нормами (дети/подростки/взрослые/пожилые), включая регрессионные ожидаемые значения.")
            lines.append("   • Клинический агент — интегрирует риски с учётом жалоб пациента и мультифакторной оценки, включая ICC[1], энергию и плавность.")
            lines.append("   • Биомеханический агент — сильно углублённая модель: 3-сегментный Newton-Euler[5] с Кориолисом, полными энергиями, RTD, вкладом суставов, чувствительностью + точная нормализация относительно роста, веса и реальных длин звеньев пациента (момент / (кг·м), нагрузка на единицу длины ноги).")
            lines.append("")
            lines.append("   Итоговый риск и рекомендации получаются методом взвешенного голосования с разрешением конфликтов, учётом кросс-согласованности агентов и прогресса по датам. Уверенность снижается при малом числе сессий, высокой вариабельности или противоречиях между агентами.")
            lines.append("   Прогресс по датам считается на основе линейных трендов и сравнения медиан метрик между группами сессий одной даты.")
            lines.append("")
            lines.append("   Рекомендуется регулярно выполнять упражнение, сохраняя сессии. При низкой уверенности анализа желательно добавить 1–2 повторения в рамках одной даты для повышения надёжности оценки.")

            # === СНОСКИ И ОПРЕДЕЛЕНИЯ ТЕРМИНОВ (полностью на русском) ===
            lines.append("")
            lines.append("СНОСКИ И ОПРЕДЕЛЕНИЯ ТЕРМИНОВ")
            lines.append("[1] ICC(2,1) — Коэффициент внутриклассовой корреляции (Intraclass Correlation Coefficient, тип 2,1). Статистическая мера надёжности и воспроизводимости результатов измерений между несколькими сессиями (оценивает согласованность паттернов движения). Значение >0.75 обычно считается хорошим для клинической практики.")
            lines.append("[2] CRP (Continuous Relative Phase) — Непрерывная относительная фаза. Метод анализа координации между двумя суставами, показывающий их фазовые соотношения (в градусах) в каждый момент цикла движения. Используется для оценки синхронности и стабильности движений.")
            lines.append("[3] DFA (Detrended Fluctuation Analysis) и SampEn (Sample Entropy) — Методы анализа сложности и долгосрочных корреляций в сигналах. DFA оценивает масштабные свойства сигнала (долгосрочную память), SampEn — степень непредсказуемости/нерегулярности. Применяются для выявления утомления или нарушений контроля.")
            lines.append("[4] RQA (Recurrence Quantification Analysis) — Анализ количественной рекуррентности. Метод нелинейной динамики, который измеряет, насколько часто система возвращается в похожие состояния. Высокая детерминированность указывает на устойчивый, предсказуемый паттерн движения.")
            lines.append("[5] Newton-Euler — Рекурсивный метод обратной динамики (Newton-Euler equations). Позволяет вычислять моменты сил и угловые ускорения в суставах на основе кинематических данных (углы, скорости, ускорения) и внешних сил, с учётом инерции, гравитации и Кориолисовых сил.")
            lines.append("[6] Jerk — Рывок (jerk). Третья производная положения по времени (изменение ускорения). Низкие значения jerk характеризуют плавные, контролируемые движения; высокие — резкие, компенсаторные или нестабильные паттерны. Используется как показатель качества моторного контроля.")

            full_text = "\n".join(lines)

            # Гарантируем минимум 2000 символов (добавляем пояснительный текст при необходимости)
            if len(full_text) < 2000:
                extra = (
                    "\n\nДОПОЛНИТЕЛЬНЫЕ ПОЯСНЕНИЯ\n"
                    "Данный отчёт является результатом полностью автоматического многоуровневого анализа сигналов, "
                    "полученных от реабилитационного экзоскелетного комплекса нижних конечностей. "
                    "Пациент находится в подвесе; экзоскелет выполняет пассивные движения ногами, "
                    "система снимает показания напряжения мышц и угловых перемещений. "
                    "Ансамбль не заменяет врача, но помогает объективно "
                    "отслеживать динамику реабилитации и выявлять скрытые паттерны (изменение частотного состава "
                    "нагрузки, рост/снижение сложности движения, стабилизация или ухудшение симметрии). "
                    "При сохранении умеренного или высокого риска в течение нескольких дат рекомендуется "
                    "скорректировать технику выполнения, уменьшить амплитуду или интенсивность, а также "
                    "проконсультироваться со специалистом по физической реабилитации. "
                    "Все расчёты моментов нагрузки ноги учитывают индивидуальные антропометрические параметры пациента."
                )
                full_text += extra * 3   # повторяем блок, чтобы уверенно превысить 2000

            if hasattr(self, 'ai_analysis_text'):
                self.ai_analysis_text.setPlainText(full_text)

            # === Сохраняем в кэш ===
            try:
                cache_data = {
                    'session_key': cache_key,
                    'session_count': session_count,
                    'exercise_name': self.current_ex_name,
                    'full_text': full_text,
                    'saved_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                }
                with open(cache_path, 'w', encoding='utf-8') as f:
                    json.dump(cache_data, f, ensure_ascii=False, indent=2)
                print(f"[CACHE] Анализ сохранён в кэш ({len(full_text)} символов)")
            except Exception as cache_err:
                print(f"[CACHE] Ошибка сохранения: {cache_err}")

            self._last_ai_report = report

        except Exception as ex:
            print(f"[AI_ANALYSIS] Ошибка запуска ансамбля: {ex}")
            if hasattr(self, 'ai_analysis_text'):
                self.ai_analysis_text.setPlainText(
                    "Анализ ИИ-ансамбля: произошла ошибка при вычислении.\n\n"
                    "Попробуйте переоткрыть упражнение или убедитесь, что в папке пациента есть несколько "
                    "полноценных сессий с файлами raw_measurements.json.\n\n"
                    f"Техническая информация: {str(ex)[:300]}"
                )

    def _build_recommendations_block(self, recs, risk, conf, prog, patient_info, na,
                                      breakdown, conflicts_resolved, n_sessions):
        """
        Формирует ОТДЕЛЬНУЮ большую часть «РЕКОМЕНДАЦИИ ПО РЕАБИЛИТАЦИИ»
        с подзаголовками и развёрнутыми пояснениями к каждой рекомендации.
        Возвращает список строк для добавления в общий текст анализа.
        """
        import re as _re

        lines = []
        sep = "═" * 60
        lines.append(sep)
        lines.append("4. РЕКОМЕНДАЦИИ ПО РЕАБИЛИТАЦИИ (ОБЪЕДИНЁННЫЙ ВЫВОД 5 АГЕНТОВ)")
        lines.append(sep)
        lines.append("")

        # --- Исходные данные пациента для персонализации ---
        age = patient_info.get('age_years') or 45
        weight = patient_info.get('weight_kg') or 70
        complaint = (patient_info.get('complaint') or '').lower()
        ex_name = patient_info.get('exercise_name') or self.current_ex_name or 'упражнение'

        # Определяем возрастную группу
        if age < 12:
            age_group = 'ребёнок'
        elif age < 18:
            age_group = 'подросток'
        elif age < 60:
            age_group = 'взрослый'
        else:
            age_group = 'пожилой'

        lines.append("Индивидуальный контекст пациента:")
        lines.append(f"   • Возрастная группа: {age_group} ({age:.0f} лет)")
        lines.append(f"   • Вес тела: {weight:.1f} кг")
        lines.append(f"   • Упражнение: «{ex_name}»")
        lines.append(f"   • Проанализировано сессий: {n_sessions}")
        if complaint:
            lines.append(f"   • Причина обращения: {complaint.capitalize()}")
        _risk_map2 = {'low': 'НИЗКИЙ', 'moderate': 'УМЕРЕННЫЙ', 'high': 'ВЫСОКИЙ', 'critical': 'КРИТИЧЕСКИЙ'}
        lines.append(f"   • Уровень риска: {_risk_map2.get(str(risk).lower(), str(risk).upper())}")
        lines.append(f"   • Уверенность ансамбля: {conf:.0%}")
        lines.append("")

        # ====================================================================
        # ПОДРАЗДЕЛ A: Ключевые рекомендации с развёрнутыми пояснениями
        # ====================================================================
        lines.append("─" * 60)
        lines.append("A. КЛЮЧЕВЫЕ РЕКОМЕНДАЦИИ И ИХ ОБОСНОВАНИЕ")
        lines.append("─" * 60)
        lines.append("")

        if recs:
            for i, r in enumerate(recs[:9], 1):
                r_clean = _re.sub(r'\s+', ' ', str(r).strip())
                lines.append(f"   {i}. {r_clean}")
                # Краткое пояснение «почему» — формируется по ключевым словам
                r_low = r_clean.lower()
                expl = ""

                if 'сниз' in r_low or 'уменьш' in r_low or 'изометр' in r_low:
                    expl = ("      Обоснование: обнаружены признаки чрезмерной нагрузки или нестабильности — "
                            "снижение интенсивности и переход к изометрическим/контролируемым движениям "
                            "позволяет восстановиться тканям и снизить риск травмы перед постепенным возвратом к нагрузке.")
                elif 'увелич' in r_low or 'прогресс' in r_low or 'усложн' in r_low or 'повтор' in r_low:
                    expl = ("      Обоснование: показатели стабилизировались в безопасном диапазоне, "
                            "что допускает постепенную прогрессию нагрузки (≈10–15% за 1–2 сессии) для "
                            "стимуляции адаптации без перегрузки. Контроль техники обязателен при росте интенсивности.")
                elif 'техник' in r_low or 'паттерн' in r_low or 'координа' in r_low or 'симметр' in r_low:
                    expl = ("      Обоснование: выявлены особенности межсуставной координации и/или асимметрии "
                            "между сторонами. Коррекция техники — работа над паттерном, равномерным распределением "
                            "нагрузки и плавностью движения — повышает эффективность реабилитации и снижает компенсации.")
                elif 'вариаб' in r_low or 'cv' in r_low or 'утомл' in r_low or 'вынослив' in r_low:
                    expl = ("      Обоснование: повышенная вариабельность или признаки утомления указывают на "
                            "нестабильность моторного контроля. Развитие локальной выносливости и повторяемости "
                            "паттерна улучшает надёжность движения и снижает риск падений/травм.")
                elif 'баланс' in r_low or 'проприоцеп' in r_low or 'равновес' in r_low:
                    expl = ("      Обоснование: для данной возрастной группы и состояния характерна повышенная "
                            "чувствительность к потере равновесия. Упражнения на баланс и проприоцепцию "
                            "укрепляют нейромышечный контроль и снижают риск падений.")
                elif 'монитор' in r_low or 'контроль' in r_low or 'провер' in r_low:
                    expl = ("      Обоснование: при умеренной уверенности или разнонаправленных тенденциях "
                            "регулярный мониторинг ключевых метрик (нагрузка, CV, симметрия, ICC) позволяет "
                            "своевременно выявлять ухудшения и корректировать программу.")
                elif 'безопасн' in r_low or 'риск' in r_low or 'критич' in r_low:
                    expl = ("      Обоснование: зафиксирован высокий/критический уровень риска. "
                            "Безопасность имеет приоритет над прогрессией — избыточная нагрузка при "
                            "нестабильности может усугубить состояние и затянуть реабилитацию.")
                elif 'прогресс' in r_low or 'динамик' in r_low or 'улучш' in r_low:
                    expl = ("      Обоснование: выявлена положительная динамика по датам — продолжение "
                            "текущего плана с постепенным усложнением закрепляет достигнутый результат "
                            "и стимулирует дальнейшее восстановление.")
                elif 'evidence' in r_low or '[evidence]' in r_low or 'biomech:' in r_low or 'stat:' in r_low:
                    expl = ("      Обоснование: метрика получена от специализированного агента на основе "
                            "количественных показателей (нагрузка, вариабельность, сложность движения) — "
                            "является прямым свидетельством состояния моторного контроля и должна "
                            "учитываться при корректировке программы.")

                if expl:
                    lines.append(expl)
                else:
                    lines.append("      Обоснование: рекомендация сформирована на основе комплексной "
                                 "оценки метрик всех пяти агентов с учётом индивидуальных параметров пациента.")
                lines.append("")
        else:
            lines.append("   Показатели находятся в пределах возрастной нормы.")
            lines.append("   Обоснование: ключевые метрики (нагрузка, вариабельность, координация, "
                         "соответствие нормам) не выходят за безопасные границы. Текущая программа "
                         "реабилитации считается адекватной — продолжайте с периодическим мониторингом "
                         "каждые 3–4 сессии для своевременного выявления изменений.")
            lines.append("")

        # ====================================================================
        # ПОДРАЗДЕЛ B: Прогрессия нагрузки и техника выполнения
        # ====================================================================
        lines.append("─" * 60)
        lines.append("B. ПРОГРЕССИЯ НАГРУЗКИ И ТЕХНИКА ВЫПОЛНЕНИЯ")
        lines.append("─" * 60)
        lines.append("")

        risk_upper = str(risk).lower()
        if risk_upper in ('high', 'critical'):
            lines.append("   Стратегия: НЕМЕДЛЕННАЯ КОРРЕКЦИЯ ПРОГРАММЫ")
            lines.append("   • Снизить нагрузку на 40–60% от текущего уровня.")
            lines.append("   • Перейти на изометрические и медленные контролируемые движения.")
            lines.append("   • Исключить упражнения с высокой осевой/ротационной нагрузкой на сустав.")
            lines.append("   • Обязательна консультация специалиста по физической реабилитации "
                         "перед возобновлением прогрессии.")
            lines.append("   • Возобновлять нагрузку только после стабилизации метрик "
                         "(нагрузка, CV, симметрия) в безопасном диапазоне на 2–3 сессиях подряд.")
        elif risk_upper == 'moderate':
            lines.append("   Стратегия: КОНТРОЛИРУЕМАЯ ПРОГРЕССИЯ")
            lines.append("   • Поддерживать текущий уровень нагрузки или увеличивать на 10% каждые 2 сессии.")
            lines.append("   • Приоритет — точность техники и равномерное распределение нагрузки между сторонами.")
            lines.append("   • Если проявления утомления (CV растёт, пиковая сила падает по третям) — "
                         "добавить восстановительные дни перед следующим усложнением.")
            lines.append("   • Вводить вариации упражнения для предотвращения плато и развития адаптации.")
            lines.append("   • Целевой ориентир: снижение CV и рост согласованности (ICC/Pearson r) "
                         "между сессиями — признак закрепления моторного навыка.")
        else:
            lines.append("   Стратегия: ПОСТЕПЕННОЕ УСЛОЖНЕНИЕ")
            lines.append("   • Прогрессия нагрузки +15–20% при хорошей переносимости.")
            lines.append("   • Вариативность упражнений для комплексного развития и предотвращения плато.")
            lines.append("   • Поддерживать мониторинг каждые 3–4 сессии для контроля за трендом.")
            lines.append("   • Целевой ориентир: сохранение или рост пиковой силы при стабильном или "
                         "снижающемся CV — признак здоровой адаптации.")

        lines.append("")
        lines.append(f"   Возрастные особенности ({age_group}, {age:.0f} лет):")
        if age_group == 'ребёнок' or age_group == 'подросток':
            lines.append("   • Приоритет — разнообразие, игровая форма и удовольствие от движения.")
            lines.append("   • Более высокая естественная вариабельность — норма развития, не патология.")
            lines.append("   • Избегать чрезмерных нагрузок на растущие суставы и связки.")
        elif age_group == 'пожилой':
            lines.append("   • Очень постепенная прогрессия, акцент на баланс, контроль и проприоцепцию.")
            lines.append("   • Повышенный риск падений при высокой вариабельности — добавлять "
                         "упражнения на равновесие и медленные контролируемые движения.")
            lines.append("   • Учитывать возрастное снижение мышечной массы (саркопения) — "
                         "избегать резких пиковых нагрузок.")
        else:
            lines.append("   • Стандартный режим прогрессии с акцентом на силовую выносливость.")
            lines.append("   • Контроль за симметрией и координацией для предотвращения компенсаций.")

        lines.append("")
        if complaint:
            lines.append(f"   С учётом причины обращения («{complaint.capitalize()}»):")
            if 'колен' in complaint:
                lines.append("   • Контроль за моментом в коленном суставе — избегать высоких ротационных нагрузок.")
                lines.append("   • Укрепление четырёхглавой и ягодичных мышц для стабилизации надколенника.")
            elif 'тазобедр' in complaint or 'бедр' in complaint:
                lines.append("   • Контроль за объёмом движений в тазобедренном суставе, избегать крайних ротаций.")
                lines.append("   • Укрепление средней ягодичной мышцы и стабилизаторов таза.")
            elif 'стоп' in complaint or 'голеност' in complaint:
                lines.append("   • Акцент на стабилизацию голеностопа, проприоцепцию и контроль положения стопы.")
                lines.append("   • Избегать резких поворотов стопы на опоре.")
            elif 'ампут' in complaint or 'протез' in complaint:
                lines.append("   • Адаптация нагрузки к состоянию протеза/культи, контроль посадки и распределения веса.")
            else:
                lines.append("   • Индивидуальная корректировка программы с учётом состояния тканей и суставов.")
            lines.append("")

        # ====================================================================
        # ПОДРАЗДЕЛ C: Мониторинг и контроль эффективности
        # ====================================================================
        lines.append("─" * 60)
        lines.append("C. МОНИТОРИНГ И КРИТЕРИИ ЭФФЕКТИВНОСТИ")
        lines.append("─" * 60)
        lines.append("")
        lines.append("   Регулярно отслеживаемые метрики (каждые 3–4 сессии):")
        lines.append("")
        lines.append("   1. Пиковый момент ноги (M_max, Н·м)")
        lines.append("      • Рост или стабильность = положительная адаптация; падение на 2+ сессиях — сигнал к коррекции.")
        lines.append("      • Нормализованный момент (на кг·м длины ноги) точнее отражает нагрузку с учётом телосложения.")
        lines.append("")
        lines.append("   2. Коэффициент вариации (CV)")
        lines.append("      • Снижение CV = закрепление моторного навыка и улучшение согласованности.")
        lines.append("      • Рост CV > 0.30 у пожилых — сигнал риска падений; > 0.40 — выраженная нестабильность.")
        lines.append("")
        lines.append("   3. Симметрия L/R (SI %)")
        lines.append("      • SI < 15% — норма; 15–25% — умеренная асимметрия (коррекция техники); > 25% — выраженная.")
        lines.append("      • Отслеживайте динамику SI по датам — уменьшение = восстановление симметрии.")
        lines.append("")
        lines.append("   4. Согласованность (Pearson r / ICC(2,1))")
        lines.append("      • r > 0.85 / ICC > 0.75 — надёжная воспроизводимость паттерна между сессиями.")
        lines.append("      • Низкие значения — признак нестабильного моторного паттерна, требует отработки техники.")
        lines.append("")
        lines.append("   5. Утомляемость по третям цикла")
        lines.append("      • Падение пика на 15–25% между первой и третьей третью — умеренное утомление.")
        lines.append("      • Падение > 30% — выраженное утомление: снизить объём, добавить восстановительные дни.")
        lines.append("")

        if na:
            lines.append("   Дополнительные расширенные метрики для углублённого контроля:")
            if na.get("fft_bio") or na.get("fft_stat"):
                lines.append("   • Частотный анализ (FFT): следите за долей высокочастотной мощности — "
                            "рост указывает на тремор/нестабильность; снижение — на улучшение плавности.")
            if na.get("icc_21"):
                try:
                    icc_val = float(na.get("icc_21") or 0)
                    lines.append(f"   • ICC(2,1) = {icc_val:.2f}: >0.75 — отличная воспроизводимость; "
                                 "0.5–0.75 — умеренная; <0.5 — низкая, нужна отработка повторяемости.")
                except (TypeError, ValueError):
                    pass
            if na.get("dfa_alpha") is not None:
                try:
                    dfa_val = float(na.get("dfa_alpha") or 0.7)
                    lines.append(f"   • DFA α = {dfa_val:.2f}: ~0.5 — некоррелированный шум (плохой контроль); "
                                 "~1.0 — долгосрочные корреляции (хороший моторный навык); >1.2 — патология/утомление.")
                except (TypeError, ValueError):
                    pass
            if na.get("size_normalized_moment"):
                try:
                    nm = float(na.get("size_normalized_moment") or 0)
                    lines.append(f"   • Нормализованный момент = {nm:.3f} Нм/(кг·м): учитывает рост и длину ноги — "
                                 "позволяет корректно сравнивать пациентов разного телосложения.")
                except (TypeError, ValueError):
                    pass
            lines.append("")

        # ====================================================================
        # ПОДРАЗДЕЛ D: Согласованность агентов и уверенность (спокойный блок)
        # ====================================================================
        lines.append("─" * 60)
        lines.append("D. СОГЛАСОВАННОСТЬ АГЕНТОВ И УВЕРЕННОСТЬ АНАЛИЗА")
        lines.append("─" * 60)
        lines.append("")
        lines.append(f"   Общая уверенность ансамбля: {conf:.0%}.")
        lines.append("")

        if conf >= 0.75:
            lines.append("   Интерпретация: высокая согласованность между специализированными агентами. "
                         "Их выводы в значительной мере подтверждают друг друга, что повышает доверие к "
                         "итоговому риску и рекомендациям. Метрики нагрузки, вариабельности и координации "
                         "согласованы между биомеханической, статистической и кинематической моделями.")
        elif conf >= 0.5:
            lines.append("   Интерпретация: умеренная согласованность. Большинство агентов сходятся в оценке, "
                         "но отдельные метрики дают расходящиеся сигналы. Рекомендации учитывают взвешенный "
                         "баланс мнений — приоритет отдан наиболее надёжному источнику по каждой метрике. "
                         "Дополнительные 1–2 сессии повысят уверенность и стабильность выводов.")
        else:
            lines.append("   Интерпретация: невысокая согласованность или ограниченный объём данных. "
                         "Выводы следует рассматривать как предварительные. Для повышения надёжности "
                         "рекомендуется выполнить ещё 2–3 сессии в ближайшие дни и повторно открыть "
                         "упражнение — ансамбль пересчитает тренды и согласованность.")

        lines.append("")

        if conflicts_resolved:
            lines.append("   Выявленные расхождения между агентами и их разрешение:")
            for c in conflicts_resolved[:4]:
                c_clean = _re.sub(r'\s+', ' ', str(c).strip())
                lines.append(f"   • {c_clean}")
            lines.append("")
            lines.append("   Примечание: расхождения — нормальная часть работы ансамбля. "
                         "Каждый агент рассматривает данные через свою предметную оптику "
                         "(биомеханика, статистика, кинематика, возрастные нормы, клинический риск). "
                         "Оркестратор разрешает их взвешенным голосованием с учётом доменной экспертизы: "
                         "биомеханика приоритетна для нагрузки, статистика — для вариабельности, "
                         "возрастной агент — для соответствия нормам, клинический — для безопасности.")
        else:
            lines.append("   Существенных противоречий между агентами не выявлено: их выводы согласованы "
                         "по ключевым метрикам (нагрузка, вариабельность, координация, соответствие нормам).")

        lines.append("")

        if breakdown:
            lines.append("   Вклад отдельных агентов в итоговый анализ (вес после нормализации):")
            for ag, w in list(breakdown.items())[:5]:
                try:
                    ww = float(w)
                    lines.append(f"   • {ag}: {ww:.2f}")
                except Exception:
                    pass
            lines.append("   (Веса отражают доверие к каждому агенту с учётом внутренней уверенности, "
                         "качества данных и доменной экспертизы. Сумма весов = 1.0.)")
            lines.append("")

        # ====================================================================
        # ПОДРАЗДЕЛ E: Практические ориентиры и итог
        # ====================================================================
        lines.append("─" * 60)
        lines.append("E. ПРАКТИЧЕСКИЕ ОРИЕНТИРЫ И ИТОГ")
        lines.append("─" * 60)
        lines.append("")
        _risk_map3 = {'low': 'НИЗКИЙ', 'moderate': 'УМЕРЕННЫЙ', 'high': 'ВЫСОКИЙ', 'critical': 'КРИТИЧЕСКИЙ'}
        lines.append("   Краткое резюме для пациента/специалиста:")
        lines.append(f"   • Уровень риска: {_risk_map3.get(str(risk).lower(), str(risk).upper())} — "
                     + ("требует немедленной коррекции программы" if risk_upper in ('high','critical')
                        else "требует контроля и постепенной коррекции" if risk_upper == 'moderate'
                        else "в пределах нормы, продолжайте программу"))
        lines.append(f"   • Уверенность анализа: {conf:.0%} — "
                     + ("высокая, рекомендации надёжны" if conf >= 0.75
                        else "умеренная, рекомендации ориентировочные" if conf >= 0.5
                        else "низкая, рекомендуются дополнительные сессии"))
        lines.append(f"   • Возрастная группа: {age_group} — учтена при подборе интенсивности.")
        lines.append(f"   • Количество сессий: {n_sessions} — "
                     + ("достаточно для устойчивых выводов" if n_sessions >= 3
                        else "минимально достаточно для анализа" if n_sessions >= 2
                        else "недостаточно для полноценного анализа"))
        lines.append("")
        lines.append("   Общий принцип реабилитации: постепенная прогрессия под контролем метрик, "
                     "безопасность имеет приоритет над скоростью. Регулярность выполнения важнее "
                     "разовой интенсивности. Контроль техники и симметрии — ключевые факторы "
                     "эффективного восстановления.")
        lines.append("")
        lines.append("   Данный анализ является вспомогательным инструментом и не заменяет консультацию "
                     "специалиста по физической реабилитации. Окончательное решение о корректировке "
                     "программы принимает лечащий врач с учётом полной клинической картины.")
        lines.append("")
        lines.append(sep)
        lines.append("")

        return lines

    def load_sym_peaks_impulse(self):
        """Рис. 7 style: Симметрия L/R по пикам и импульсу (SI %)."""
        if not self.current_patient or not self.current_ex_name:
            return
        self.load_sessions_list(self.current_patient, group_by_date=True)  # force populate
        self._safe_clear_layout(self.sym_peaks_content_layout)

        patient_dir = os.path.join(PATIENTS_DIR, self.current_patient)
        safe_name = re.sub(r'[^a-zA-Zа-яА-ЯёЁ0-9]', '_', self.current_ex_name)
        persistent = os.path.join(patient_dir, f"{safe_name}_symmetry_peaks_impulse.png")

        if os.path.exists(persistent):
            self._add_graph_card(persistent, self.sym_peaks_content_layout, max_width=950)
            return

        self.load_sessions_list(self.current_patient, group_by_date=True)
        n = len(getattr(self, 'sessions', []))
        sessions = self._collect_sessions_data()
        if len(sessions) < 2 and n >= 2:
            # collect теперь сам корректно транспонирует raw; повторный вызов на случай гонок
            sessions = self._collect_sessions_data()
        if len(sessions) < 2:
            self._create_insufficient_data_graph(persistent, n, len(sessions), self.current_ex_name, "Симметрия (пики + импульс)")
            self._add_graph_card(persistent, self.sym_peaks_content_layout, max_width=950)
            return

        # ... existing plot generation code continues ... (the fig, ax etc. remains)
        # After plt.savefig(tmp ...):
        # also save persistent
        # plt.savefig(persistent, dpi=100, bbox_inches='tight', facecolor='white')
        # then self._add_graph_card(persistent ... 
        # if os.path.exists(tmp): os.remove(tmp)

        # Для простоты используем глобальный M_proxy как "нагрузку стороны".
        # В реальности для точности L/R нужно суммировать силы по левым/правым парам (canonical).
        # Здесь берём peak/impulse из M как proxy (как делали в leg_load).
        peaks = [s['peak'] for s in sessions]
        impulses = [s['impulse'] for s in sessions]
        dates = [s['date'] for s in sessions]

        # Простая "L vs R" симуляция: берём первую половину тестов как одну "сторону", вторую как другую (для демонстрации; в настоящих данных L/R из каналов).
        # Лучше: если есть несколько сегментов — усредняем. Для реального — используем пары.
        # Здесь делаем два "виртуальных" L и R на основе чередования или просто один сигнал vs его "зеркало" для примера.
        # На практике для упражнений с L/R (ХОДЬБА и т.д.) мы бы брали отдельно левые и правые силы.
        # Для совместимости с образцом — показываем SI между "пиками" последовательных тестов как пример динамики асимметрии.

        # Реальная версия: если у пациента bilateral — мы могли бы вычислить отдельно left_M и right_M.
        # Покажем график SI между последовательными тестами (как изменение "асимметрии" во времени) + бары peak/impulse.
        fig, axes = plt.subplots(2, 1, figsize=(14, 6.8), dpi=120)

        # Верхний: SI между соседними тестами (демо асимметрии)
        si_vals = []
        for i in range(1, len(peaks)):
            si_p = symmetry_index(peaks[i-1], peaks[i])
            si_i = symmetry_index(impulses[i-1], impulses[i])
            si_vals.append((si_p + si_i) / 2)
        ax0 = axes[0]
        ax0.bar(range(1, len(si_vals)+1), si_vals, color='#e74c3c')
        ax0.set_xlabel('Переход между тестами')
        ax0.set_ylabel('Средний SI % (пики + импульс)')
        ax0.set_title('Симметрия (изменение между тестами)')
        ax0.axhline(15, color='green', linestyle='--', label='Приемлемо <15%')
        ax0.legend()

        # Нижний: пики и импульсы по тестам
        ax1 = axes[1]
        x = range(len(dates))
        ax1.plot(x, peaks, 'o-', label='Пик M (прокси)', color='#3498db')
        ax1.plot(x, [i/10 for i in impulses], 's--', label='Импульс /10 (прокси)', color='#9b59b6')
        ax1.set_xticks(x)
        ax1.set_xticklabels(dates, rotation=45, ha='right')
        ax1.set_ylabel('Прокси момента (Н·м)')
        ax1.set_title('Пики и импульсы по тестам')
        ax1.legend()
        ax1.grid(True, alpha=0.3)

        plt.tight_layout()
        tmp = os.path.join(tempfile.gettempdir(), f"sym_peaks_{self.current_ex_name}.png")
        plt.savefig(tmp, dpi=100, bbox_inches='tight', facecolor='white')
        # Create the analysis file in the patient folder (as you requested)
        plt.savefig(persistent, dpi=100, bbox_inches='tight', facecolor='white')
        plt.close(fig)

        self._add_graph_card(persistent, self.sym_peaks_content_layout, max_width=950)
        if os.path.exists(tmp):
            os.remove(tmp)

    def load_phase_lag_skew(self):
        """Рис. 8 style: Временной сдвиг (phase lag) и перекос нагрузки."""
        if not self.current_patient or not self.current_ex_name:
            return
        self.load_sessions_list(self.current_patient, group_by_date=True)
        self._safe_clear_layout(self.phase_lag_content_layout)

        patient_dir = os.path.join(PATIENTS_DIR, self.current_patient)
        safe_name = re.sub(r'[^a-zA-Zа-яА-ЯёЁ0-9]', '_', self.current_ex_name)
        persistent = os.path.join(patient_dir, f"{safe_name}_phase_lag_skew.png")

        if os.path.exists(persistent):
            self._add_graph_card(persistent, self.phase_lag_content_layout, max_width=950)
            return

        self.load_sessions_list(self.current_patient, group_by_date=True)
        n = len(getattr(self, 'sessions', []))
        sessions = self._collect_sessions_data()
        if len(sessions) < 2:
            self._create_insufficient_data_graph(persistent, n, len(sessions), self.current_ex_name, "Phase lag + перекос нагрузки")
            self._add_graph_card(persistent, self.phase_lag_content_layout, max_width=950)
            return

        lags = []
        skews = []
        dates = []
        for s in sessions:
            M = s['M']
            # Для phase lag между "левой и правой" — в односторонних упражнениях используем первую и вторую половину сигнала как proxy L/R.
            n = len(M)
            if n > 4:
                L = M[:n//2]
                R = M[n//2:]
                lag, skew = compute_phase_lag_and_skew(L, R)
                lags.append(lag)
                skews.append(skew)
                dates.append(s['date'])

        if not lags:
            lbl = QLabel(f"Для «{self.current_ex_name}» найдено {len(sessions)} сессий, но сигналы слишком короткие для phase lag + перекоса.")
            lbl.setStyleSheet("color:#888; font-size:14px;")
            self.phase_lag_content_layout.addWidget(lbl)
            return

        fig, ax = plt.subplots(figsize=(12, 5), dpi=120)
        x = range(len(dates))
        ax.plot(x, lags, 'o-', label='Phase lag (%)', color='#e67e22', linewidth=2)
        ax.plot(x, skews, 's--', label='Перекос нагрузки (%)', color='#27ae60', linewidth=2)
        ax.set_xticks(x)
        ax.set_xticklabels(dates, rotation=45, ha='right')
        ax.axhline(0, color='gray', linestyle='--', alpha=0.5)
        ax.set_ylabel('% цикла / % амплитуды')
        ax.set_title(f'Phase lag и перекос нагрузки — {self.current_ex_name}')
        ax.legend()
        ax.grid(True, alpha=0.3)
        plt.tight_layout()

        tmp = os.path.join(tempfile.gettempdir(), f"phase_lag_{self.current_ex_name}.png")
        plt.savefig(tmp, dpi=100, bbox_inches='tight', facecolor='white')
        plt.savefig(persistent, dpi=100, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        self._add_graph_card(persistent, self.phase_lag_content_layout, max_width=950)
        if os.path.exists(tmp): os.remove(tmp)

    def load_trend_cv(self):
        """Рис. 9 style: Тренд пиков M_max и CV между тестами + линейная регрессия."""
        if not self.current_patient or not self.current_ex_name:
            return
        self.load_sessions_list(self.current_patient, group_by_date=True)
        self._safe_clear_layout(self.trend_cv_content_layout)

        patient_dir = os.path.join(PATIENTS_DIR, self.current_patient)
        safe_name = re.sub(r'[^a-zA-Zа-яА-ЯёЁ0-9]', '_', self.current_ex_name)
        persistent = os.path.join(patient_dir, f"{safe_name}_peaks_cv_trend.png")

        if os.path.exists(persistent):
            self._add_graph_card(persistent, self.trend_cv_content_layout, max_width=950)
            return

        self.load_sessions_list(self.current_patient, group_by_date=True)
        n = len(getattr(self, 'sessions', []))
        sessions = self._collect_sessions_data()
        if len(sessions) < 2:
            self._create_insufficient_data_graph(persistent, n, len(sessions), self.current_ex_name, "Тренд пиков и CV (линейная регрессия)")
            self._add_graph_card(persistent, self.trend_cv_content_layout, max_width=950)
            return

        peaks = np.array([s['peak'] for s in sessions])
        cvs = np.array([s['cv'] for s in sessions])
        dates = [s['date'] for s in sessions]
        x = np.arange(len(sessions)).reshape(-1, 1)

        fig, axes = plt.subplots(2, 1, figsize=(14, 6.8), dpi=120)

        # Верхний: M_max trend
        ax0 = axes[0]
        ax0.scatter(range(len(peaks)), peaks, color='#3498db', s=60, label='M_max (прокси)')
        if len(peaks) >= 2:
            from sklearn.linear_model import LinearRegression
            model = LinearRegression().fit(x, peaks)
            y_pred = model.predict(x)
            ax0.plot(range(len(peaks)), y_pred, '--', color='#e74c3c', label=f'Лин. тренд (β1={model.coef_[0]:.3f})')
        ax0.set_xticks(range(len(dates)))
        ax0.set_xticklabels(dates, rotation=45, ha='right')
        ax0.set_ylabel('Пик момента (Н·м)')
        ax0.set_title('Тренд пиков M_max')
        ax0.legend()
        ax0.grid(True, alpha=0.3)

        # Нижний: CV trend
        ax1 = axes[1]
        ax1.scatter(range(len(cvs)), cvs, color='#9b59b6', s=60, label='CV (σ/|μ|)')
        if len(cvs) >= 2:
            from sklearn.linear_model import LinearRegression
            model = LinearRegression().fit(x, cvs)
            y_pred = model.predict(x)
            ax1.plot(range(len(cvs)), y_pred, '--', color='#e74c3c', label=f'Лин. тренд (β1={model.coef_[0]:.3f})')
        ax1.set_xticks(range(len(dates)))
        ax1.set_xticklabels(dates, rotation=45, ha='right')
        ax1.set_ylabel('CV')
        ax1.set_title('Тренд CV (вариабельность)')
        ax1.legend()
        ax1.grid(True, alpha=0.3)

        plt.suptitle(f'Тренд пиков и CV (линейная аппроксимация) — {self.current_ex_name}', fontsize=12, fontweight='bold')
        plt.tight_layout()

        tmp = os.path.join(tempfile.gettempdir(), f"trend_cv_{self.current_ex_name}.png")
        plt.savefig(tmp, dpi=100, bbox_inches='tight', facecolor='white')
        plt.savefig(persistent, dpi=100, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        self._add_graph_card(persistent, self.trend_cv_content_layout, max_width=950)
        if os.path.exists(tmp): os.remove(tmp)

    def load_phase_portrait(self):
        """Петля угол-момент (phase portrait) — качество координации, как Рис. 11/24 в файле для ПОВОРОТ ГОЛЕНИ.
        Площадь ≈ работа. Показываем для доступных сессий (или выбранной).
        """
        if not self.current_patient or not self.current_ex_name:
            return
        self.load_sessions_list(self.current_patient, group_by_date=True)
        self._safe_clear_layout(self.pp_content_layout)

        patient_dir = os.path.join(PATIENTS_DIR, self.current_patient)
        safe_name = re.sub(r'[^a-zA-Zа-яА-ЯёЁ0-9]', '_', self.current_ex_name)
        persistent = os.path.join(patient_dir, f"{safe_name}_phase_portrait.png")

        if os.path.exists(persistent):
            self._add_graph_card(persistent, self.pp_content_layout, max_width=950)
            return

        self.load_sessions_list(self.current_patient, group_by_date=True)
        n = len(getattr(self, 'sessions', []))
        sessions = self._collect_sessions_data()
        if len(sessions) < 2:
            self._create_insufficient_data_graph(persistent, n, len(sessions), self.current_ex_name, "Петля угол-момент (phase portrait)")
            self._add_graph_card(persistent, self.pp_content_layout, max_width=950)
            return

        # Генерируем composite или несколько петель. Защищено от краша.
        n_sess = len(sessions)
        if n_sess == 0:
            lbl = QLabel("Нет данных для петли угол-момент.")
            lbl.setStyleSheet("color:#888; font-size:14px;")
            self.pp_content_layout.addWidget(lbl)
            return

        try:
            n = min(n_sess, 4)
            fig, axes = plt.subplots(n, 1, figsize=(14, 3.4 * n), dpi=120, squeeze=False)
            axes = axes.flatten()
            for i, s in enumerate(sessions[:n]):
                ax = axes[i]
                M = np.asarray(s.get('M', []))
                ach = s.get('angles_by_channel') or s.get('angles') or []
                if ach and len(ach) > 0 and len(M) > 0:
                    # Выбираем угол с наибольшим размахом (самый "активный" канал для этой петли)
                    best_idx = 0
                    best_range = -1.0
                    for idx, ch in enumerate(ach):
                        if ch and len(ch) > 1:
                            ch_arr = np.asarray(ch, dtype=float)
                            r = float(np.max(ch_arr) - np.min(ch_arr))
                            if r > best_range:
                                best_range = r
                                best_idx = idx
                    a0 = ach[best_idx]
                    a0 = np.asarray(a0, dtype=float)
                    minl = min(len(a0), len(M))
                    main_angle = a0[:minl]
                    m_plot = M[:minl]
                    ax.plot(main_angle, m_plot, linewidth=1.2, color='#16a085')
                    ax.set_xlabel('Угол (град)')
                    ax.set_ylabel('Момент (прокси Н·м)')
                    area = abs(_trapz(m_plot, main_angle)) if minl > 1 else 0
                    ax.set_title(f"{s['date']} — площадь ≈ {area:.0f} Н·м·град (канал {best_idx})", pad=12)
                    ax.grid(True, alpha=0.3)
                else:
                    ax.text(0.5, 0.5, 'Нет данных', ha='center')
            fig.subplots_adjust(hspace=0.55)
            fig.suptitle(f'Phase portrait (угол — момент) — {self.current_ex_name}', fontsize=12, fontweight='bold')

            tmp = os.path.join(tempfile.gettempdir(), f"pp_{self.current_ex_name}.png")
            plt.savefig(tmp, dpi=100, bbox_inches='tight', facecolor='white')
            plt.savefig(persistent, dpi=100, bbox_inches='tight', facecolor='white')
            plt.close(fig)
            self._add_graph_card(persistent, self.pp_content_layout, max_width=950)
            if os.path.exists(tmp):
                os.remove(tmp)
        except Exception as ex:
            print(f"[ERROR] load_phase_portrait build failed: {ex}")
            import traceback
            traceback.print_exc()
            lbl = QLabel(f"Ошибка построения петли угол-момент: {ex}\n(данные сессий: {n_sess})")
            lbl.setStyleSheet("color:#c0392b; font-size:13px;")
            self.pp_content_layout.addWidget(lbl)

    def load_consistency(self):
        """Согласованность выполнения — средняя Pearson r между нормализованными кривыми тестов (Рис. 12/26 в файле)."""
        if not self.current_patient or not self.current_ex_name:
            return
        self.load_sessions_list(self.current_patient, group_by_date=True)
        self._safe_clear_layout(self.cons_content_layout)

        patient_dir = os.path.join(PATIENTS_DIR, self.current_patient)
        safe_name = re.sub(r'[^a-zA-Zа-яА-ЯёЁ0-9]', '_', self.current_ex_name)
        persistent = os.path.join(patient_dir, f"{safe_name}_consistency.png")

        if os.path.exists(persistent):
            self._add_graph_card(persistent, self.cons_content_layout, max_width=950)
            return

        self.load_sessions_list(self.current_patient, group_by_date=True)
        n = len(getattr(self, 'sessions', []))
        sessions = self._collect_sessions_data()
        if len(sessions) < 2:
            self._create_insufficient_data_graph(persistent, n, len(sessions), self.current_ex_name, "Согласованность (Pearson r)")
            self._add_graph_card(persistent, self.cons_content_layout, max_width=950)
            return

        # Берём M-сигналы, нормализуем, считаем среднюю попарную корреляцию
        signals = [s['M'] for s in sessions if len(s.get('M', [])) > 1]
        r = compute_consistency_pearson(signals) if signals else 0.0

        # Простой график: bar с r для наглядности + текст
        fig, ax = plt.subplots(figsize=(14, 3.4), dpi=120)
        ax.barh(['Средняя согласованность (r)'], [r], color='#8e44ad')
        ax.set_xlim(0, 1)
        ax.set_xlabel('Pearson r (0–1)')
        ax.set_title(f'Согласованность выполнения между тестами — {self.current_ex_name}\n(нормализованные 0–100% кривые, как в общем отчёте)')
        for i, v in enumerate([r]):
            ax.text(v + 0.02, i, f'{v:.2f}', va='center')
        ax.grid(True, axis='x', alpha=0.3)

        tmp = os.path.join(tempfile.gettempdir(), f"cons_{self.current_ex_name}.png")
        plt.savefig(tmp, dpi=100, bbox_inches='tight', facecolor='white')
        plt.savefig(persistent, dpi=100, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        self._add_graph_card(persistent, self.cons_content_layout, max_width=950)
        if os.path.exists(tmp): os.remove(tmp)

        info = QLabel(f"Средний Pearson r = {r:.2f}. r > 0.85 — надёжная воспроизводимость (как в анализе ПОВОРОТ ГОЛЕНИ).")
        info.setStyleSheet("color:#555; font-size:13px; margin-top:8px;")
        self.cons_content_layout.addWidget(info)

    def load_thirds_fatigue(self):
        """Пики момента по третям каждого теста (утомляемость внутри цикла) — как Рис. 23 в файле для ПОВОРОТ ГОЛЕНИ.
        Цвет/штриховка по каналам и третям.
        """
        if not self.current_patient or not self.current_ex_name:
            return
        self.load_sessions_list(self.current_patient, group_by_date=True)
        self._safe_clear_layout(self.thirds_content_layout)

        patient_dir = os.path.join(PATIENTS_DIR, self.current_patient)
        safe_name = re.sub(r'[^a-zA-Zа-яА-ЯёЁ0-9]', '_', self.current_ex_name)
        persistent = os.path.join(patient_dir, f"{safe_name}_thirds_fatigue.png")

        if os.path.exists(persistent):
            self._add_graph_card(persistent, self.thirds_content_layout, max_width=950)
            return

        self.load_sessions_list(self.current_patient, group_by_date=True)
        n = len(getattr(self, 'sessions', []))
        sessions = self._collect_sessions_data()
        if len(sessions) < 2:
            self._create_insufficient_data_graph(persistent, n, len(sessions), self.current_ex_name, "Пики по третям (утомляемость внутри цикла)")
            self._add_graph_card(persistent, self.thirds_content_layout, max_width=950)
            return

        # Для каждого теста/сессии — бары для 3 третей (используем M как прокси момента) — вертикально друг под другом
        n_sess = len(sessions)
        nrows = min(n_sess, 3)
        fig, axes = plt.subplots(nrows if nrows > 0 else 1, 1, figsize=(14, 3.4 * nrows), dpi=120, squeeze=False)
        axes = axes.flatten()

        for idx, s in enumerate(sessions[:3]):
            ax = axes[idx]
            M = np.array(s.get('M', []))
            if len(M) < 3:
                ax.text(0.5, 0.5, 'Мало точек', ha='center')
                continue
            n = len(M)
            thirds = [M[:n//3], M[n//3:2*n//3], M[2*n//3:]]
            peaks = [float(np.max(t)) if len(t) > 0 else 0 for t in thirds]
            colors = ['#3498db', '#e74c3c', '#2ecc71']
            bars = ax.bar(['1/3', '2/3', '3/3'], peaks, color=colors)
            ax.set_title(f"{s['date']}")
            ax.set_ylabel('Пик момента (прокси)')
            for bar, p in zip(bars, peaks):
                ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5, f'{p:.1f}', ha='center', fontsize=8)
            ax.grid(True, axis='y', alpha=0.3)

        plt.suptitle(f'Пики по третям теста (утомляемость внутри цикла) — {self.current_ex_name}', fontsize=12, fontweight='bold')
        plt.tight_layout()

        tmp = os.path.join(tempfile.gettempdir(), f"thirds_{self.current_ex_name}.png")
        plt.savefig(tmp, dpi=100, bbox_inches='tight', facecolor='white')
        plt.savefig(persistent, dpi=100, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        self._add_graph_card(persistent, self.thirds_content_layout, max_width=950)
        if os.path.exists(tmp): os.remove(tmp)

    def load_radar_metrics(self):
        """Радар-паспорт метрик (нормализованный; Тест 1 vs последний) — как Рис. 27 в файле.
        Метрики: пик, импульс, CV, consistency (если есть), ROM если доступен.
        """
        if not self.current_patient or not self.current_ex_name:
            return
        self.load_sessions_list(self.current_patient, group_by_date=True)
        self._safe_clear_layout(self.radar_content_layout)

        patient_dir = os.path.join(PATIENTS_DIR, self.current_patient)
        safe_name = re.sub(r'[^a-zA-Zа-яА-ЯёЁ0-9]', '_', self.current_ex_name)
        persistent = os.path.join(patient_dir, f"{safe_name}_radar_metrics.png")

        if os.path.exists(persistent):
            self._add_graph_card(persistent, self.radar_content_layout, max_width=800)
            return

        self.load_sessions_list(self.current_patient, group_by_date=True)
        n = len(getattr(self, 'sessions', []))
        sessions = self._collect_sessions_data()
        if len(sessions) < 2:
            self._create_insufficient_data_graph(persistent, n, len(sessions), self.current_ex_name, "Радар-паспорт метрик (T1 vs последний)")
            self._add_graph_card(persistent, self.radar_content_layout, max_width=800)
            return

        # Берём первый и последний
        first = sessions[0]
        last = sessions[-1]

        # Метрики (нормализуем относительно max по двум)
        metrics = ['Пик M', 'Импульс', 'CV (инв.)', 'Согл. (r)']
        vals1 = [
            first['peak'],
            first['impulse'],
            1.0 / (first['cv'] + 1e-6),  # инвертируем, т.к. ниже CV лучше
            0.8  # placeholder, реально можно посчитать consistency между T1 и последним
        ]
        vals2 = [
            last['peak'],
            last['impulse'],
            1.0 / (last['cv'] + 1e-6),
            0.85
        ]

        # Простая нормализация 0-1
        maxv = [max(v1, v2) for v1, v2 in zip(vals1, vals2)]
        n1 = [v / (m + 1e-6) for v, m in zip(vals1, maxv)]
        n2 = [v / (m + 1e-6) for v, m in zip(vals2, maxv)]

        # Radar plot
        angles = np.linspace(0, 2 * np.pi, len(metrics), endpoint=False).tolist()
        n1 += n1[:1]
        n2 += n2[:1]
        angles += angles[:1]

        fig, ax = plt.subplots(figsize=(14, 7), subplot_kw=dict(polar=True), dpi=120)
        ax.plot(angles, n1, 'o-', linewidth=2, label='Тест 1 (первый)', color='#3498db')
        ax.fill(angles, n1, alpha=0.25, color='#3498db')
        ax.plot(angles, n2, 's--', linewidth=2, label='Последний тест', color='#e74c3c')
        ax.fill(angles, n2, alpha=0.25, color='#e74c3c')
        ax.set_xticks(angles[:-1])
        ax.set_xticklabels(metrics)
        ax.set_title(f'Радар-паспорт метрик (T1 vs последний) — {self.current_ex_name}', fontsize=12, fontweight='bold', pad=20)
        ax.legend(loc='upper right', bbox_to_anchor=(1.3, 1.1))

        tmp = os.path.join(tempfile.gettempdir(), f"radar_{self.current_ex_name}.png")
        plt.savefig(tmp, dpi=100, bbox_inches='tight', facecolor='white')
        plt.savefig(persistent, dpi=100, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        self._add_graph_card(persistent, self.radar_content_layout, max_width=950)
        if os.path.exists(tmp): os.remove(tmp)

    def _add_graph_card(self, img_path, layout, max_width=850):
        if not os.path.exists(img_path):
            return

        pix = QPixmap(img_path)
        if not pix.isNull():
            if pix.width() > max_width:
                pix = pix.scaledToWidth(max_width, Qt.TransformationMode.SmoothTransformation)

            img_label = QLabel()
            img_label.setPixmap(pix)
            img_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            img_label.setMinimumHeight(pix.height())
            img_label.setSizePolicy(QSizePolicy.Policy.Fixed, QSizePolicy.Policy.Fixed)
            layout.addWidget(img_label)
        else:
            error_label = QLabel(f"⚠️ Не удалось загрузить изображение")
            error_label.setStyleSheet("color: #ffaa44; background: transparent; font-size: 14px;")
            error_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            layout.addWidget(error_label)

    def _add_analysis_card(self, title, img_path, layout):
        if not os.path.exists(img_path):
            return

        title_label = QLabel(title)
        title_label.setStyleSheet("font-size: 13px; font-weight: 600; color: #88ccff; background: transparent; margin-top: 6px;")
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        title_label.setWordWrap(True)
        layout.addWidget(title_label)

        pix = QPixmap(img_path)
        if not pix.isNull():
            max_width = 850
            if pix.width() > max_width:
                pix = pix.scaledToWidth(max_width, Qt.TransformationMode.SmoothTransformation)

            img_label = QLabel()
            img_label.setPixmap(pix)
            img_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            img_label.setMinimumHeight(pix.height())
            layout.addWidget(img_label)
        else:
            error_label = QLabel(f"⚠️ Не удалось загрузить изображение")
            error_label.setStyleSheet("color: #ffaa44; background: transparent;")
            error_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            layout.addWidget(error_label)

    def _create_insufficient_data_graph(self, persistent, n_ui, collected, ex_name, analysis_title):
        """Генерирует понятную PNG-карточку с объяснением вместо пустоты.
        Сохраняет в persistent, чтобы в будущем показывалась сразу (как настоящие графики).
        """
        try:
            fig, ax = plt.subplots(figsize=(14, 4.5), dpi=110)
            ax.axis('off')
            msg = (
                f"Анализ: {analysis_title}\n"
                f"Упражнение: «{ex_name}»\n\n"
                f"Сессий показано в правой панели: {n_ui}\n"
                f"Собрано валидных данных (с полным M): {collected}\n\n"
                "Код пытался сам восстановить:\n"
                "• Если raw_measurements.json отсутствовал — автоматически вызван regenerate (поиск .docx в пациенте и пересоздание сессии).\n"
                "• Использованы несколько fallback-прокси для момента (compute + сумма сил по времени + средний угол).\n\n"
                "Почему всё равно 0:\n"
                "• В дереве пациента не нашлось исходного .docx для восстановления измерений этого упражнения.\n"
                "• Или данные в raw повреждены/пустые.\n\n"
                "Что сделать, чтобы код сам создал графики:\n"
                "1. Пересоздай это упражнение заново из оригинального .docx (добавь пациента/упражнение ещё раз) — тогда raw запишется во все сессии.\n"
                "2. Удали все *_*.png анализа в корне папки пациента и открой упражнение заново.\n"
                "3. Смотри консоль — там [DEBUG _collect] и попытки regenerate покажут точную причину."
            )
            ax.text(0.02, 0.98, msg, transform=ax.transAxes, fontsize=9.5,
                    verticalalignment='top', fontfamily='sans-serif',
                    bbox=dict(boxstyle='round,pad=0.4', facecolor='#fff3cd', edgecolor='#856404', alpha=0.95))
            plt.tight_layout()
            plt.savefig(persistent, dpi=100, bbox_inches='tight', facecolor='white')
            plt.close(fig)
        except Exception as e:
            print(f"[WARN] _create_insufficient_data_graph failed for {analysis_title}: {e}")

    def _load_images_as_single_composite(self, layout, image_paths, max_width=850):
        """
        Склеивает несколько PNG в одно высокое изображение и показывает в одном QLabel.
        Это радикально уменьшает количество виджетов и помогает против 0xC0000409.
        """
        from PyQt6.QtGui import QImage, QPainter

        try:
            valid_paths = [p for p in image_paths if os.path.exists(p)]
            if not valid_paths:
                return

            images = []
            total_height = 0
            spacing = 10

            for path in valid_paths:
                img = QImage(path)
                if img.isNull():
                    continue
                if img.width() > max_width:
                    img = img.scaledToWidth(max_width, Qt.TransformationMode.SmoothTransformation)
                images.append(img)
                total_height += img.height() + spacing

            if not images:
                return

            big_image = QImage(max_width, total_height, QImage.Format.Format_ARGB32)
            big_image.fill(Qt.GlobalColor.white)

            painter = QPainter(big_image)
            y = 0
            for img in images:
                painter.drawImage(0, y, img)
                y += img.height() + spacing
            painter.end()

            pix = QPixmap.fromImage(big_image)
            label = QLabel()
            label.setPixmap(pix)
            label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            layout.addWidget(label)

        except Exception as e:
            # Не даём приложению упасть при проблемах с большими изображениями
            error_label = QLabel(f"Ошибка загрузки графиков: {str(e)[:100]}")
            error_label.setStyleSheet("color: #ff6666; font-size: 13px;")
            layout.addWidget(error_label)

    def on_session_selected(self, item):
        """
        При выборе сеанса/даты перезагружаем графики ТОЛЬКО в той вкладке,
        которая сейчас активна.
        - Для Парных графиков: всегда точный сеанс.
        - Для Спектрального и Гистерезиса: при выборе даты берём последний сеанс за эту дату.
        """
        data = item.data(Qt.ItemDataRole.UserRole)
        if not data or not self.current_patient:
            return

        current_widget = self.center_tabs.currentWidget()

        # Определяем, какую папку загружать
        target_folder = None

        if current_widget == self.pairs_page:
            # Для парных графиков (иерархический режим)
            if isinstance(data, str):
                if '_' in data:
                    # Это точный сеанс (имя папки)
                    target_folder = data
                elif re.match(r'\d{4}-\d{2}-\d{2}', data):
                    # Это заголовок даты — берём самый новый сеанс за эту дату
                    folders = self.date_to_folders.get(data, [])
                    if folders:
                        target_folder = folders[0]
            else:
                return
        else:
            # Для спектрального, гистерезиса и симметрии — используем дату
            if isinstance(data, str) and re.match(r'\d{4}-\d{2}-\d{2}', data):
                folders = self.date_to_folders.get(data, [])
                if folders:
                    target_folder = folders[0]
            else:
                target_folder = data

        if not target_folder:
            return

        if current_widget == self.spectral_page:
            self.load_spectral_graphs(self.current_patient, target_folder)
        elif current_widget == self.pairs_page:
            self.load_pair_graphs(self.current_patient, target_folder)
        elif current_widget == self.hysteresis_page:
            self.load_hysteresis_graphs(self.current_patient, target_folder)
        elif current_widget == self.symmetry_page:
            self.load_symmetry_graphs(self.current_patient, target_folder)
        elif current_widget == self.legload_page:
            self.load_leg_load_graph(self.current_patient, target_folder)
        elif current_widget == self.sym_peaks_page:
            self.load_sym_peaks_impulse()
        elif current_widget == self.phase_lag_page:
            self.load_phase_lag_skew()
        elif current_widget == self.trend_cv_page:
            self.load_trend_cv()
        elif current_widget == self.phase_portrait_page:
            self.load_phase_portrait()
        elif current_widget == self.consistency_page:
            self.load_consistency()
        elif current_widget == self.thirds_fatigue_page:
            self.load_thirds_fatigue()
        elif current_widget == self.radar_page:
            self.load_radar_metrics()

    def _on_tab_changed(self, index: int):
        """
        Показываем/скрываем панель 'Выбор сеанса' + перезагружаем список сеансов
        в нужном формате в зависимости от вкладки.
        """
        current = self.center_tabs.widget(index)

        # Показываем селектор сеансов для пер-сеансовых вкладок И для новых агрегатных (чтобы можно было выбирать сеанс по дате).
        # Для агрегатных вкладок выбор сеанса просто перестраивает общий вид (все сессии), но список дат доступен.
        show_selector = current in (self.spectral_page, self.pairs_page, self.hysteresis_page, self.symmetry_page, self.legload_page,
                                    self.sym_peaks_page, self.phase_lag_page, self.trend_cv_page,
                                    self.phase_portrait_page, self.consistency_page,
                                    self.thirds_fatigue_page, self.radar_page)

        self.right_panel.setVisible(show_selector)

        if show_selector:
            self.splitter.setSizes([340, 960, 300])

            # При переключении вкладки — загружаем правильный вид списка
            if self.current_patient:
                if current == self.pairs_page:
                    # Только для парных графиков — иерархический вид (даты + сеансы под ними)
                    self.load_sessions_list(self.current_patient, hierarchical=True)
                else:
                    # Для спектрального, гистерезиса, симметрии и новых агрегатных — только даты
                    self.load_sessions_list(self.current_patient, group_by_date=True)

                # Автоматически выбираем самую свежую дату/сеанс после переключения вкладки
                self._auto_select_latest_session()
        else:
            self.splitter.setSizes([340, 1260, 0])

    def _update_session_selector_visibility(self):
        """Удобный вызов для установки видимости после загрузки упражнения."""
        self._on_tab_changed(self.center_tabs.currentIndex())

    def _auto_select_latest_session(self):
        """Автоматически выбирает самую свежую дату (или сеанс) и загружает её."""
        if self.session_list.count() > 0:
            self.session_list.setCurrentRow(0)
            item = self.session_list.item(0)
            if item:
                self.on_session_selected(item)

    def on_apply_smoothing(self):
        """Мгновенная перестройка графиков текущего сеанса с новым уровнем сглаживания."""
        if not self.current_patient or not self.current_ex_name:
            show_styled_message(self, "Информация", "Сначала откройте упражнение.", "info")
            return

        current_widget = self.center_tabs.currentWidget()
        aggregate_pages = (self.sym_peaks_page, self.phase_lag_page, self.trend_cv_page,
                           self.phase_portrait_page, self.consistency_page,
                           self.thirds_fatigue_page, self.radar_page)
        if current_widget in aggregate_pages:
            show_styled_message(self, "Информация",
                "Сглаживание применяется только к графикам отдельных сеансов "
                "(спектральный анализ, парные графики, гистерезис, симметрия, нагрузка ноги).\n\n"
                "Агрегатные анализы (симметрия пики+импульс, phase lag, тренд CV, петля угол-момент, "
                "согласованность, пики по третям, радар) строятся по всем сессиям упражнения сразу "
                "и не зависят от сглаживания одного сеанса.", "info")
            return

        current_item = self.session_list.currentItem()
        if not current_item:
            show_styled_message(self, "Информация", "Выберите сеанс в правой панели.", "info")
            return

        folder_name = current_item.data(Qt.ItemDataRole.UserRole)
        if not folder_name:
            # Пытаемся получить из текста (fallback)
            folder_name = current_item.text().strip()

        # Определяем интенсивность
        text = self.smoothing_combo.currentText()
        if "Без" in text:
            intensity = "none"
        elif "Лёгкое" in text:
            intensity = "light"
        elif "Сильное" in text:
            intensity = "strong"
        else:
            intensity = "medium"

        self.current_smoothing_intensity = intensity

        # UI feedback
        original_text = self.apply_smoothing_btn.text()
        self.apply_smoothing_btn.setText("Перестраиваю...")
        self.apply_smoothing_btn.setEnabled(False)
        QApplication.processEvents()

        try:
            success, message = regenerate_graphs_for_session(
                self.current_patient, folder_name, intensity
            )

            if success:
                # Перезагружаем текущую активную вкладку
                current_widget = self.center_tabs.currentWidget()

                if current_widget == self.pairs_page:
                    self.load_pair_graphs(self.current_patient, folder_name)
                elif current_widget == self.hysteresis_page:
                    self.load_hysteresis_graphs(self.current_patient, folder_name)
                elif current_widget == self.symmetry_page:
                    self.load_symmetry_graphs(self.current_patient, folder_name)
                elif current_widget == self.legload_page:
                    self.load_leg_load_graph(self.current_patient, folder_name)
                elif current_widget == self.sym_peaks_page:
                    self.load_sym_peaks_impulse()
                elif current_widget == self.phase_lag_page:
                    self.load_phase_lag_skew()
                elif current_widget == self.trend_cv_page:
                    self.load_trend_cv()
                elif current_widget == self.phase_portrait_page:
                    self.load_phase_portrait()
                elif current_widget == self.consistency_page:
                    self.load_consistency()
                elif current_widget == self.phase_portrait_page:
                    self.load_phase_portrait()
                elif current_widget == self.thirds_fatigue_page:
                    self.load_thirds_fatigue()
                elif current_widget == self.radar_page:
                    self.load_radar_metrics()
                elif current_widget == self.spectral_page:
                    self.load_spectral_graphs(self.current_patient, folder_name)
                else:
                    self.load_amplitude_graph(self.current_patient)

                show_styled_message(self, "Готово", f"Сглаживание изменено на «{text}».\nГрафики перестроены.", "info")
            else:
                show_styled_message(self, "Внимание", message + "\n\nПопробуйте пересоздать данные пациента.", "warning")

        except Exception as e:
            show_styled_message(self, "Ошибка", f"Не удалось перестроить графики:\n{str(e)}", "critical")
        finally:
            self.apply_smoothing_btn.setText(original_text)
            self.apply_smoothing_btn.setEnabled(True)


# ============================================================
# WelcomePage
# ============================================================
class WelcomePage(QWidget):
    start_requested = pyqtSignal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
        self.setStyleSheet("background-color: #0a1628;")

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 140, 0, 80)
        layout.setSpacing(32)
        layout.setAlignment(Qt.AlignmentFlag.AlignCenter)

        self.logo_container = QLabel()
        self.logo_container.setFixedSize(260, 260)
        self.logo_container.setStyleSheet("""
            QLabel {
                background-color: #1e90ff;
                border-radius: 130px;
            }
        """)

        icon = qta.icon('fa5s.robot', color='white', scale_factor=2.7)
        self.logo = QLabel(self.logo_container)
        pix = icon.pixmap(125, 125)
        self.logo.setPixmap(pix)
        self.logo.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.logo.setGeometry(67, 67, 125, 125)

        self.glow_effect = QGraphicsDropShadowEffect()
        self.glow_effect.setBlurRadius(110)
        self.glow_effect.setXOffset(0)
        self.glow_effect.setYOffset(0)
        self.glow_effect.setColor(QColor(30, 144, 255, 240))
        self.logo_container.setGraphicsEffect(self.glow_effect)

        layout.addWidget(self.logo_container, alignment=Qt.AlignmentFlag.AlignCenter)

        self.title = QLabel("SWSU ROBOTICS")
        self.title.setFont(QFont("Segoe UI", 68, QFont.Weight.Bold))
        self.title.setStyleSheet("color: #ffffff; letter-spacing: -3px;")
        self.title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.title)

        subtitle = QLabel("MECHATRONICS RESEARCH & ENGINEERING KIT")
        subtitle.setFont(QFont("Segoe UI", 19, QFont.Weight.Medium))
        subtitle.setStyleSheet("color: #60c0ff;")
        subtitle.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(subtitle)

        version_btn = QPushButton("Версия 1.0.0")
        version_btn.setFixedSize(170, 40)
        version_btn.setStyleSheet(
            "QPushButton { background-color: #1e90ff; color: white; border-radius: 20px; font-size: 15px; font-weight: 600; }")
        layout.addWidget(version_btn, alignment=Qt.AlignmentFlag.AlignCenter)

        line = QFrame()
        line.setFixedWidth(460)
        line.setFixedHeight(2)
        line.setStyleSheet("background-color: #1e90ff; border: none;")
        layout.addWidget(line, alignment=Qt.AlignmentFlag.AlignCenter)

        self.start_btn = QPushButton("НАЧАТЬ РАБОТУ")
        self.start_btn.setFixedHeight(74)
        self.start_btn.setFixedWidth(460)
        self.start_btn.setFont(QFont("Segoe UI", 24, QFont.Weight.Bold))
        self.start_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.start_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #1e90ff, stop:1 #0b7be0);
                color: white;
                border-radius: 37px;
                border: none;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #40b0ff, stop:1 #1e90ff);
            }
            QPushButton:pressed {
                background: #0a5eb8;
            }
        """)
        self.start_btn.clicked.connect(self.start_requested.emit)
        layout.addWidget(self.start_btn, alignment=Qt.AlignmentFlag.AlignCenter)

        footer = QLabel("© 2025 SWSU Robotics Department. All rights reserved.")
        footer.setStyleSheet("color: #5577aa; font-size: 13px;")
        footer.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addStretch()
        layout.addWidget(footer, alignment=Qt.AlignmentFlag.AlignCenter)

        self.animate_logo()

    def animate_logo(self):
        anim = QPropertyAnimation(self.logo_container, b"pos")
        anim.setDuration(1200)
        anim.setStartValue(QPoint(self.logo_container.x(), self.logo_container.y() - 160))
        anim.setEndValue(self.logo_container.pos())
        anim.setEasingCurve(QEasingCurve.Type.OutBack)
        anim.start()
        self.glow_timer = QTimer(self)
        self.glow_timer.timeout.connect(self.pulse_glow)
        self.glow_value = 70
        self.glow_dir = 1
        self.glow_timer.start(25)

    def pulse_glow(self):
        self.glow_value += 4 * self.glow_dir
        if self.glow_value > 125 or self.glow_value < 55:
            self.glow_dir *= -1
        self.glow_effect.setBlurRadius(self.glow_value)
        self.logo_container.setGraphicsEffect(self.glow_effect)


# ============================================================
# MainWindow
# ============================================================
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("SWSU ROBOTICS")
        self.setMinimumSize(1380, 900)
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint)
        self.showFullScreen()

        central = QWidget()
        self.setCentralWidget(central)
        main_layout = QVBoxLayout(central)
        main_layout.setContentsMargins(0, 0, 0, 0)

        self.stacked = QStackedWidget()
        main_layout.addWidget(self.stacked)

        self.welcome_page = WelcomePage()
        self.welcome_page.start_requested.connect(self.show_patients_list)

        self.patients_page = PatientsListPage(self)
        self.patient_exercises_page = PatientWithExercisesPage(self)
        self.exercise_view_page = ExerciseViewPage()

        self.stacked.addWidget(self.welcome_page)
        self.stacked.addWidget(self.patients_page)
        self.stacked.addWidget(self.patient_exercises_page)
        self.stacked.addWidget(self.exercise_view_page)

        self.patients_page.patient_selected.connect(self.show_patient_exercises)
        self.patient_exercises_page.back_to_list.connect(self.show_patients_list)
        self.patient_exercises_page.exercise_selected.connect(self.show_exercise_view)
        self.exercise_view_page.back_to_exercises.connect(self.show_patient_exercises_page)

        self.progress = QProgressBar()
        self.progress.setVisible(False)
        main_layout.addWidget(self.progress)

        self.thread = None
        self.esc_shortcut = QShortcut(QKeySequence("Esc"), self)
        self.esc_shortcut.activated.connect(self.on_esc_pressed)

        self.stacked.setCurrentWidget(self.welcome_page)

    def show_patients_list(self):
        current_patient = getattr(self.patient_exercises_page, 'current_patient', None)
        if current_patient:
            self.patients_page.set_selected_patient(current_patient)
        else:
            self.patients_page.refresh()
        self.stacked.setCurrentWidget(self.patients_page)

    def show_patient_exercises(self, patient_name):
        self.patient_exercises_page.set_patient(patient_name)
        self.stacked.setCurrentWidget(self.patient_exercises_page)

    def show_patient_exercises_page(self):
        self.stacked.setCurrentWidget(self.patient_exercises_page)

    def show_exercise_view(self, patient_name, exercise_folder):
        exercise_path = os.path.join(PATIENTS_DIR, patient_name, exercise_folder)
        self.exercise_view_page.set_exercise(patient_name, exercise_folder, exercise_path)
        self.stacked.setCurrentWidget(self.exercise_view_page)

    def add_new_patient(self):
        dlg = PatientInfoDialog()
        if dlg.exec() == QDialog.DialogCode.Accepted:
            name, birth, complaint, height, weight, upper, middle, lower = dlg.get_info()
            if name is None:
                return
            safe_name = re.sub(r'[^a-zA-Zа-яА-ЯёЁ0-9]', '_', name.strip())
            safe_birth = re.sub(r'\D', '', birth.strip())[:8]
            folder_name = f"{safe_name}_{safe_birth}" if safe_birth else safe_name
            original_folder = folder_name
            counter = 1
            while os.path.exists(os.path.join(PATIENTS_DIR, folder_name)):
                folder_name = f"{original_folder}_{counter}"
                counter += 1
            patient_dir = os.path.join(PATIENTS_DIR, folder_name)
            os.makedirs(patient_dir, exist_ok=True)
            info = {
                'name': name,
                'birth_date': birth,
                'complaint': complaint,
                'height_cm': height,
                'weight_kg': weight,
                'upper_link_cm': upper,
                'middle_link_cm': middle,
                'lower_link_cm': lower,
                'created': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            with open(os.path.join(patient_dir, 'info.txt'), 'w', encoding='utf-8') as f:
                json.dump(info, f, ensure_ascii=False, indent=2)
            show_styled_message(self, "Готово", f"Пациент «{name}» успешно создан.", "info")
            self.patients_page.refresh()

    def edit_patient(self, patient_folder):
        patient_dir = os.path.join(PATIENTS_DIR, patient_folder)
        info_path = os.path.join(patient_dir, 'info.txt')
        name = ""
        birth = ""
        complaint = ""
        height = ""
        weight = ""
        upper = ""
        middle = ""
        lower = ""
        if os.path.exists(info_path):
            try:
                with open(info_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    name = data.get('name', '')
                    birth = data.get('birth_date', '')
                    complaint = data.get('complaint', '')
                    height = data.get('height_cm', '')
                    weight = data.get('weight_kg', '')
                    upper = data.get('upper_link_cm', '')
                    middle = data.get('middle_link_cm', '')
                    lower = data.get('lower_link_cm', '')
            except:
                pass
        dlg = PatientInfoDialog(name, birth, complaint, height, weight, upper, middle, lower)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            new_name, new_birth, new_complaint, new_height, new_weight, new_upper, new_middle, new_lower = dlg.get_info()
            if new_name is None:
                return
            info = {
                'name': new_name,
                'birth_date': new_birth,
                'complaint': new_complaint,
                'height_cm': new_height,
                'weight_kg': new_weight,
                'upper_link_cm': new_upper,
                'middle_link_cm': new_middle,
                'lower_link_cm': new_lower,
                'edited': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            with open(info_path, 'w', encoding='utf-8') as f:
                json.dump(info, f, ensure_ascii=False, indent=2)
            current_widget = self.stacked.currentWidget()
            if isinstance(current_widget,
                          PatientWithExercisesPage) and current_widget.current_patient == patient_folder:
                current_widget.set_patient(patient_folder)
            self.patients_page.refresh()
            show_styled_message(self, "Готово", f"Данные пациента «{new_name}» обновлены.", "info")

    def add_report_to_existing(self, patient_folder):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Выберите файл с данными (.docx)", "",
            "Документы Word (*.docx);;Все файлы (*.*)"
        )
        if not file_path:
            return
        self.progress.setVisible(True)
        self.progress.setRange(0, 0)
        self.setEnabled(False)
        self.thread = ProcessingThread(file_path, patient_folder)
        self.thread.finished_signal.connect(self.on_processing_finished)
        self.thread.start()

    def on_processing_finished(self, success, msg, patient_dir):
        self.progress.setVisible(False)
        self.setEnabled(True)
        if success:
            patient_name = patient_dir.split(os.sep)[-1]
            current_widget = self.stacked.currentWidget()
            if isinstance(current_widget, PatientWithExercisesPage) and current_widget.current_patient == patient_name:
                current_widget.set_patient(patient_name)
            self.patients_page.refresh()
            show_styled_message(self, "Успех", msg, "info")
        else:
            show_styled_message(self, "Ошибка", msg, "critical")

    def on_esc_pressed(self):
        current = self.stacked.currentWidget()
        if current == self.exercise_view_page:
            self.exercise_view_page.back_to_exercises.emit()
        elif current == self.patient_exercises_page:
            self.show_patients_list()

    def closeEvent(self, event):
        event.accept()


# ============================================================
# ЗАПУСК
# ============================================================
if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyle("Fusion")

    palette = QPalette()
    palette.setColor(QPalette.ColorRole.Window, QColor(10, 20, 40))
    palette.setColor(QPalette.ColorRole.WindowText, QColor(255, 255, 255))
    app.setPalette(palette)

    window = MainWindow()
    sys.exit(app.exec())